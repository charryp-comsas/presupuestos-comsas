"""
App de Presupuestos COMSAS - Fase 1 + Fase 2 + Fase 3 + Fase 4 + Fase 5
------------------------------------------------------------
Fase 1: crear un presupuesto, buscar y agregar items del catalogo
real (ya cargado en Supabase), editar/eliminar y ver el total.
Fase 2: cuadrillas y mano de obra -- crear cuadrillas nuevas o editar
la composicion (cargo y cantidad) de las existentes (sub-pestaña
"Cuadrillas" en Mantenimiento de precios, equivalente a
CREAR_CUADRILLAS.txt), y asignar/reasignar la cuadrilla de un APU
desde el editor de receta (equivalente a ASIGNAR_CUADRILLAS.txt /
CAMBIAR_CUADRILLA.txt). El costo/dia de cada cuadrilla divide el
costo del MAESTRO entre frentes_maestro, igual que la hoja
CUADRILLAS del Excel maestro.
Fase 3: calcular AIU sobre el costo directo, generar la propuesta
en Word a partir de la plantilla real, y generar un Excel de
respaldo por presupuesto.
Fase 4: mantenimiento de precios del catalogo -- equivalente a
MANTENIMIENTO_PRECIOS.txt (semaforo de vigencia + escalamiento por
indice ICOCED + puente de precios desde facturas, todo con bitacora)
y RECOSTEAR_MANO_OBRA.txt (cambio de tarifas de personal y recalculo
de mano_obra por cuadrilla). El semaforo, el ICOCED y el puente de
facturas viven a nivel de INSUMO (tablas insumos / apu_insumos /
apu_materiales_fijos, extraidas de la hoja MATERIALES y de las
formulas de cada hoja de APU en BASE APU COSTOS 2026.xlsm -- ver
fase4_precios/schema_fase4.sql), y cualquier cambio de precio de un
insumo recalcula solo el 'materiales' de catalogo_apu para los APUs
que lo usan. Las lineas de materiales que en el Excel no enlazan
directo a MATERIALES (ej. formulas de concreto que vienen de
SUB-ACTIVIDADES) quedan en apu_materiales_fijos: se suman al costo
pero no se recalculan automaticamente todavia.

Fase 5: login con Supabase Auth y dos roles -- 'administrador' (todos
los permisos) y 'operador' (puede usar todas las pestañas y trabajar
presupuestos libremente, pero no puede guardar cambios que alteren el
catalogo en general: precios de insumos, cuadrillas, tarifas de
personal, ni la receta de un APU). La restriccion real vive en las
politicas RLS (ver fase5_login/schema_fase5.sql); la interfaz solo
deshabilita esos botones para que el operador no se lleve una sorpresa.

Esto sigue siendo un PROTOTIPO DE VALIDACION frente al Excel maestro,
pero ya se puede correr en mas de una maquina a la vez contra el mismo
Supabase (cada quien con su propio login).
"""

import re
import unicodedata
from datetime import date, datetime
from io import BytesIO
from pathlib import Path

import streamlit as st
from supabase import create_client

import base64

ASSETS_DIR = Path(__file__).parent / "assets"
LOGO_PATH = ASSETS_DIR / "logo_comsas.png"
LOGO_ICON_PATH = ASSETS_DIR / "logo_icon.png"
LOGO_WATERMARK_PATH = ASSETS_DIR / "logo_watermark.png"

st.set_page_config(
    page_title="Presupuestos COMSAS",
    page_icon=str(LOGO_ICON_PATH) if LOGO_ICON_PATH.exists() else "🏗️",
    layout="wide",
)


def _b64_file(path: Path) -> str | None:
    if not path.exists():
        return None
    return base64.b64encode(path.read_bytes()).decode()


# ---------------------------------------------------------------------
# Branding COMSAS: logo en el sidebar (tal cual, sin fondos forzados) +
# marca de agua muy sutil del logo en el contenido + acento dorado en
# el titulo y en la pestaña activa. Nada de bloques de color solido.
# ---------------------------------------------------------------------
if LOGO_PATH.exists():
    st.logo(str(LOGO_PATH), size="large")

_watermark_b64 = _b64_file(LOGO_WATERMARK_PATH)

_watermark_css = (
    f"""
    /* Marca de agua dentro de cada pestaña (hoja), no solo en el fondo general */
    .stTabs [data-baseweb="tab-panel"] {{
        background-image: url("data:image/png;base64,{_watermark_b64}");
        background-repeat: no-repeat;
        background-position: bottom 16px right 24px;
        background-size: 200px auto;
    }}
    """
    if _watermark_b64
    else ""
)

st.markdown(
    f"""
    <style>
    :root {{
        --comsas-gold: #C6A15B;
        --comsas-gold-dark: #9C7A2E;
    }}

    {_watermark_css}

    /* Acento dorado sutil en el borde del sidebar, sin cambiar su color base */
    section[data-testid="stSidebar"] {{
        border-right: 1px solid rgba(198, 161, 91, 0.45);
    }}
    div[data-testid="stSidebarHeader"] {{
        border-bottom: 1px solid rgba(198, 161, 91, 0.35);
        padding-bottom: 12px;
        margin-bottom: 8px;
    }}

    /* Barra de pestañas: cajas reales con color, no solo texto subrayado.
    Confirmado por inspeccion en la app publicada (2026-08-01): cada pestaña
    es un <div data-testid="stTab" aria-selected="..."> -- NO un
    <button data-baseweb="tab"> como en versiones viejas de Streamlit. Se
    dejan ambos selectores por compatibilidad, pero el que realmente
    aplica ahora es data-testid="stTab". */
    .stTabs [data-baseweb="tab-list"],
    div[data-testid="stTabs"] [role="tablist"] {{
        gap: 6px;
        background-color: transparent;
        border-bottom: 2px solid rgba(198, 161, 91, 0.55);
        padding-left: 2px;
    }}
    .stTabs [data-baseweb="tab"],
    div[data-testid="stTab"] {{
        background-color: #F5F1E8 !important;
        border: 1px solid rgba(198, 161, 91, 0.4) !important;
        border-bottom: none !important;
        border-radius: 10px 10px 0 0 !important;
        padding: 10px 22px;
        color: #4a4740 !important;
        font-weight: 500;
        transition: background-color 0.15s ease, color 0.15s ease;
    }}
    .stTabs [data-baseweb="tab"]:hover,
    div[data-testid="stTab"]:hover {{
        background-color: #EFE3C4 !important;
        color: #161412 !important;
    }}
    .stTabs [aria-selected="true"],
    div[data-testid="stTab"][aria-selected="true"] {{
        background-color: #161412 !important;
        border-color: #161412 !important;
        color: var(--comsas-gold) !important;
        font-weight: 600;
    }}
    .stTabs [aria-selected="true"] p,
    .stTabs [aria-selected="true"] *,
    div[data-testid="stTab"][aria-selected="true"] p,
    div[data-testid="stTab"][aria-selected="true"] * {{
        color: var(--comsas-gold) !important;
    }}
    /* la barra deslizante de subrayado ya no hace falta con pestañas en caja */
    .stTabs [data-baseweb="tab-highlight"],
    div[data-testid="stTabs"] [data-baseweb="tab-highlight"] {{
        display: none;
    }}
    .stTabs [data-baseweb="tab-border"],
    div[data-testid="stTabs"] [data-baseweb="tab-border"] {{
        display: none;
    }}
    .stTabs [data-baseweb="tab-panel"],
    div[data-testid="stTabs"] [data-testid="stTabsContent"] {{
        border: 1px solid rgba(198, 161, 91, 0.4);
        border-top: none;
        border-radius: 0 0 10px 10px;
        padding: 20px;
    }}

    /* Tarjetas (st.container(border=True)): hairline dorado en vez del gris por defecto */
    div[data-testid="stVerticalBlockBorderWrapper"] {{
        border-color: rgba(198, 161, 91, 0.4) !important;
    }}

    /* Tablas (st.dataframe): marco dorado -- el header interno es un
    canvas y no se puede pintar celda por celda desde CSS, pero el
    theme.secondaryBackgroundColor del config.toml ya lo deja en tono
    crema acorde a la paleta. */
    div[data-testid="stDataFrame"] {{
        border: 1px solid rgba(198, 161, 91, 0.4);
        border-radius: 8px;
        overflow: hidden;
    }}

    /* Botones primarios: negro con borde y texto dorado, se invierte al pasar el mouse */
    button[kind="primary"], button[kind="primaryFormSubmit"] {{
        background-color: #161412 !important;
        border: 1px solid var(--comsas-gold) !important;
        color: var(--comsas-gold) !important;
        box-shadow: none !important;
    }}
    button[kind="primary"]:hover, button[kind="primaryFormSubmit"]:hover {{
        background-color: var(--comsas-gold) !important;
        color: #161412 !important;
        border-color: var(--comsas-gold) !important;
    }}

    /* Botones secundarios: solo el borde se dora al pasar el mouse */
    button[kind="secondary"]:hover, button[kind="secondaryFormSubmit"]:hover {{
        border-color: var(--comsas-gold) !important;
        color: var(--comsas-gold-dark) !important;
    }}
    </style>
    """,
    unsafe_allow_html=True,
)


# ---------------------------------------------------------------------
# Conexion a Supabase (Fase 5: un cliente POR SESION, no compartido)
# ---------------------------------------------------------------------
def get_client():
    """OJO: antes esto usaba @st.cache_resource, que crea UNA sola
    instancia compartida por todo el proceso -- bien mientras solo tu
    corrias la app, pero con login cada usuario tiene su propia sesion
    de Supabase Auth (token de acceso), y compartir el cliente mezclaria
    la sesion de un usuario con la de otro. Por eso ahora vive en
    st.session_state, que Streamlit ya aisla por usuario/pestaña."""
    if "sb_client" not in st.session_state:
        url = st.secrets["SUPABASE_URL"]
        key = st.secrets["SUPABASE_ANON_KEY"]
        st.session_state.sb_client = create_client(url, key)
    return st.session_state.sb_client


sb = get_client()

# URL publica de la app (necesaria para armar el link de recuperacion de
# contraseña que Supabase manda por correo -- debe coincidir con una de
# las "Redirect URLs" configuradas en Supabase, ver fase5_login/LEEME.txt).
APP_URL = "https://presupuesto-udgpw8evderlmem9wxitkd.streamlit.app/"


# ---------------------------------------------------------------------
# Login (Fase 5)
# ---------------------------------------------------------------------
def usuario_actual():
    return st.session_state.get("usuario")


def es_admin():
    u = usuario_actual()
    return bool(u) and u.get("rol") == "administrador"


def es_gestor():
    """Administrador o Cotizador: ambos pueden crear/editar catalogo, APUs,
    cuadrillas, tarifas de personal y la receta de un APU. Las unicas dos
    acciones que quedan reservadas solo a 'administrador' son escalar
    precios (ICOCED) y aplicar precios desde facturas -- esas siguen
    usando es_admin() directamente."""
    u = usuario_actual()
    return bool(u) and u.get("rol") in ("administrador", "cotizador")


def _usuario_bitacora():
    """Correo del usuario conectado, para dejar registrado en la bitacora
    de cambios de precios/APU quien hizo cada modificacion. 'sistema' es
    un respaldo por si algo llama esto sin sesion activa (no deberia
    pasar, la app exige login antes de llegar a estas pantallas)."""
    u = usuario_actual()
    return (u or {}).get("email") or "sistema"


def cerrar_sesion():
    try:
        sb.auth.sign_out()
    except Exception:
        pass
    for clave in ("usuario", "sb_client"):
        st.session_state.pop(clave, None)
    st.rerun()


def pantalla_login():
    _, col_form, _ = st.columns([1, 1.1, 1])
    with col_form:
        if LOGO_PATH.exists():
            st.image(str(LOGO_PATH), width=220)
        st.markdown(
            '<h2 style="text-align:center;margin-top:0">'
            'Presupuestos <span style="color:#9C7A2E">COMSAS</span></h2>',
            unsafe_allow_html=True,
        )
        st.markdown(
            "<p style='text-align:center;color:var(--comsas-gold-dark,#9C7A2E)'>"
            "Inicia sesion para continuar.</p>",
            unsafe_allow_html=True,
        )
        with st.form("form_login"):
            email = st.text_input("Correo")
            password = st.text_input("Contraseña", type="password")
            enviado = st.form_submit_button("Entrar", type="primary")

        with st.expander("¿Olvidaste tu contraseña?"):
            email_recuperar = st.text_input(
                "Escribe el correo con el que entras a la app", key="email_recuperar"
            )
            if st.button("Enviar link de recuperacion"):
                if not email_recuperar:
                    st.error("Escribe tu correo.")
                else:
                    try:
                        sb.auth.reset_password_for_email(
                            email_recuperar, {"redirect_to": APP_URL}
                        )
                    except Exception:
                        pass
                    # Mensaje generico a proposito (no confirma ni niega si el
                    # correo existe -- evita que alguien use este formulario
                    # para averiguar que correos estan registrados).
                    st.success(
                        "Si ese correo esta registrado, te llega un link para "
                        "poner una contraseña nueva (revisa tambien spam). "
                        "El link vale por un tiempo limitado."
                    )
    if enviado:
        if not email or not password:
            st.error("Escribe correo y contraseña.")
        else:
            try:
                auth_resp = sb.auth.sign_in_with_password(
                    {"email": email, "password": password}
                )
            except Exception:
                st.error("Correo o contraseña incorrectos.")
                auth_resp = None
            auth_user_id = (
                auth_resp.user.id if auth_resp is not None and auth_resp.user else None
            )
            if auth_user_id:
                filas_perfil = (
                    sb.table("usuarios")
                    .select("nombre, email, rol, activo")
                    .eq("auth_user_id", auth_user_id)
                    .limit(1)
                    .execute()
                    .data
                )
                perfil = filas_perfil[0] if filas_perfil else None
                if not perfil or not perfil.get("activo"):
                    st.error(
                        "Tu cuenta no esta registrada (o esta inactiva) en la app. "
                        "Pide a un administrador que te de de alta en la tabla 'usuarios'."
                    )
                    try:
                        sb.auth.sign_out()
                    except Exception:
                        pass
                else:
                    st.session_state.usuario = perfil
                    st.rerun()
    st.stop()


def _limpiar_query_params():
    for clave in list(st.query_params.keys()):
        del st.query_params[clave]


def pantalla_restablecer_password(token_hash):
    _, col_form, _ = st.columns([1, 1.1, 1])
    with col_form:
        if LOGO_PATH.exists():
            st.image(str(LOGO_PATH), width=220)
        st.markdown(
            '<h2 style="text-align:center;margin-top:0">Poner una contraseña nueva</h2>',
            unsafe_allow_html=True,
        )

        if "recovery_verificado" not in st.session_state:
            try:
                sb.auth.verify_otp({"token_hash": token_hash, "type": "recovery"})
                st.session_state.recovery_verificado = True
            except Exception:
                st.error(
                    "Este link ya no es valido (puede que haya expirado o que ya "
                    "se haya usado). Vuelve a la pantalla de inicio y pide uno nuevo "
                    "desde '¿Olvidaste tu contraseña?'."
                )
                if st.button("Ir a iniciar sesion"):
                    _limpiar_query_params()
                    st.rerun()
                st.stop()

        with st.form("form_nueva_password"):
            nueva = st.text_input("Contraseña nueva", type="password")
            confirmar = st.text_input("Confirma la contraseña nueva", type="password")
            enviado = st.form_submit_button("Guardar contraseña nueva", type="primary")

    if enviado:
        if not nueva or len(nueva) < 6:
            st.error("La contraseña debe tener al menos 6 caracteres.")
        elif nueva != confirmar:
            st.error("Las dos contraseñas no coinciden.")
        else:
            try:
                sb.auth.update_user({"password": nueva})
            except Exception as e:
                st.error(f"No se pudo guardar la contraseña nueva: {e}")
            else:
                st.session_state.pop("recovery_verificado", None)
                try:
                    sb.auth.sign_out()
                except Exception:
                    pass
                _limpiar_query_params()
                st.success("Contraseña actualizada. Ya puedes iniciar sesion con ella.")
                if st.button("Ir a iniciar sesion"):
                    st.rerun()
                st.stop()
    st.stop()


_qp = st.query_params
if _qp.get("type") == "recovery" and _qp.get("token_hash") and not usuario_actual():
    pantalla_restablecer_password(_qp.get("token_hash"))

if not usuario_actual():
    pantalla_login()

with st.sidebar:
    _rol_actual = usuario_actual()["rol"]
    st.markdown(
        f"""
        <div style="margin-bottom:10px">
            <span style="color:var(--text-color, inherit)">Conectado como: <b>{usuario_actual()['nombre']}</b></span><br>
            <span style="
                display:inline-block;margin-top:4px;padding:2px 10px;
                border-radius:999px;background-color:#161412;
                color:#C6A15B;font-size:12px;font-weight:600;letter-spacing:0.3px;
            ">{_rol_actual.upper()}</span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if st.button("Cerrar sesion"):
        cerrar_sesion()
    if _rol_actual == "operador":
        st.caption(
            "Modo operador: puedes usar todas las pestañas y trabajar presupuestos "
            "libremente. Los cambios que alteran el catalogo en general (precios de "
            "insumos, cuadrillas, tarifas de personal, receta de un APU) estan "
            "reservados a un administrador o cotizador -- veras esos botones "
            "deshabilitados."
        )
    elif _rol_actual == "cotizador":
        st.caption(
            "Modo cotizador: mismos permisos que administrador (catalogo, APUs, "
            "cuadrillas, tarifas, presupuestos), excepto escalar precios (ICOCED) "
            "y aplicar precios desde facturas -- esos dos botones quedan "
            "reservados a un administrador."
        )
    st.divider()


def money(v):
    if v is None:
        return "$ 0"
    return f"$ {v:,.0f}".replace(",", ".")


def sanitizar_nombre_archivo(nombre):
    """Quita caracteres invalidos para un nombre de archivo en Windows."""
    invalidos = '\\/:*?"<>|'
    for ch in invalidos:
        nombre = nombre.replace(ch, "-")
    return nombre.strip()


def elegir_carpeta_con_dialogo(carpeta_inicial=None):
    """Abre el selector nativo de carpetas del sistema operativo (igual al
    Application.FileDialog de la macro de Excel). Solo funciona corriendo
    la app en tu propia maquina (streamlit run app.py), no en la nube --
    ahi no hay pantalla para mostrar el dialogo."""
    try:
        import tkinter as tk
        from tkinter import filedialog

        root = tk.Tk()
        root.withdraw()
        root.wm_attributes("-topmost", 1)
        carpeta = filedialog.askdirectory(initialdir=carpeta_inicial or None)
        root.destroy()
        return carpeta or None
    except Exception:
        return None


# ---------------------------------------------------------------------
# Fase 3: calculo de AIU
# (misma logica que PLANTILLA_AIU.xlsx, hoja ENTRADAS)
# ---------------------------------------------------------------------
def calcular_aiu(costo_directo, administracion_pct, imprevistos_pct, utilidad_pct, iva_utilidad_pct):
    administracion = costo_directo * administracion_pct
    imprevistos = costo_directo * imprevistos_pct
    utilidad = costo_directo * utilidad_pct
    iva_utilidad = utilidad * iva_utilidad_pct
    aiu_total_pct = administracion_pct + imprevistos_pct + utilidad_pct + utilidad_pct * iva_utilidad_pct
    aiu_total_valor = administracion + imprevistos + utilidad + iva_utilidad
    valor_total = costo_directo * (1 + aiu_total_pct)
    return {
        "administracion": administracion,
        "imprevistos": imprevistos,
        "utilidad": utilidad,
        "iva_utilidad": iva_utilidad,
        "aiu_total_pct": aiu_total_pct,
        "aiu_total_valor": aiu_total_valor,
        "valor_total": valor_total,
    }


# ---------------------------------------------------------------------
# Fase 6: margen real de la empresa -- distinto del AIU que ve el
# cliente. Regla de negocio (2026-08-03): el costo de cada APU (de
# COSTOS, tal como esta en catalogo_apu) se aumenta con este margen
# ANTES de aplicar el AIU. Ese costo aumentado es el "Costo Directo" que
# entra al AIU y termina en el precio que ve el cliente. El costo SIN
# aumentar es el "Costo real" que de verdad le cuesta a la empresa
# ejecutar la obra -- nunca se le muestra al cliente.
# ---------------------------------------------------------------------
def obtener_margen_real_pct(sb, presupuesto):
    """% de margen real a usar para este presupuesto: el que tenga
    guardado el presupuesto (override por proyecto) o, si no tiene, el
    default global de parametros.margen_real_pct."""
    valor_presupuesto = presupuesto.get("margen_real_pct") if presupuesto else None
    if valor_presupuesto is not None:
        return float(valor_presupuesto)
    return float(obtener_parametro(sb, "margen_real_pct", "25")) / 100


def aplicar_margen_items(items, margen_real_pct):
    """Devuelve una copia de 'items' con precio_unitario_snapshot
    aumentado por el margen real -- esto es lo que se le presenta al
    cliente (Word, Excel de respaldo, cronograma/flujo de caja). Los
    'items' originales (sin aumentar) siguen siendo el costo real."""
    factor = 1 + float(margen_real_pct)
    ajustados = []
    for it in items:
        nuevo = dict(it)
        nuevo["precio_unitario_snapshot"] = round(float(it["precio_unitario_snapshot"]) * factor, 0)
        ajustados.append(nuevo)
    return ajustados


def agrupar_por_capitulo(items):
    por_capitulo = {}
    for it in items:
        nombre_cap = (it.get("presupuesto_capitulos") or {}).get("nombre", "Sin capitulo")
        por_capitulo.setdefault(nombre_cap, []).append(it)
    return por_capitulo


def calcular_apu_precio_presentacion(apu, margen_real_pct):
    """Transforma un APU de COSTOS (catalogo_apu) en su version 'a
    precio de presentacion' para el cliente:
      1. La supervision (SISO/Residente/Director) se absorbe dentro de
         mano de obra y se quita como linea aparte -- el cliente ve un
         solo costo de mano de obra por actividad, no el detalle de
         quien la supervisa.
      2. Sobre ese costo (ya con la supervision adentro) se distribuye
         el margen real de la empresa, aumentando equipo/materiales/
         transporte/mano_obra en la misma proporcion.
    El total resultante coincide con precio_unitario_snapshot * (1 +
    margen_real_pct) que se usa en el resto del presupuesto -- solo
    cambia como se reparte ese total entre las columnas."""
    factor = 1 + float(margen_real_pct)
    mano_obra_con_supervision = float(apu.get("mano_obra") or 0) + float(apu.get("personal_supervision") or 0)
    return {
        "equipo": round(float(apu.get("equipo") or 0) * factor, 0),
        "materiales": round(float(apu.get("materiales") or 0) * factor, 0),
        "transporte": round(float(apu.get("transporte") or 0) * factor, 0),
        "mano_obra": round(mano_obra_con_supervision * factor, 0),
        "total": round(
            (
                float(apu.get("equipo") or 0)
                + float(apu.get("materiales") or 0)
                + float(apu.get("transporte") or 0)
                + mano_obra_con_supervision
            )
            * factor,
            0,
        ),
    }


def numero_a_letras(valor):
    """Convierte un valor en pesos a letras, igual a la funcion Nlet del Excel."""
    try:
        from num2words import num2words

        entero = int(round(valor))
        texto = num2words(entero, lang="es").upper()
        return f"{texto} PESOS M/CTE"
    except Exception:
        return ""


def obtener_items_y_capitulos(sb, presupuesto_id):
    """Trae los items de un presupuesto agrupados por capitulo (orden de captura)."""
    items = (
        sb.table("presupuesto_items")
        .select("*, presupuesto_capitulos(nombre, orden)")
        .eq("presupuesto_id", presupuesto_id)
        .order("id")
        .execute()
        .data
    )
    por_capitulo = {}
    for it in items:
        nombre_cap = (it.get("presupuesto_capitulos") or {}).get("nombre", "Sin capitulo")
        por_capitulo.setdefault(nombre_cap, []).append(it)
    return items, por_capitulo


def preparar_refresco_precios(sb, presupuesto_id):
    """Compara el precio_unitario_snapshot de cada item enlazado a un APU
    (apu_codigo no nulo) contra el catalogo_apu.total ACTUAL. Solo mira
    items de ESTE presupuesto -- ningun otro presupuesto se toca. Los
    items sin apu_codigo (importados o agregados a mano) se dejan igual,
    no hay de donde refrescarlos."""
    items = (
        sb.table("presupuesto_items")
        .select("id, apu_codigo, descripcion_snapshot, unidad_snapshot, cantidad, precio_unitario_snapshot")
        .eq("presupuesto_id", presupuesto_id)
        .not_.is_("apu_codigo", "null")
        .execute()
        .data
    )
    if not items:
        return []
    codigos = list({it["apu_codigo"] for it in items})
    catalogo = (
        sb.table("catalogo_apu")
        .select("codigo, descripcion, unidad, total")
        .in_("codigo", codigos)
        .execute()
        .data
    )
    catalogo_por_codigo = {c["codigo"]: c for c in catalogo}

    cambios = []
    for it in items:
        actual = catalogo_por_codigo.get(it["apu_codigo"])
        if actual is None:
            continue  # el APU ya no existe en el catalogo -- se deja igual
        precio_nuevo = float(actual["total"])
        precio_viejo = float(it["precio_unitario_snapshot"])
        if round(precio_nuevo, 2) == round(precio_viejo, 2):
            continue
        cambios.append(
            {
                "item_id": it["id"],
                "apu_codigo": it["apu_codigo"],
                "descripcion": actual["descripcion"] or it["descripcion_snapshot"],
                "unidad": actual["unidad"] or it["unidad_snapshot"],
                "cantidad": float(it["cantidad"]),
                "precio_viejo": precio_viejo,
                "precio_nuevo": precio_nuevo,
            }
        )
    return cambios


def aplicar_refresco_precios(sb, presupuesto_id, cambios):
    """Actualiza precio_unitario_snapshot (y descripcion/unidad_snapshot)
    de los items indicados, SOLO en este presupuesto, y recalcula su
    costo_directo. No toca la bitacora de precios del catalogo -- esto
    es un refresco puntual de un presupuesto, no un cambio de precio
    maestro."""
    for c in cambios:
        sb.table("presupuesto_items").update(
            {
                "precio_unitario_snapshot": c["precio_nuevo"],
                "descripcion_snapshot": c["descripcion"],
                "unidad_snapshot": c["unidad"],
            }
        ).eq("id", c["item_id"]).execute()

    items_actualizados, _ = obtener_items_y_capitulos(sb, presupuesto_id)
    nuevo_costo_directo = sum(
        float(it["cantidad"]) * float(it["precio_unitario_snapshot"]) for it in items_actualizados
    )
    sb.table("presupuestos").update({"costo_directo": nuevo_costo_directo}).eq(
        "id", presupuesto_id
    ).execute()
    return nuevo_costo_directo


# ---------------------------------------------------------------------
# Fase 3: generar la propuesta en Word a partir de la plantilla real
# ---------------------------------------------------------------------
def _reemplazar_tokens_simples(paragraph, mapa):
    texto_completo = "".join(r.text for r in paragraph.runs)
    if not texto_completo or "[[" not in texto_completo:
        return
    nuevo_texto = texto_completo
    cambiado = False
    for token, valor in mapa.items():
        if token in nuevo_texto:
            nuevo_texto = nuevo_texto.replace(token, valor)
            cambiado = True
    if not cambiado:
        return
    for i, run in enumerate(paragraph.runs):
        run.text = nuevo_texto if i == 0 else ""


def _insertar_parrafo_despues(paragraph):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    nuevo_p = OxmlElement("w:p")
    paragraph._p.addnext(nuevo_p)
    return Paragraph(nuevo_p, paragraph._parent)


def _eliminar_parrafo(paragraph):
    p = paragraph._p
    p.getparent().remove(p)


# ---------------------------------------------------------------------
# Fase 3: formato institucional (igual al Excel "PRESUPUESTO DE OBRA
# CIVIL" que se le presenta al cliente). Esta lista de filas es la unica
# fuente de verdad: la usan tanto el Excel de respaldo como la imagen
# del anexo en la propuesta de Word, para que los dos documentos
# muestren siempre el mismo numero.
# ---------------------------------------------------------------------
def construir_filas_institucional(presupuesto, por_capitulo, costo_directo, aiu):
    filas = [
        {"tipo": "titulo", "texto": f"PRESUPUESTO DE OBRA CIVIL - {(presupuesto.get('proyecto') or '').upper()}"},
        {"tipo": "encabezado", "valores": ["ITEM", "ACTIVIDAD", "UN", "CANT", "VR UNITARIO", "VR TOTAL"]},
    ]

    num_cap = 0
    for nombre_cap, items_cap in por_capitulo.items():
        num_cap += 1
        filas.append({"tipo": "capitulo", "num": str(num_cap), "nombre": nombre_cap.upper()})
        subtotal_cap = 0.0
        for i, it in enumerate(items_cap, start=1):
            cantidad = float(it["cantidad"])
            precio = float(it["precio_unitario_snapshot"])
            subtotal = cantidad * precio
            subtotal_cap += subtotal
            filas.append(
                {
                    "tipo": "item",
                    "num": f"{num_cap}.{i}",
                    "descripcion": it["descripcion_snapshot"],
                    "unidad": it["unidad_snapshot"],
                    "cantidad": cantidad,
                    "vr_unitario": precio,
                    "vr_total": subtotal,
                }
            )
        filas.append(
            {"tipo": "subtotal", "etiqueta": f"Sub-Total {nombre_cap.upper()}", "valor": subtotal_cap}
        )

    filas.append({"tipo": "total_cd", "etiqueta": "TOTAL COSTOS DIRECTOS", "valor": costo_directo})
    filas.append(
        {
            "tipo": "aiu",
            "etiqueta": "ADMINISTRACION",
            "pct": float(presupuesto.get("administracion_pct") or 0),
            "valor": aiu["administracion"],
        }
    )
    filas.append(
        {
            "tipo": "aiu",
            "etiqueta": "IMPREVISTOS",
            "pct": float(presupuesto.get("imprevistos_pct") or 0),
            "valor": aiu["imprevistos"],
        }
    )
    filas.append(
        {
            "tipo": "aiu",
            "etiqueta": "UTILIDAD",
            "pct": float(presupuesto.get("utilidad_pct") or 0),
            "valor": aiu["utilidad"],
        }
    )
    filas.append(
        {
            "tipo": "aiu",
            "etiqueta": "IVA SOBRE UTILIDAD",
            "pct": float(presupuesto.get("iva_utilidad_pct") or 0),
            "valor": aiu["iva_utilidad"],
        }
    )
    filas.append({"tipo": "total_obra", "etiqueta": "TOTAL OBRA", "valor": aiu["valor_total"]})
    filas.append({"tipo": "son", "texto": f"SON: {numero_a_letras(aiu['valor_total'])}"})
    return filas


def generar_imagen_presupuesto_institucional(presupuesto, por_capitulo, total_general, aiu):
    """Dibuja el presupuesto en formato institucional (item numerado,
    capitulos, subtotales, AIU y SON: valor en letras) como una imagen,
    para pegarla en el anexo de la propuesta de Word -- igual a como se
    pegaba una imagen del rango de Excel en la macro GENERAR_PROPUESTA."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    filas = construir_filas_institucional(presupuesto, por_capitulo, total_general, aiu)

    anchos = [0.07, 0.45, 0.07, 0.10, 0.155, 0.155]
    bordes_x = [0.0]
    for a in anchos:
        bordes_x.append(bordes_x[-1] + a)

    alto_fila = 0.28
    alto_titulo = 0.45
    alturas = [alto_titulo if f["tipo"] == "titulo" else alto_fila for f in filas]
    alto_total = sum(alturas) + 0.15

    fig, ax = plt.subplots(figsize=(11, alto_total))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, alto_total)
    ax.axis("off")

    for x in bordes_x[1:-1]:
        ax.plot([x, x], [0, alto_total], color="#cccccc", linewidth=0.5, zorder=0)

    fontsize_item = 7.5
    limite_desc = bordes_x[2] - 0.015  # borde derecho de la columna ACTIVIDAD, con margen
    textos_descripcion = []  # (objeto de texto, limite_x_derecho) -- se ajustan al final

    y = alto_total
    for f, alto in zip(filas, alturas):
        y -= alto
        tipo = f["tipo"]
        y_centro = y + alto / 2

        if tipo == "titulo":
            ax.text(0.5, y_centro, f["texto"], ha="center", va="center", fontsize=12, fontweight="bold")
            ax.plot([0, 1], [y, y], color="black", linewidth=1.2)
        elif tipo == "encabezado":
            for i, texto in enumerate(f["valores"]):
                cx = (bordes_x[i] + bordes_x[i + 1]) / 2
                ax.text(cx, y_centro, texto, ha="center", va="center", fontsize=8.5, fontweight="bold")
            ax.plot([0, 1], [y, y], color="black", linewidth=1)
        elif tipo == "capitulo":
            ax.add_patch(plt.Rectangle((0, y), 1, alto, facecolor="#D9D9D9", edgecolor="none", zorder=0))
            ax.text(bordes_x[0] + 0.01, y_centro, f["num"], ha="left", va="center", fontsize=8.5, fontweight="bold")
            ax.text(bordes_x[1] + 0.01, y_centro, f["nombre"], ha="left", va="center", fontsize=8.5, fontweight="bold")
        elif tipo == "item":
            ax.text(bordes_x[0] + 0.01, y_centro, f["num"], ha="left", va="center", fontsize=fontsize_item)
            texto_desc = ax.text(
                bordes_x[1] + 0.01, y_centro, f["descripcion"], ha="left", va="center", fontsize=fontsize_item
            )
            textos_descripcion.append((texto_desc, limite_desc))
            ax.text((bordes_x[2] + bordes_x[3]) / 2, y_centro, f["unidad"], ha="center", va="center", fontsize=fontsize_item)
            ax.text((bordes_x[3] + bordes_x[4]) / 2, y_centro, f"{f['cantidad']:g}", ha="center", va="center", fontsize=fontsize_item)
            ax.text(bordes_x[5] - 0.01, y_centro, money(f["vr_unitario"]), ha="right", va="center", fontsize=fontsize_item)
            ax.text(0.99, y_centro, money(f["vr_total"]), ha="right", va="center", fontsize=fontsize_item)
        elif tipo == "subtotal":
            ax.plot([bordes_x[4], 1], [y + alto, y + alto], color="black", linewidth=0.8)
            ax.text(bordes_x[4] - 0.01, y_centro, f["etiqueta"], ha="right", va="center", fontsize=8, fontweight="bold")
            ax.text(0.99, y_centro, money(f["valor"]), ha="right", va="center", fontsize=8, fontweight="bold")
        elif tipo == "total_cd":
            ax.plot([0, 1], [y + alto, y + alto], color="black", linewidth=1.4)
            ax.text(bordes_x[4] - 0.01, y_centro, f["etiqueta"], ha="right", va="center", fontsize=9, fontweight="bold")
            ax.text(0.99, y_centro, money(f["valor"]), ha="right", va="center", fontsize=9, fontweight="bold")
        elif tipo == "aiu":
            ax.text(bordes_x[4] - 0.01, y_centro, f["etiqueta"], ha="right", va="center", fontsize=8)
            ax.text((bordes_x[4] + bordes_x[5]) / 2, y_centro, f"{f['pct'] * 100:g}%", ha="center", va="center", fontsize=8)
            ax.text(0.99, y_centro, money(f["valor"]), ha="right", va="center", fontsize=8)
        elif tipo == "total_obra":
            ax.plot([0, 1], [y + alto, y + alto], color="black", linewidth=1.6)
            ax.plot([0, 1], [y, y], color="black", linewidth=1.6)
            ax.text(bordes_x[4] - 0.01, y_centro, f["etiqueta"], ha="right", va="center", fontsize=10, fontweight="bold")
            ax.text(0.99, y_centro, money(f["valor"]), ha="right", va="center", fontsize=10, fontweight="bold")
        elif tipo == "son":
            ax.text(0.5, y_centro, f["texto"], ha="center", va="center", fontsize=8, style="italic", wrap=True)

    ax.add_patch(plt.Rectangle((0, 0), 1, alto_total, fill=False, edgecolor="black", linewidth=1.2))

    # Ajustar las descripciones que se salen de su columna (en vez de un
    # truncado fijo por caracteres, que se ve distinto segun cuantas letras
    # anchas tenga cada texto): primero se reduce el tamano de letra hasta un
    # minimo, y si aun no cabe, se trunca con puntos suspensivos.
    fig.canvas.draw()
    renderer = fig.canvas.get_renderer()
    fontsize_min = 6.0
    for texto_obj, limite_x in textos_descripcion:
        tam = texto_obj.get_fontsize()
        while tam > fontsize_min:
            bbox = texto_obj.get_window_extent(renderer=renderer).transformed(ax.transData.inverted())
            if bbox.x1 <= limite_x:
                break
            tam -= 0.5
            texto_obj.set_fontsize(tam)
            fig.canvas.draw()

        bbox = texto_obj.get_window_extent(renderer=renderer).transformed(ax.transData.inverted())
        texto_actual = texto_obj.get_text()
        while bbox.x1 > limite_x and len(texto_actual) > 5:
            texto_actual = texto_actual[:-2].rstrip()
            texto_obj.set_text(texto_actual + "...")
            fig.canvas.draw()
            bbox = texto_obj.get_window_extent(renderer=renderer).transformed(ax.transData.inverted())

    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=200, bbox_inches="tight", pad_inches=0.1)
    plt.close(fig)
    buffer.seek(0)
    return buffer


def _insertar_resumen_capitulos(doc, por_capitulo):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    target = None
    for p in doc.paragraphs:
        if "[[CAPITULOS]]" in p.text:
            target = p
            break
    if target is None:
        return
    ancla = target
    for nombre_cap, filas in por_capitulo.items():
        nuevo = _insertar_parrafo_despues(ancla)
        nuevo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_nombre = nuevo.add_run(nombre_cap.upper())
        run_nombre.bold = True
        descripciones = [f["descripcion_snapshot"] for f in filas]
        if descripciones:
            nuevo.add_run(": " + ", ".join(descripciones) + ".")
        ancla = nuevo
    _eliminar_parrafo(target)


def _insertar_anexo_presupuesto(doc, imagen_buffer):
    """Pega la imagen del presupuesto (formato institucional) donde estaba
    el token [[ANEXO]], en una pagina aparte (salto de pagina antes),
    separada de la pagina donde queda el titulo ANEXOS."""
    target = None
    for p in doc.paragraphs:
        if "[[ANEXO]]" in p.text:
            target = p
            break
    if target is None:
        return

    for run in target.runs:
        run.text = ""
    target.paragraph_format.page_break_before = True

    seccion = doc.sections[0]
    ancho_disponible = seccion.page_width - seccion.left_margin - seccion.right_margin

    run = target.add_run()
    run.add_picture(imagen_buffer, width=ancho_disponible)


def generar_propuesta_docx(plantilla_path, presupuesto, items, por_capitulo, total_general, aiu):
    import docx

    doc = docx.Document(plantilla_path)

    meses_es = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    hoy = date.today()
    fecha_txt = f"{hoy.day} de {meses_es[hoy.month - 1]} de {hoy.year}"
    plazo_dias = presupuesto.get("plazo_dias")
    plazo_dias = int(round(float(plazo_dias))) if plazo_dias else int(round(float(presupuesto.get("plazo_meses") or 0) * 30))
    valor_total = aiu["valor_total"]
    letras = numero_a_letras(valor_total)
    valor_txt = (
        f"El valor de los trabajos a realizar es de {letras + ' ' if letras else ''}"
        f"({money(valor_total)}). El detalle de este valor se presenta en los anexos."
    )

    mapa = {
        "[[CLIENTE]]": (presupuesto.get("cliente") or "").upper(),
        "[[PROYECTO]]": (presupuesto.get("proyecto") or "").upper(),
        "[[FECHA]]": fecha_txt,
        "[[ATENCION]]": presupuesto.get("atencion") or "",
        "[[UBIC]]": presupuesto.get("ubicacion") or "",
        "[[VALOR]]": valor_txt,
        "[[PLAZO]]": str(plazo_dias),
        "[[ANTICIPO]]": f"{float(presupuesto.get('anticipo_pct') or 0) * 100:g}%",
        "[[PAGO2]]": f"{float(presupuesto.get('pago2_pct') or 0) * 100:g}%",
        "[[AVANCE2]]": f"{float(presupuesto.get('avance2_pct') or 0) * 100:g}%",
        "[[PAGOFIN]]": f"{float(presupuesto.get('pagofin_pct') or 0) * 100:g}%",
        "[[VIGENCIA]]": str(int(presupuesto.get("vigencia_dias") or 30)),
    }

    for p in doc.paragraphs:
        _reemplazar_tokens_simples(p, mapa)

    _insertar_resumen_capitulos(doc, por_capitulo)
    imagen_anexo = generar_imagen_presupuesto_institucional(presupuesto, por_capitulo, total_general, aiu)
    _insertar_anexo_presupuesto(doc, imagen_anexo)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------
# Fase 3: generar el Excel de respaldo de un presupuesto
# (una sola hoja, formato institucional: igual al que ve el cliente)
# ---------------------------------------------------------------------
def _escribir_filas_institucional(ws, filas, usar_formula_letras=False):
    """Escribe las filas del formato institucional en una hoja ya creada,
    con formulas reales (item = cant*vr_unitario, subtotales = SUM(...),
    AIU = %*costo_directo, etc). Si usar_formula_letras=True, la fila
    SON: queda como formula que llama a la macro Nlet() del Excel
    maestro (letras vivas); si no, queda como texto fijo."""
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    negrita = Font(bold=True)
    negrita_grande = Font(bold=True, size=12)
    cursiva = Font(italic=True)
    relleno_cap = PatternFill("solid", fgColor="D9D9D9")
    borde_top = Border(top=Side(style="thin"))
    borde_top_grueso = Border(top=Side(style="double"))

    # Seguimiento de filas para poder escribir formulas reales (no valores
    # fijos): cada capitulo, subtotal, AIU y total quedan enlazados entre
    # si -- si alguien cambia una cantidad o un porcentaje en el Excel,
    # todo lo demas se recalcula solo, igual que en el institucional.
    fila_actual = 0
    primera_fila_items = None
    ultima_fila_items = None
    filas_subtotal = []
    fila_total_cd = None
    fila_administracion = None
    fila_imprevistos = None
    fila_utilidad = None
    fila_iva = None
    fila_total_obra = None

    for f in filas:
        fila_actual += 1
        tipo = f["tipo"]

        if tipo == "titulo":
            ws.cell(row=fila_actual, column=1, value=f["texto"])
            ws.merge_cells(start_row=fila_actual, start_column=1, end_row=fila_actual, end_column=6)
            celda = ws.cell(row=fila_actual, column=1)
            celda.font = Font(bold=True, size=13)
            celda.alignment = Alignment(horizontal="center")

        elif tipo == "encabezado":
            for col, texto in enumerate(f["valores"], start=1):
                c = ws.cell(row=fila_actual, column=col, value=texto)
                c.font = negrita
                c.alignment = Alignment(horizontal="center")

        elif tipo == "capitulo":
            ws.cell(row=fila_actual, column=1, value=f["num"])
            ws.cell(row=fila_actual, column=2, value=f["nombre"])
            for col in range(1, 7):
                c = ws.cell(row=fila_actual, column=col)
                c.font = negrita
                c.fill = relleno_cap
            primera_fila_items = None
            ultima_fila_items = None

        elif tipo == "item":
            ws.cell(row=fila_actual, column=1, value=f["num"])
            ws.cell(row=fila_actual, column=2, value=f["descripcion"])
            ws.cell(row=fila_actual, column=3, value=f["unidad"])
            ws.cell(row=fila_actual, column=4, value=f["cantidad"])
            ws.cell(row=fila_actual, column=5, value=f["vr_unitario"]).number_format = "$ #,##0"
            c_total = ws.cell(row=fila_actual, column=6)
            c_total.value = f"=D{fila_actual}*E{fila_actual}"
            c_total.number_format = "$ #,##0"
            if primera_fila_items is None:
                primera_fila_items = fila_actual
            ultima_fila_items = fila_actual

        elif tipo == "subtotal":
            ws.cell(row=fila_actual, column=1, value=f["etiqueta"])
            ws.merge_cells(start_row=fila_actual, start_column=1, end_row=fila_actual, end_column=5)
            c_val = ws.cell(row=fila_actual, column=6)
            if primera_fila_items is not None:
                c_val.value = f"=SUM(F{primera_fila_items}:F{ultima_fila_items})"
            else:
                c_val.value = 0
            c_val.number_format = "$ #,##0"
            for col in range(1, 7):
                c = ws.cell(row=fila_actual, column=col)
                c.font = negrita
                c.border = borde_top
            filas_subtotal.append(fila_actual)

        elif tipo == "total_cd":
            ws.cell(row=fila_actual, column=1, value=f["etiqueta"])
            ws.merge_cells(start_row=fila_actual, start_column=1, end_row=fila_actual, end_column=5)
            c_val = ws.cell(row=fila_actual, column=6)
            if filas_subtotal:
                c_val.value = "=" + "+".join(f"F{r}" for r in filas_subtotal)
            else:
                c_val.value = 0
            c_val.number_format = "$ #,##0"
            for col in range(1, 7):
                c = ws.cell(row=fila_actual, column=col)
                c.font = negrita
                c.border = borde_top_grueso
            fila_total_cd = fila_actual

        elif tipo == "aiu":
            ws.cell(row=fila_actual, column=1, value=f["etiqueta"])
            ws.merge_cells(start_row=fila_actual, start_column=1, end_row=fila_actual, end_column=4)
            c_pct = ws.cell(row=fila_actual, column=5, value=f["pct"])
            c_pct.number_format = "0%"
            c_pct.alignment = Alignment(horizontal="center")
            c_val = ws.cell(row=fila_actual, column=6)
            etiqueta = f["etiqueta"].upper()
            if etiqueta == "IVA SOBRE UTILIDAD" and fila_utilidad is not None:
                # el IVA es un % de la UTILIDAD, no del costo directo
                c_val.value = f"=E{fila_actual}*F{fila_utilidad}"
            elif fila_total_cd is not None:
                c_val.value = f"=E{fila_actual}*$F${fila_total_cd}"
            else:
                c_val.value = f["valor"]
            c_val.number_format = "$ #,##0"
            if etiqueta == "ADMINISTRACION":
                fila_administracion = fila_actual
            elif etiqueta == "IMPREVISTOS":
                fila_imprevistos = fila_actual
            elif etiqueta == "UTILIDAD":
                fila_utilidad = fila_actual
            elif etiqueta == "IVA SOBRE UTILIDAD":
                fila_iva = fila_actual

        elif tipo == "total_obra":
            ws.cell(row=fila_actual, column=1, value=f["etiqueta"])
            ws.merge_cells(start_row=fila_actual, start_column=1, end_row=fila_actual, end_column=5)
            c_val = ws.cell(row=fila_actual, column=6)
            componentes = [r for r in (fila_total_cd, fila_administracion, fila_imprevistos, fila_utilidad, fila_iva) if r]
            if componentes:
                c_val.value = "=" + "+".join(f"F{r}" for r in componentes)
            else:
                c_val.value = f["valor"]
            c_val.number_format = "$ #,##0"
            for col in range(1, 7):
                c = ws.cell(row=fila_actual, column=col)
                c.font = negrita_grande
                c.border = borde_top_grueso
            fila_total_obra = fila_actual

        elif tipo == "son":
            celda_son = ws.cell(row=fila_actual, column=1)
            if usar_formula_letras and fila_total_obra is not None:
                # Letras vivas: llama a la macro Nlet() del Excel maestro.
                # Si se cambia una cantidad o un porcentaje despues, este
                # texto se recalcula solo (requiere abrir el archivo con
                # macros habilitadas).
                celda_son.value = f'="SON: "&Nlet(F{fila_total_obra})'
            else:
                # Texto fijo con el valor en letras al momento de generar
                # -- no se puede formular sin macros.
                celda_son.value = f["texto"]
            ws.merge_cells(start_row=fila_actual, start_column=1, end_row=fila_actual, end_column=6)
            celda_son.font = cursiva

    for col, ancho in zip("ABCDEF", (8, 50, 8, 10, 15, 15)):
        ws.column_dimensions[col].width = ancho

    return fila_total_obra


def _insertar_logo_excel(ws, celda_ancla="H2", ancho_px=140):
    """Pega el logo COMSAS como imagen flotante en una hoja de Excel, sin
    tocar el formato institucional de la tabla (se ancla fuera de las
    columnas A-F que ocupa el presupuesto). No falla el resto de la
    generacion si el logo no existe o Pillow no esta disponible."""
    if not LOGO_PATH.exists():
        return
    try:
        from openpyxl.drawing.image import Image as XLImage

        img = XLImage(str(LOGO_PATH))
        escala = ancho_px / img.width
        img.width = ancho_px
        img.height = int(img.height * escala)
        ws.add_image(img, celda_ancla)
    except Exception:
        pass


def _escribir_hoja_apus_presentacion(sb, wb, items, margen_real_pct):
    """Agrega al libro una hoja 'APUS' con el desglose de cada actividad
    del presupuesto A PRECIO DE PRESENTACION: con el margen real de la
    empresa ya incluido y la supervision (SISO/Residente/Director)
    absorbida dentro de mano de obra (sin mostrarla como linea aparte).
    El total de cada APU aqui coincide con el precio unitario que se le
    presenta al cliente en el presupuesto -- son los mismos APUs de
    COSTOS del catalogo, pasados por calcular_apu_precio_presentacion()."""
    from openpyxl.styles import Border, Font, PatternFill, Side

    negrita = Font(bold=True)
    negrita_grande = Font(bold=True, size=13)
    relleno_encabezado = PatternFill("solid", fgColor="D9D9D9")
    borde_top = Border(top=Side(style="thin"))

    codigos_apu = sorted({it["apu_codigo"] for it in items if it.get("apu_codigo")})
    ws = wb.create_sheet("APUS")
    ws.cell(
        row=1, column=1,
        value="APUs A PRECIO DE PRESENTACION (incluye el margen real de la empresa)",
    ).font = negrita_grande
    ws.cell(
        row=2, column=1,
        value=(
            "La supervision (SISO/Residente/Director) queda incluida dentro de MANO DE "
            "OBRA -- no se muestra como linea aparte. Estos totales ya incluyen el margen "
            "real de la empresa y coinciden con los precios unitarios de este presupuesto."
        ),
    )
    fila = 4
    for codigo in codigos_apu:
        apu = obtener_apu_detalle(sb, codigo)
        if not apu:
            ws.cell(row=fila, column=1, value=f"{codigo} -- no encontrado en el catalogo")
            fila += 2
            continue
        pres = calcular_apu_precio_presentacion(apu, margen_real_pct)
        ws.cell(
            row=fila, column=1, value=f"{codigo} - {apu['descripcion']} ({apu['unidad']})"
        ).font = negrita_grande
        fila += 1
        ws.append(["Concepto", "Valor"])
        for c in ws[fila]:
            c.font = negrita
            c.fill = relleno_encabezado
        fila += 1
        for etiqueta, valor in (
            ("EQUIPO", pres["equipo"]),
            ("MATERIALES", pres["materiales"]),
            ("TRANSPORTE", pres["transporte"]),
            ("MANO DE OBRA (incluye supervision)", pres["mano_obra"]),
        ):
            ws.cell(row=fila, column=1, value=etiqueta)
            c_val = ws.cell(row=fila, column=2, value=valor)
            c_val.number_format = "$ #,##0"
            fila += 1
        ws.cell(row=fila, column=1, value=f"TOTAL {codigo} (precio de presentacion)").font = negrita
        c_total = ws.cell(row=fila, column=2, value=pres["total"])
        c_total.font = negrita
        c_total.number_format = "$ #,##0"
        c_total.border = borde_top
        fila += 3

    ws.column_dimensions["A"].width = 55
    ws.column_dimensions["B"].width = 16


def generar_excel_respaldo(sb, presupuesto, items, por_capitulo, total_general, aiu, margen_real_pct):
    """Respaldo simple en .xlsx (sin macros): todo con formulas reales,
    excepto el texto SON: que queda fijo (Excel no tiene forma de
    escribir un numero en letras sin una macro). Incluye una hoja APUS
    con el detalle de cada actividad a precio de presentacion."""
    import openpyxl

    filas = construir_filas_institucional(presupuesto, por_capitulo, total_general, aiu)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "PRESUPUESTO"
    _escribir_filas_institucional(ws, filas, usar_formula_letras=False)
    _insertar_logo_excel(ws)
    _escribir_hoja_apus_presentacion(sb, wb, items, margen_real_pct)

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def generar_excel_respaldo_macro(
    sb, presupuesto, items, por_capitulo, total_general, aiu, plantilla_macro_path, margen_real_pct
):
    """Respaldo en .xlsm partiendo de PLANTILLA_MACRO_LETRAS.xlsm -- un
    archivo "portador" que el usuario crea una sola vez en Excel y que
    SOLO contiene la macro Nlet (copiada de tu ModArmar.bas), sin ninguna
    hoja del modulo maestro. Asi el respaldo no carga datos internos
    (precios, clientes, parametros) por dentro, ni siquiera ocultos --
    si algun dia se comparte por error con un cliente, no hay nada que
    exponer. Las hojas que traiga ese portador (idealmente una sola,
    vacia) se ocultan y se protege la estructura del libro; la nueva
    hoja con el presupuesto queda como la unica visible. El valor en
    letras queda como formula viva: '=\"SON: \"&Nlet(total)'."""
    import openpyxl

    filas = construir_filas_institucional(presupuesto, por_capitulo, total_general, aiu)
    wb = openpyxl.load_workbook(plantilla_macro_path, keep_vba=True)

    for nombre in list(wb.sheetnames):
        wb[nombre].sheet_state = "hidden"

    nombre_hoja = "PRESUPUESTO_APP" if "PRESUPUESTO" in wb.sheetnames else "PRESUPUESTO"
    ws = wb.create_sheet(nombre_hoja, 0)
    ws.sheet_state = "visible"
    wb.active = wb.sheetnames.index(nombre_hoja)

    _escribir_filas_institucional(ws, filas, usar_formula_letras=True)
    _insertar_logo_excel(ws)
    _escribir_hoja_apus_presentacion(sb, wb, items, margen_real_pct)

    # Ademas de ocultarlas, protege la ESTRUCTURA del libro (no el
    # contenido de las celdas) para que la opcion "Mostrar" de Excel
    # quede deshabilitada mientras el libro este protegido -- asi las
    # hojas ocultas no reaparecen por un clic accidental. Se puede
    # desproteger desde Revisar > Proteger libro si alguna vez hace falta
    # entrar a esas hojas.
    from openpyxl.workbook.protection import WorkbookProtection

    if wb.security is None:
        wb.security = WorkbookProtection()
    wb.security.lockStructure = True

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------
# Excel de manejo interno (Cronograma, Lista de materiales, Flujo de
# Caja, Dashboard y APUs de las actividades del presupuesto).
# Usa los datos reales de Fase 2 (cuadrillas/cargos) y Fase 4
# (insumos/apu_insumos/apu_materiales_fijos) para que TODO quede con
# formulas de verdad -- si cambias un precio de insumo, una tarifa de
# cargo, o una cantidad, se recalcula solo en cadena.
# NO se manda al cliente: es para uso interno del equipo.
# ---------------------------------------------------------------------
def generar_excel_manejo_interno(sb, presupuesto, items, por_capitulo, costo_directo, aiu, usar_curva_s=False):
    import math

    import openpyxl
    from openpyxl.chart import BarChart, LineChart, Reference
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    negrita = Font(bold=True)
    negrita_grande = Font(bold=True, size=13)
    titulo_hoja = Font(bold=True, size=14)
    relleno_encabezado = PatternFill("solid", fgColor="D9D9D9")
    borde_top = Border(top=Side(style="thin"))

    # -------------------------------------------------------------
    # 1. Recolectar datos de Fase 2 / Fase 4 para los APUs usados
    # -------------------------------------------------------------
    codigos_apu = sorted({it["apu_codigo"] for it in items if it.get("apu_codigo")})
    detalle_apus = {c: obtener_apu_detalle(sb, c) for c in codigos_apu}
    insumos_por_apu = {c: obtener_insumos_de_apu(sb, c) for c in codigos_apu}
    fijos_por_apu = {c: obtener_fijos_de_apu(sb, c) for c in codigos_apu}

    cuadrillas_usadas = sorted(
        {
            d["cuadrilla_codigo"]
            for d in detalle_apus.values()
            if d and d.get("cuadrilla_codigo")
        }
    )
    composicion_cuadrillas = {cc: obtener_composicion_cuadrilla(sb, cc) for cc in cuadrillas_usadas}
    cuadrillas_info = {c["codigo"]: c for c in obtener_cuadrillas(sb) if c["codigo"] in cuadrillas_usadas}
    cargos_usados = sorted({comp["cargo"] for lst in composicion_cuadrillas.values() for comp in lst})
    cargos_personal = {c["cargo"]: c for c in obtener_cargos_personal(sb) if c["cargo"] in cargos_usados}

    insumos_info = {}
    for lst in insumos_por_apu.values():
        for f in lst:
            insumos_info[f["insumo_codigo"]] = {
                "descripcion": f["descripcion"],
                "unidad": f["unidad"],
                "precio": f["precio"],
            }
    insumos_usados = sorted(insumos_info.keys())

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # -------------------------------------------------------------
    # 2. Hoja BASE: precios de insumos, tarifas de cargos y costo/dia
    #    de cuadrillas -- todo lo demas referencia estas celdas.
    # -------------------------------------------------------------
    ws_base = wb.create_sheet("BASE")
    ws_base.sheet_properties.tabColor = "BFBFBF"
    fila = 1
    ws_base.cell(row=fila, column=1, value="INSUMOS USADOS EN ESTE PRESUPUESTO").font = negrita_grande
    fila += 2
    ws_base.append(["Codigo", "Descripcion", "Unidad", "Precio"])
    for c in ws_base[fila]:
        c.font = negrita
    fila += 1
    fila_insumo = {}
    for codigo in insumos_usados:
        info = insumos_info[codigo]
        ws_base.append([codigo, info["descripcion"], info["unidad"], info["precio"]])
        ws_base.cell(row=fila, column=4).number_format = "$ #,##0"
        fila_insumo[codigo] = fila
        fila += 1

    fila += 2
    ws_base.cell(row=fila, column=1, value="CARGOS DE PERSONAL (tarifa/dia)").font = negrita_grande
    fila += 2
    ws_base.append(["Cargo", "Tarifa/dia"])
    for c in ws_base[fila]:
        c.font = negrita
    fila += 1
    fila_cargo = {}
    for cargo, info in cargos_personal.items():
        ws_base.append([cargo, float(info["tarifa_dia"])])
        ws_base.cell(row=fila, column=2).number_format = "$ #,##0"
        fila_cargo[cargo] = fila
        fila += 1

    fila += 2
    ws_base.cell(row=fila, column=1, value="CUADRILLAS (costo/dia)").font = negrita_grande
    fila += 2
    fila_costo_dia_cuadrilla = {}
    for codigo_cua in cuadrillas_usadas:
        info_cua = cuadrillas_info.get(codigo_cua, {})
        frentes = max(int(info_cua.get("frentes_maestro") or 1), 1)
        ws_base.cell(row=fila, column=1, value=f"{codigo_cua} - {info_cua.get('nombre', '')}").font = negrita
        fila += 1
        ws_base.append(["Cargo", "Cantidad", "Subtotal/dia"])
        for c in ws_base[fila]:
            c.font = negrita
        fila += 1
        fila_ini_comp = fila
        for comp in composicion_cuadrillas.get(codigo_cua, []):
            cargo = comp["cargo"]
            cantidad = float(comp["cantidad"])
            fref = fila_cargo.get(cargo)
            if fref:
                formula = f"=B{fila}*BASE!B{fref}"
                if cargo == "MAESTRO":
                    formula = f"=(B{fila}*BASE!B{fref})/{frentes}"
            else:
                formula = cantidad
            ws_base.append([cargo, cantidad, formula])
            ws_base.cell(row=fila, column=3).number_format = "$ #,##0"
            fila += 1
        fila_fin_comp = fila - 1
        ws_base.cell(row=fila, column=1, value="COSTO/DIA").font = negrita
        if fila_fin_comp >= fila_ini_comp:
            ws_base.cell(row=fila, column=3, value=f"=SUM(C{fila_ini_comp}:C{fila_fin_comp})")
        else:
            ws_base.cell(row=fila, column=3, value=0)
        ws_base.cell(row=fila, column=3).font = negrita
        ws_base.cell(row=fila, column=3).number_format = "$ #,##0"
        fila_costo_dia_cuadrilla[codigo_cua] = fila
        fila += 2

    for col, ancho in zip("ABCD", (16, 45, 10, 14)):
        ws_base.column_dimensions[col].width = ancho

    # -------------------------------------------------------------
    # 3. Hoja APUS: desglose de cada actividad usada en el presupuesto
    # -------------------------------------------------------------
    ws_apus = wb.create_sheet("APUS")
    fila = 1
    fila_total_apu = {}
    for codigo in codigos_apu:
        det = detalle_apus.get(codigo)
        if not det:
            ws_apus.cell(row=fila, column=1, value=f"{codigo} -- no encontrado en el catalogo")
            fila += 2
            continue

        ws_apus.cell(
            row=fila, column=1, value=f"{codigo} - {det['descripcion']}  ({det['unidad']})"
        ).font = negrita_grande
        fila += 1
        ws_apus.append(["Descripcion", "Unidad", "Cantidad", "Precio unit.", "Subtotal"])
        for c in ws_apus[fila]:
            c.font = negrita
            c.fill = relleno_encabezado
        fila += 1

        fila_ini_mat = fila
        for ins in insumos_por_apu.get(codigo, []):
            fref = fila_insumo.get(ins["insumo_codigo"])
            precio = f"=BASE!D{fref}" if fref else ins["precio"]
            ws_apus.append([ins["descripcion"], ins["unidad"], ins["cantidad"], precio, None])
            ws_apus.cell(row=fila, column=5, value=f"=C{fila}*D{fila}")
            ws_apus.cell(row=fila, column=4).number_format = "$ #,##0.00"
            ws_apus.cell(row=fila, column=5).number_format = "$ #,##0"
            fila += 1
        for fijo in fijos_por_apu.get(codigo, []):
            ws_apus.append(
                [fijo["descripcion"], fijo["unidad"], fijo["cantidad"], fijo["precio_unitario"], None]
            )
            ws_apus.cell(row=fila, column=5, value=f"=C{fila}*D{fila}")
            ws_apus.cell(row=fila, column=4).number_format = "$ #,##0.00"
            ws_apus.cell(row=fila, column=5).number_format = "$ #,##0"
            fila += 1
        fila_fin_mat = fila - 1

        ws_apus.cell(row=fila, column=1, value="Subtotal MATERIALES").font = negrita
        celda_mat = ws_apus.cell(row=fila, column=5)
        celda_mat.value = f"=SUM(E{fila_ini_mat}:E{fila_fin_mat})" if fila_fin_mat >= fila_ini_mat else 0
        celda_mat.font = negrita
        celda_mat.number_format = "$ #,##0"
        celda_mat.border = borde_top
        fila_materiales = fila
        fila += 1

        ws_apus.cell(row=fila, column=1, value="MANO DE OBRA")
        celda_mo = ws_apus.cell(row=fila, column=5)
        rendimiento = float(det.get("rendimiento_dia") or 0)
        fref_cua = fila_costo_dia_cuadrilla.get(det.get("cuadrilla_codigo"))
        if fref_cua and rendimiento:
            celda_mo.value = f"=BASE!C{fref_cua}/{rendimiento}"
        else:
            celda_mo.value = float(det.get("mano_obra") or 0)
        celda_mo.number_format = "$ #,##0"
        fila_mano_obra = fila
        fila += 1

        etiquetas_estaticas = [
            ("EQUIPO", "equipo"),
            ("TRANSPORTE", "transporte"),
            ("PERSONAL SUPERVISION", "personal_supervision"),
        ]
        filas_estaticas = {}
        for etiqueta, campo in etiquetas_estaticas:
            ws_apus.cell(row=fila, column=1, value=etiqueta)
            c_val = ws_apus.cell(row=fila, column=5, value=float(det.get(campo) or 0))
            c_val.number_format = "$ #,##0"
            filas_estaticas[campo] = fila
            fila += 1

        ws_apus.cell(row=fila, column=1, value=f"TOTAL {codigo}").font = negrita
        c_total = ws_apus.cell(
            row=fila,
            column=5,
            value=(
                f"=E{fila_materiales}+E{fila_mano_obra}+E{filas_estaticas['equipo']}"
                f"+E{filas_estaticas['transporte']}+E{filas_estaticas['personal_supervision']}"
            ),
        )
        c_total.font = negrita
        c_total.number_format = "$ #,##0"
        c_total.border = borde_top
        fila_total_apu[codigo] = fila
        fila += 3

    for col, ancho in zip("ABCDE", (50, 10, 12, 14, 14)):
        ws_apus.column_dimensions[col].width = ancho

    # -------------------------------------------------------------
    # 4. Hoja LISTA DE MATERIALES: detalle por item + resumen por
    #    insumo (cantidad total a comprar), con SUMIF sobre el detalle.
    # -------------------------------------------------------------
    ws_lista = wb.create_sheet("LISTA DE MATERIALES")
    ws_lista.append(
        [
            "Capitulo", "Item", "Cantidad presup.", "Insumo", "Descripcion",
            "Unidad", "Consumo x unidad", "Cantidad requerida", "Precio", "Subtotal",
        ]
    )
    for c in ws_lista[1]:
        c.font = negrita
        c.fill = relleno_encabezado
    fila = 2
    fila_detalle_ini = fila
    for nombre_cap, filas_cap in por_capitulo.items():
        for it in filas_cap:
            codigo = it.get("apu_codigo")
            if not codigo or (not insumos_por_apu.get(codigo) and not fijos_por_apu.get(codigo)):
                continue
            cant_presupuesto = float(it["cantidad"])
            fila_item_cantidad = fila
            primero = True
            filas_material = list(insumos_por_apu.get(codigo, []))
            for ins in filas_material:
                fref = fila_insumo.get(ins["insumo_codigo"])
                ws_lista.cell(row=fila, column=1, value=nombre_cap)
                ws_lista.cell(row=fila, column=2, value=f"{codigo} - {it['descripcion_snapshot'][:45]}")
                ws_lista.cell(
                    row=fila, column=3,
                    value=cant_presupuesto if primero else f"=C{fila_item_cantidad}",
                )
                ws_lista.cell(row=fila, column=4, value=ins["insumo_codigo"])
                ws_lista.cell(row=fila, column=5, value=ins["descripcion"])
                ws_lista.cell(row=fila, column=6, value=ins["unidad"])
                ws_lista.cell(row=fila, column=7, value=ins["cantidad"])
                ws_lista.cell(row=fila, column=8, value=f"=C{fila}*G{fila}")
                ws_lista.cell(row=fila, column=9, value=f"=BASE!D{fref}" if fref else ins["precio"])
                ws_lista.cell(row=fila, column=10, value=f"=H{fila}*I{fila}")
                ws_lista.cell(row=fila, column=9).number_format = "$ #,##0.00"
                ws_lista.cell(row=fila, column=10).number_format = "$ #,##0"
                primero = False
                fila += 1
            for fijo in fijos_por_apu.get(codigo, []):
                ws_lista.cell(row=fila, column=1, value=nombre_cap)
                ws_lista.cell(row=fila, column=2, value=f"{codigo} - {it['descripcion_snapshot'][:45]}")
                ws_lista.cell(
                    row=fila, column=3,
                    value=cant_presupuesto if primero else f"=C{fila_item_cantidad}",
                )
                ws_lista.cell(row=fila, column=4, value="(fijo, no es insumo)")
                ws_lista.cell(row=fila, column=5, value=fijo["descripcion"])
                ws_lista.cell(row=fila, column=6, value=fijo["unidad"])
                ws_lista.cell(row=fila, column=7, value=fijo["cantidad"])
                ws_lista.cell(row=fila, column=8, value=f"=C{fila}*G{fila}")
                ws_lista.cell(row=fila, column=9, value=fijo["precio_unitario"])
                ws_lista.cell(row=fila, column=10, value=f"=H{fila}*I{fila}")
                ws_lista.cell(row=fila, column=9).number_format = "$ #,##0.00"
                ws_lista.cell(row=fila, column=10).number_format = "$ #,##0"
                primero = False
                fila += 1
    fila_detalle_fin = fila - 1

    fila += 2
    ws_lista.cell(row=fila, column=1, value="RESUMEN POR INSUMO (cantidad total a comprar)").font = negrita_grande
    fila += 2
    ws_lista.append(["Insumo", "Descripcion", "Unidad", "Cantidad total", "Precio", "Subtotal"])
    for c in ws_lista[fila]:
        c.font = negrita
        c.fill = relleno_encabezado
    fila += 1
    fila_resumen_ini = fila
    for codigo_ins in insumos_usados:
        info = insumos_info[codigo_ins]
        fref = fila_insumo.get(codigo_ins)
        ws_lista.cell(row=fila, column=1, value=codigo_ins)
        ws_lista.cell(row=fila, column=2, value=info["descripcion"])
        ws_lista.cell(row=fila, column=3, value=info["unidad"])
        if fila_detalle_fin >= fila_detalle_ini:
            ws_lista.cell(
                row=fila, column=4,
                value=f"=SUMIF(D{fila_detalle_ini}:D{fila_detalle_fin},A{fila},H{fila_detalle_ini}:H{fila_detalle_fin})",
            )
        else:
            ws_lista.cell(row=fila, column=4, value=0)
        ws_lista.cell(row=fila, column=5, value=f"=BASE!D{fref}" if fref else info["precio"])
        ws_lista.cell(row=fila, column=6, value=f"=D{fila}*E{fila}")
        ws_lista.cell(row=fila, column=5).number_format = "$ #,##0.00"
        ws_lista.cell(row=fila, column=6).number_format = "$ #,##0"
        fila += 1
    fila_resumen_fin = fila - 1
    ws_lista.cell(row=fila, column=1, value="TOTAL MATERIALES A COMPRAR").font = negrita
    if fila_resumen_fin >= fila_resumen_ini:
        ws_lista.cell(row=fila, column=6, value=f"=SUM(F{fila_resumen_ini}:F{fila_resumen_fin})")
    else:
        ws_lista.cell(row=fila, column=6, value=0)
    ws_lista.cell(row=fila, column=6).font = negrita
    ws_lista.cell(row=fila, column=6).number_format = "$ #,##0"
    fila_total_materiales = fila

    for col, ancho in zip("ABCDEFGHIJ", (18, 45, 16, 16, 30, 10, 16, 16, 12, 14)):
        ws_lista.column_dimensions[col].width = ancho

    # -------------------------------------------------------------
    # 5. Hoja CRONOGRAMA: reparte cada capitulo en el plazo, EN DIAS
    #    (regla de negocio 2026-08-03: ya no se reparte por mes ni por
    #    semana). Dos formas de repartir cada dia dentro de la duracion
    #    del capitulo:
    #      - lineal: el valor del capitulo se divide en partes iguales
    #        entre sus dias.
    #      - curva S: el avance sigue una curva lenta-rapido-lenta
    #        (smoothstep) tipica de obra, tambien calculada dia a dia.
    #    La duracion de cada capitulo sale de dias reales (cantidad/
    #    rendimiento_dia de la cuadrilla) cuando estan disponibles; si
    #    no, un reparto proporcional al peso del capitulo sobre el
    #    costo directo.
    # -------------------------------------------------------------
    ws_cron = wb.create_sheet("CRONOGRAMA")
    plazo_dias_presu = float(presupuesto.get("plazo_dias") or 0)
    if not plazo_dias_presu:
        plazo_dias_presu = float(presupuesto.get("plazo_meses") or 1) * 30
    num_dias = max(1, math.ceil(plazo_dias_presu))
    ws_cron.cell(row=1, column=1, value="CRONOGRAMA (reparto estimado por dia)").font = titulo_hoja
    ws_cron.cell(
        row=2, column=1,
        value=(
            f"Reparto {'curva S' if usar_curva_s else 'lineal'} -- duracion de cada capitulo: dias "
            "reales (cantidad/rendimiento de la cuadrilla) cuando hay cuadrilla asignada; si no, "
            "proporcional al peso del capitulo sobre el costo directo."
        ),
    )
    fila = 4
    encabezados_cron = ["Capitulo", "Subtotal", "% del total", "Duracion (dias)"] + [
        f"Dia {d}" for d in range(1, num_dias + 1)
    ]
    ws_cron.append(encabezados_cron)
    for c in ws_cron[fila]:
        c.font = negrita
        c.fill = relleno_encabezado
    fila += 1

    def _fraccion_curva_s(t):
        """Smoothstep: avance acumulado 0..1 lento-rapido-lento, la curva
        S clasica de obra. t es el avance de tiempo relativo (0..1)
        dentro de la duracion del capitulo."""
        t = min(1.0, max(0.0, t))
        return t * t * (3 - 2 * t)

    dia_actual_inicio = 0.0
    filas_capitulos_cron = []
    for nombre_cap, filas_cap in por_capitulo.items():
        subtotal_cap = sum(float(x["cantidad"]) * float(x["precio_unitario_snapshot"]) for x in filas_cap)
        peso = (subtotal_cap / costo_directo) if costo_directo else 0

        dias_reales = 0.0
        todos_con_rendimiento = bool(filas_cap)
        for it in filas_cap:
            codigo = it.get("apu_codigo")
            det = detalle_apus.get(codigo) if codigo else None
            rendimiento = float(det.get("rendimiento_dia") or 0) if det else 0
            if det and det.get("cuadrilla_codigo") and rendimiento:
                dias_reales += float(it["cantidad"]) / rendimiento
            else:
                todos_con_rendimiento = False
        if todos_con_rendimiento and dias_reales > 0:
            duracion_dias = round(dias_reales, 2)
        else:
            duracion_dias = round(max(peso * num_dias, 1.0), 2)

        dia_inicio = dia_actual_inicio
        dia_fin = dia_inicio + duracion_dias
        dia_actual_inicio = dia_fin

        fila_cap_row = fila
        ws_cron.cell(row=fila, column=1, value=nombre_cap)
        ws_cron.cell(row=fila, column=2, value=subtotal_cap).number_format = "$ #,##0"
        ws_cron.cell(row=fila, column=3, value=f"=B{fila}/{costo_directo}" if costo_directo else 0)
        ws_cron.cell(row=fila, column=3).number_format = "0.0%"
        ws_cron.cell(row=fila, column=4, value=duracion_dias)

        if usar_curva_s:
            acumulado_anterior = 0.0
            for d in range(1, num_dias + 1):
                t = (d - dia_inicio) / duracion_dias if duracion_dias else 0
                acumulado_hoy = _fraccion_curva_s(t)
                valor_dia = round((acumulado_hoy - acumulado_anterior) * subtotal_cap, 2)
                acumulado_anterior = acumulado_hoy
                ws_cron.cell(row=fila, column=4 + d, value=valor_dia).number_format = "$ #,##0"
        else:
            for d in range(1, num_dias + 1):
                # fraccion de este dia que cae dentro de [dia_inicio, dia_fin)
                solape = max(0.0, min(d, dia_fin) - max(d - 1, dia_inicio))
                fraccion = min(1.0, max(0.0, solape))
                valor_dia = round(fraccion * (subtotal_cap / duracion_dias), 2) if duracion_dias else 0
                ws_cron.cell(row=fila, column=4 + d, value=valor_dia).number_format = "$ #,##0"
        filas_capitulos_cron.append(fila_cap_row)
        fila += 1

    fila_totales_dia = fila
    ws_cron.cell(row=fila, column=1, value="TOTAL DIA").font = negrita
    for d in range(1, num_dias + 1):
        col = 4 + d
        letra = openpyxl.utils.get_column_letter(col)
        if filas_capitulos_cron:
            ws_cron.cell(
                row=fila, column=col,
                value=f"=SUM({letra}{filas_capitulos_cron[0]}:{letra}{filas_capitulos_cron[-1]})",
            )
        else:
            ws_cron.cell(row=fila, column=col, value=0)
        ws_cron.cell(row=fila, column=col).font = negrita
        ws_cron.cell(row=fila, column=col).number_format = "$ #,##0"
    fila += 1

    fila_acumulado = fila
    ws_cron.cell(row=fila, column=1, value="ACUMULADO (%)")
    for d in range(1, num_dias + 1):
        col = 4 + d
        letra = openpyxl.utils.get_column_letter(col)
        letra_ant = openpyxl.utils.get_column_letter(col - 1) if d > 1 else None
        if costo_directo:
            if d == 1:
                formula = f"={letra}{fila_totales_dia}/{costo_directo}"
            else:
                formula = f"={letra_ant}{fila}+{letra}{fila_totales_dia}/{costo_directo}"
        else:
            formula = 0
        ws_cron.cell(row=fila, column=col, value=formula)
        ws_cron.cell(row=fila, column=col).number_format = "0.0%"

    ws_cron.column_dimensions["A"].width = 32
    for col in range(2, 5 + num_dias):
        ws_cron.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 14

    # -------------------------------------------------------------
    # 6. Hoja FLUJO DE CAJA: costos diarios (del cronograma) vs
    #    ingresos segun las condiciones de pago del presupuesto --
    #    tambien dia a dia, no por mes/semana.
    # -------------------------------------------------------------
    ws_flujo = wb.create_sheet("FLUJO DE CAJA")
    ws_flujo.cell(row=1, column=1, value="FLUJO DE CAJA (estimado, diario)").font = titulo_hoja
    ws_flujo.cell(
        row=2, column=1,
        value=(
            "Cada ingreso se descompone proporcionalmente en Costo Directo, "
            "Administracion, Imprevistos, Utilidad e IVA sobre Utilidad, segun "
            "el peso de cada componente en el valor total del contrato -- base "
            "para el control de costos."
        ),
    ).font = Font(italic=True, size=9)

    # columnas: A Dia | B Costos | C Ingresos | D Costo Directo (del ingreso)
    # | E Administracion | F Imprevistos | G Utilidad | H IVA Utilidad
    # | I Flujo neto | J Acumulado
    fila = 4
    ws_flujo.append([
        "Dia", "Costos", "Ingresos", "Costo Directo", "Administracion",
        "Imprevistos", "Utilidad", "IVA Utilidad", "Flujo neto", "Acumulado",
    ])
    for c in ws_flujo[fila]:
        c.font = negrita
        c.fill = relleno_encabezado
    fila += 1
    fila_ini_flujo = fila

    avance2_pct = float(presupuesto.get("avance2_pct") or 0.5)
    anticipo_pct = float(presupuesto.get("anticipo_pct") or 0.45)
    pago2_pct = float(presupuesto.get("pago2_pct") or 0.45)
    pagofin_pct = float(presupuesto.get("pagofin_pct") or 0.10)
    valor_total = aiu["valor_total"]

    # proporcion fija de cada componente dentro del valor total del
    # contrato -- se aplica a cualquier ingreso (anticipo, pago
    # intermedio, pago final) porque el AIU es un % constante sobre el
    # costo directo.
    ratio_cd = (costo_directo / valor_total) if valor_total else 0
    ratio_admin = (aiu["administracion"] / valor_total) if valor_total else 0
    ratio_imprev = (aiu["imprevistos"] / valor_total) if valor_total else 0
    ratio_util = (aiu["utilidad"] / valor_total) if valor_total else 0
    ratio_iva_util = (aiu["iva_utilidad"] / valor_total) if valor_total else 0

    for d in range(1, num_dias + 1):
        col_cron = 4 + d
        letra_acum_cron = openpyxl.utils.get_column_letter(col_cron)
        letra_costo_dia = openpyxl.utils.get_column_letter(col_cron)
        fila_r = fila_ini_flujo + d - 1
        ws_flujo.cell(row=fila_r, column=1, value=d)
        ws_flujo.cell(
            row=fila_r, column=2,
            value=f"=CRONOGRAMA!{letra_costo_dia}{fila_totales_dia}",
        )
        ws_flujo.cell(row=fila_r, column=2).number_format = "$ #,##0"

        # ingreso: anticipo el dia 1; saldo final el ultimo dia; el pago
        # intermedio se paga el primer dia en que el acumulado del
        # cronograma alcanza el % de avance pactado.
        condiciones = []
        if d == 1:
            condiciones.append(f"{anticipo_pct}*{valor_total}")
        condiciones.append(
            f'IF(AND(CRONOGRAMA!{letra_acum_cron}{fila_acumulado}>={avance2_pct},'
            f'OR({d}=1,CRONOGRAMA!{openpyxl.utils.get_column_letter(col_cron-1) if d>1 else letra_acum_cron}{fila_acumulado}<{avance2_pct})),'
            f"{pago2_pct}*{valor_total},0)"
        )
        if d == num_dias:
            condiciones.append(f"{pagofin_pct}*{valor_total}")
        formula_ingreso = "=" + "+".join(condiciones)
        ws_flujo.cell(row=fila_r, column=3, value=formula_ingreso)
        ws_flujo.cell(row=fila_r, column=3).number_format = "$ #,##0"

        # desglose proporcional del ingreso del dia (columna C) en los
        # componentes del AIU -- clave para el control de costos.
        for col_destino, ratio in (
            (4, ratio_cd),
            (5, ratio_admin),
            (6, ratio_imprev),
            (7, ratio_util),
            (8, ratio_iva_util),
        ):
            ws_flujo.cell(row=fila_r, column=col_destino, value=f"=C{fila_r}*{ratio}")
            ws_flujo.cell(row=fila_r, column=col_destino).number_format = "$ #,##0"

        ws_flujo.cell(row=fila_r, column=9, value=f"=C{fila_r}-B{fila_r}")
        ws_flujo.cell(row=fila_r, column=9).number_format = "$ #,##0"
        if d == 1:
            ws_flujo.cell(row=fila_r, column=10, value=f"=I{fila_r}")
        else:
            ws_flujo.cell(row=fila_r, column=10, value=f"=J{fila_r - 1}+I{fila_r}")
        ws_flujo.cell(row=fila_r, column=10).number_format = "$ #,##0"

    fila_fin_flujo = fila_ini_flujo + num_dias - 1
    fila_total_flujo = fila_fin_flujo + 2
    ws_flujo.cell(row=fila_total_flujo, column=1, value="TOTAL").font = negrita
    for col in range(2, 10):
        letra_col = openpyxl.utils.get_column_letter(col)
        if col == 9:
            continue
        ws_flujo.cell(
            row=fila_total_flujo, column=col,
            value=f"=SUM({letra_col}{fila_ini_flujo}:{letra_col}{fila_fin_flujo})",
        ).font = negrita
        ws_flujo.cell(row=fila_total_flujo, column=col).number_format = "$ #,##0"

    for col, ancho in zip("ABCDEFGHIJ", (8, 16, 16, 16, 16, 16, 16, 16, 16, 16)):
        ws_flujo.column_dimensions[col].width = ancho

    # -------------------------------------------------------------
    # 7. Hoja DASHBOARD: resumen ejecutivo + graficas nativas de Excel
    #    (se recalculan solas porque referencian celdas de las otras
    #    hojas, no imagenes fijas).
    # -------------------------------------------------------------
    ws_dash = wb.create_sheet("DASHBOARD", 0)
    ws_dash.cell(row=1, column=1, value=f"{presupuesto.get('proyecto', '')}").font = Font(bold=True, size=16)
    ws_dash.cell(row=2, column=1, value=f"Cliente: {presupuesto.get('cliente', '')}")
    ws_dash.cell(row=3, column=1, value=f"Estado: {(presupuesto.get('estado') or 'borrador').upper()}")

    fila = 5
    resumen = [
        ("Costo directo", costo_directo),
        ("Administracion", aiu["administracion"]),
        ("Imprevistos", aiu["imprevistos"]),
        ("Utilidad", aiu["utilidad"]),
        ("IVA utilidad", aiu["iva_utilidad"]),
        ("Valor del contrato", aiu["valor_total"]),
        ("Total materiales a comprar", f"=('LISTA DE MATERIALES'!F{fila_total_materiales})"),
        ("Plazo (dias)", num_dias),
    ]
    for etiqueta, valor in resumen:
        ws_dash.cell(row=fila, column=1, value=etiqueta).font = negrita
        c = ws_dash.cell(row=fila, column=2, value=valor)
        if etiqueta != "Plazo (dias)":
            c.number_format = "$ #,##0"
        fila += 1

    # tabla de apoyo para las graficas: capitulo + subtotal
    fila_tabla_cap = fila + 2
    ws_dash.cell(row=fila_tabla_cap - 1, column=1, value="Costo por capitulo").font = negrita
    ws_dash.append  # noop, mantenemos append manual mas abajo
    f2 = fila_tabla_cap
    ws_dash.cell(row=f2, column=1, value="Capitulo")
    ws_dash.cell(row=f2, column=2, value="Subtotal")
    for nombre_cap, filas_cap in por_capitulo.items():
        f2 += 1
        subtotal_cap = sum(float(x["cantidad"]) * float(x["precio_unitario_snapshot"]) for x in filas_cap)
        ws_dash.cell(row=f2, column=1, value=nombre_cap)
        ws_dash.cell(row=f2, column=2, value=subtotal_cap).number_format = "$ #,##0"
    fila_tabla_cap_fin = f2

    grafico_cap = BarChart()
    grafico_cap.title = "Costo por capitulo"
    datos = Reference(ws_dash, min_col=2, min_row=fila_tabla_cap, max_row=fila_tabla_cap_fin)
    categorias = Reference(ws_dash, min_col=1, min_row=fila_tabla_cap + 1, max_row=fila_tabla_cap_fin)
    grafico_cap.add_data(datos, titles_from_data=True)
    grafico_cap.set_categories(categorias)
    grafico_cap.height, grafico_cap.width = 8, 16
    ws_dash.add_chart(grafico_cap, f"D{fila_tabla_cap}")

    grafico_curva = LineChart()
    grafico_curva.title = "Curva de avance acumulado % (diaria)"
    datos_curva = Reference(
        ws_dash.parent["CRONOGRAMA"], min_col=5, max_col=4 + num_dias, min_row=fila_acumulado
    )
    grafico_curva.add_data(datos_curva, titles_from_data=False)
    grafico_curva.height, grafico_curva.width = 8, 16
    ws_dash.add_chart(grafico_curva, f"D{fila_tabla_cap + 18}")

    ws_dash.column_dimensions["A"].width = 26
    ws_dash.column_dimensions["B"].width = 18

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------
# Fase 3: importar un presupuesto ya armado desde el Excel maestro
# (hoja "PRESUPUESTO" de MODULO_PRESUPUESTOS_COMSAS_V4.xlsm, o cualquier
# copia con las mismas columnas: COD APU | ITEM | ACTIVIDAD | UN | CANT |
# VR UNITARIO | VR TOTAL). Misma logica de deteccion de capitulos que
# EscribirResumenCapitulos en GENERAR_PROPUESTA.txt: una fila es
# "capitulo" cuando tiene ACTIVIDAD pero NO tiene UN.
# ---------------------------------------------------------------------
def cargar_libro_excel(archivo):
    import openpyxl

    return openpyxl.load_workbook(archivo, data_only=True)


def leer_datos_generales_excel(wb):
    """Lee cliente/proyecto/ubicacion/plazo de la hoja INICIO (y CRONOGRAMA
    para el plazo), igual que GENERAR_PROPUESTA.txt (wi.Range C5/C6/C8,
    wc.Range F2). Devuelve valores vacios si no encuentra la hoja o celda,
    para que el usuario los revise/complete a mano antes de crear."""
    datos = {"cliente": "", "proyecto": "", "ubicacion": "", "plazo_meses": 0.0}
    if "INICIO" in wb.sheetnames:
        wi = wb["INICIO"]
        datos["proyecto"] = str(wi["C5"].value or "").strip()
        datos["cliente"] = str(wi["C6"].value or "").strip()
        datos["ubicacion"] = str(wi["C8"].value or "").strip()
    if "CRONOGRAMA" in wb.sheetnames:
        valor_f2 = wb["CRONOGRAMA"]["F2"].value
        if isinstance(valor_f2, (int, float)) and valor_f2 > 0:
            datos["plazo_meses"] = round(valor_f2 / 30, 2)  # F2 esta en dias habiles
    return datos


def parsear_excel_presupuesto(wb, hoja=None):
    if hoja and hoja in wb.sheetnames:
        ws = wb[hoja]
    elif "PRESUPUESTO" in wb.sheetnames:
        ws = wb["PRESUPUESTO"]
    else:
        ws = wb.worksheets[0]

    fila_header = None
    for r in range(1, min(ws.max_row, 20) + 1):
        val = ws.cell(row=r, column=1).value
        if val and "COD APU" in str(val).upper():
            fila_header = r
            break
    if fila_header is None:
        fila_header = 4  # respaldo: fila conocida del template estandar

    filas_parseadas = []
    capitulo_actual = None
    for r in range(fila_header + 1, ws.max_row + 1):
        codigo = ws.cell(row=r, column=1).value
        actividad = ws.cell(row=r, column=3).value
        unidad = ws.cell(row=r, column=4).value
        cantidad = ws.cell(row=r, column=5).value
        vr_unitario = ws.cell(row=r, column=6).value

        cname = str(actividad).strip() if actividad is not None else ""
        un = str(unidad).strip() if unidad is not None else ""

        if "COSTOS DIRECTOS" in cname.upper():
            break
        if not cname:
            continue
        if "SUBTOTAL" in cname.upper() or "TOTAL" in cname.upper():
            continue

        if not un:
            capitulo_actual = cname
            continue

        if capitulo_actual is None:
            capitulo_actual = "SIN CAPITULO"

        filas_parseadas.append(
            {
                "capitulo": capitulo_actual,
                "codigo": str(codigo).strip() if codigo is not None else "",
                "descripcion": cname,
                "unidad": un,
                "cantidad": float(cantidad) if isinstance(cantidad, (int, float)) else 0.0,
                "precio_unitario": float(vr_unitario) if isinstance(vr_unitario, (int, float)) else 0.0,
            }
        )
    return filas_parseadas


def verificar_en_catalogo(sb, codigos):
    codigos_unicos = [c for c in set(codigos) if c]
    if not codigos_unicos:
        return set()
    resultados = (
        sb.table("catalogo_apu")
        .select("codigo")
        .in_("codigo", codigos_unicos)
        .execute()
        .data
    )
    return {r["codigo"] for r in resultados}


def importar_filas_a_presupuesto(sb, presupuesto_id, filas, codigos_encontrados):
    """Crea (o reutiliza) los capitulos y agrega los items parseados de un
    Excel a un presupuesto. Devuelve cuantos items quedaron insertados."""
    capitulos_existentes = (
        sb.table("presupuesto_capitulos")
        .select("id, nombre")
        .eq("presupuesto_id", presupuesto_id)
        .execute()
        .data
    )
    capitulos_cache = {c["nombre"].upper(): c["id"] for c in capitulos_existentes}

    insertados = 0
    for f in filas:
        clave_cap = f["capitulo"].upper()
        if clave_cap not in capitulos_cache:
            nuevo_cap = (
                sb.table("presupuesto_capitulos")
                .insert(
                    {"presupuesto_id": presupuesto_id, "codigo": None, "nombre": f["capitulo"]}
                )
                .execute()
            )
            capitulos_cache[clave_cap] = nuevo_cap.data[0]["id"]

        sb.table("presupuesto_items").insert(
            {
                "presupuesto_id": presupuesto_id,
                "capitulo_id": capitulos_cache[clave_cap],
                "apu_codigo": f["codigo"] if f["codigo"] in codigos_encontrados else None,
                "descripcion_snapshot": f["descripcion"],
                "unidad_snapshot": f["unidad"],
                "cantidad": f["cantidad"],
                "precio_unitario_snapshot": f["precio_unitario"],
            }
        ).execute()
        insertados += 1
    return insertados


# ---------------------------------------------------------------------
# Fase 4: mantenimiento de precios
# ---------------------------------------------------------------------
def calcular_vigencia(fecha_cotizacion, origen_precio):
    """Misma logica de la formula de PrepararSemaforoPrecios (Excel):
    SIN FECHA / ESTIMADO (ICOCED) / VENCIDO (>180 dias) / POR VENCER
    (>90 dias) / VIGENTE."""
    if origen_precio == "ICOCED":
        return "ESTIMADO"
    if not fecha_cotizacion:
        return "SIN FECHA"
    if isinstance(fecha_cotizacion, str):
        fecha_cotizacion = date.fromisoformat(fecha_cotizacion[:10])
    dias = (date.today() - fecha_cotizacion).days
    if dias > 180:
        return "VENCIDO"
    if dias > 90:
        return "POR VENCER"
    return "VIGENTE"


def obtener_insumos_para_semaforo(sb):
    """El semaforo de vigencia vive a nivel de INSUMO (hoja MATERIALES del
    Excel) -- catalogo_apu.materiales es solo el agregado por APU."""
    query = (
        sb.table("insumos")
        .select("codigo, descripcion, unidad, precio, proveedor, fecha_cotizacion, origen_precio, activo")
        .eq("activo", True)
        .order("codigo")
    )
    filas = _fetch_todas_las_filas(query)
    for f in filas:
        f["vigencia"] = calcular_vigencia(f.get("fecha_cotizacion"), f.get("origen_precio"))
    return filas


def recalcular_materiales_apus_afectados(sb, insumo_codigos, origen, usuario_nota=None):
    """Tras cambiar el precio de uno o mas insumos, recalcula
    catalogo_apu.materiales para todos los APUs que los usan (via
    apu_insumos + apu_materiales_fijos, con la vista v_apu_materiales_calculado)
    y registra cada cambio de APU en la bitacora. Devuelve cuantos APUs
    quedaron actualizados."""
    if not insumo_codigos:
        return 0
    apus_afectados = (
        sb.table("apu_insumos")
        .select("apu_codigo")
        .in_("insumo_codigo", list(insumo_codigos))
        .execute()
        .data
    )
    codigos_apu = sorted({a["apu_codigo"] for a in apus_afectados})
    if not codigos_apu:
        return 0

    calculado = (
        sb.table("v_apu_materiales_calculado")
        .select("apu_codigo, materiales_calculado")
        .in_("apu_codigo", codigos_apu)
        .execute()
        .data
    )
    nuevo_por_apu = {c["apu_codigo"]: round(float(c["materiales_calculado"] or 0), 0) for c in calculado}

    actuales = (
        sb.table("catalogo_apu")
        .select("codigo, materiales")
        .in_("codigo", codigos_apu)
        .execute()
        .data
    )
    viejo_por_apu = {a["codigo"]: float(a["materiales"] or 0) for a in actuales}

    aplicados = 0
    for codigo in codigos_apu:
        nuevo = nuevo_por_apu.get(codigo)
        viejo = viejo_por_apu.get(codigo, 0.0)
        if nuevo is None or nuevo == viejo:
            continue
        sb.table("catalogo_apu").update({"materiales": nuevo}).eq("codigo", codigo).execute()
        sb.table("bitacora_precios").insert(
            {
                "apu_codigo": codigo,
                "campo": "materiales",
                "valor_anterior": viejo,
                "valor_nuevo": nuevo,
                "origen": origen,
                "usuario": _usuario_bitacora(),
            }
        ).execute()
        aplicados += 1
    return aplicados


def es_escalable_icoced_insumo(fila, mes_actual):
    """Un insumo es candidato a escalamiento ICOCED si tiene precio > 0 y
    no fue refrescado ya este mismo mes (misma regla que EsEscalable en
    la macro: no toca lo ya fresco este mes)."""
    precio = float(fila.get("precio") or 0)
    if precio <= 0:
        return False
    fecha_cot = fila.get("fecha_cotizacion")
    if fecha_cot:
        if isinstance(fecha_cot, str):
            fecha_cot = date.fromisoformat(fecha_cot[:10])
        if (fecha_cot.year, fecha_cot.month) == mes_actual:
            return False
    return True


def aplicar_escalamiento_icoced(sb, insumos_candidatos, pct):
    """Sube el precio de cada insumo candidato un pct%, estampa
    fecha_cotizacion=hoy y origen_precio='ICOCED', registra en bitacora
    (a nivel de insumo) -- igual que EscalarPorICOCED en la macro -- y
    recalcula materiales de los APUs que usan esos insumos."""
    hoy = date.today().isoformat()
    aplicados = 0
    codigos_tocados = []
    for fila in insumos_candidatos:
        viejo = float(fila["precio"] or 0)
        nuevo = round(viejo * (1 + pct / 100), 0)
        sb.table("insumos").update(
            {"precio": nuevo, "fecha_cotizacion": hoy, "origen_precio": "ICOCED"}
        ).eq("codigo", fila["codigo"]).execute()
        sb.table("bitacora_precios").insert(
            {
                "insumo_codigo": fila["codigo"],
                "campo": "precio",
                "valor_anterior": viejo,
                "valor_nuevo": nuevo,
                "origen": "ICOCED",
                "usuario": _usuario_bitacora(),
            }
        ).execute()
        aplicados += 1
        codigos_tocados.append(fila["codigo"])
    apus_recalculados = recalcular_materiales_apus_afectados(sb, codigos_tocados, "ICOCED")
    return aplicados, apus_recalculados


def buscar_variacion_icoced_dane():
    """Intenta descargar el anexo mensual del ICOCED (DANE, en
    https://www.dane.gov.co/.../anex-ICOCED-<mes><anio>.xlsx) y ubicar la
    variacion mensual del grupo de costo 'Materiales'.

    IMPORTANTE: el DANE publica esto en un Excel cuyo diseño puede
    cambiar de un mes a otro y no hay forma de probar este parseo contra
    el archivo real desde donde se escribio este codigo (sin acceso a
    internet en ese momento) -- por eso esto es una SUGERENCIA para que
    el usuario revise y confirme, nunca se aplica solo. Si el DANE
    cambia el formato, esto puede devolver una lista vacia o un valor
    equivocado -- siempre revisa contra el boletin en pdf antes de usar
    el numero.
    """
    import io
    import urllib.error
    import urllib.request

    import openpyxl

    meses = {
        1: "ene", 2: "feb", 3: "mar", 4: "abr", 5: "may", 6: "jun",
        7: "jul", 8: "ago", 9: "sep", 10: "oct", 11: "nov", 12: "dic",
    }
    hoy = date.today()
    y, m = hoy.year, hoy.month
    intentos = []
    for _ in range(4):  # el boletin sale con ~1 mes de rezago; prueba hacia atras
        intentos.append((meses[m], y))
        m -= 1
        if m == 0:
            m, y = 12, y - 1

    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
    for mes_abr, anio in intentos:
        url = f"https://www.dane.gov.co/files/operaciones/ICOCED/anex-ICOCED-{mes_abr}{anio}.xlsx"
        try:
            req = urllib.request.Request(url, headers=headers)
            with urllib.request.urlopen(req, timeout=15) as resp:
                contenido = resp.read()
        except (urllib.error.URLError, TimeoutError, OSError):
            continue  # ese mes no esta publicado todavia (o fallo la red) -- prueba el anterior

        try:
            # sin read_only: el anexo es chico (no las ~2300 hojas del
            # maestro), y aqui SI hace falta acceso aleatorio hacia
            # arriba en cada columna para leer el encabezado.
            wb = openpyxl.load_workbook(io.BytesIO(contenido), data_only=True, read_only=False)
        except Exception:
            continue

        candidatos = []
        for nombre_hoja in wb.sheetnames:
            ws = wb[nombre_hoja]
            for fila in ws.iter_rows():
                celda_texto = None
                for celda in fila:
                    if isinstance(celda.value, str) and "material" in celda.value.strip().lower():
                        celda_texto = celda
                        break
                if celda_texto is None:
                    continue
                for otra in fila:
                    v = otra.value
                    if isinstance(v, (int, float)) and otra.column != celda_texto.column and -15 <= float(v) <= 15:
                        # busca el encabezado de esa columna mirando hacia
                        # arriba (hasta 20 filas) -- suele ser algo como
                        # "Mensual" / "Año corrido" / "Doce meses", a veces
                        # en dos niveles (grupo + periodo)
                        encabezados = []
                        for r_arriba in range(celda_texto.row - 1, max(0, celda_texto.row - 20), -1):
                            val_arriba = ws.cell(row=r_arriba, column=otra.column).value
                            if isinstance(val_arriba, str) and val_arriba.strip():
                                encabezados.append(val_arriba.strip())
                                if len(encabezados) >= 2:
                                    break
                        candidatos.append(
                            {
                                "hoja": nombre_hoja,
                                "celda": f"{openpyxl.utils.get_column_letter(otra.column)}{otra.row}",
                                "encabezado": " > ".join(reversed(encabezados)) if encabezados else "(sin encabezado detectado)",
                                "etiqueta_fila": celda_texto.value.strip(),
                                "valor": round(float(v), 2),
                            }
                        )
        if candidatos:
            return {"mes": mes_abr, "anio": anio, "url": url, "candidatos": candidatos}
    return None


# ---------------------------------------------------------------------
# Fase 4: puente de precios desde facturas (Control de Costos -> insumos)
# equivalente a ActualizarPreciosDesdeFacturas en MANTENIMIENTO_PRECIOS.txt
# ---------------------------------------------------------------------
def parsear_gastos_control_costos(archivo):
    """Lee la hoja GASTOS del Control de Costos del proyecto -- mismas
    columnas que usa la macro: A=ID, B=FECHA, F=PROVEEDOR, K=BASE,
    AI=COD INSUMO, AJ=CANT. Devuelve una fila por gasto con COD INSUMO
    y CANT validos."""
    import openpyxl

    wb = openpyxl.load_workbook(archivo, data_only=True)
    if "GASTOS" not in wb.sheetnames:
        raise ValueError("El archivo no tiene una hoja 'GASTOS'.")
    ws = wb["GASTOS"]

    gastos = []
    for r in range(4, ws.max_row + 1):
        gid = ws.cell(row=r, column=1).value
        gfecha = ws.cell(row=r, column=2).value
        gproveedor = ws.cell(row=r, column=6).value
        gbase = ws.cell(row=r, column=11).value
        gcod = ws.cell(row=r, column=35).value
        gcant = ws.cell(row=r, column=36).value

        if not gid or not gcod:
            continue
        gcod = str(gcod).strip().upper()
        if not isinstance(gcant, (int, float)) or not isinstance(gbase, (int, float)):
            continue
        if gcant <= 0 or gbase <= 0:
            continue
        if isinstance(gfecha, datetime):
            gfecha = gfecha.date()
        elif not isinstance(gfecha, date):
            gfecha = None
        gastos.append(
            {
                "id_gasto": str(gid),
                "codigo": gcod,
                "cantidad": float(gcant),
                "base": float(gbase),
                "precio_unitario": round(float(gbase) / float(gcant), 2),
                "fecha": gfecha,
                "proveedor": str(gproveedor) if gproveedor else "",
            }
        )
    return gastos


def preparar_propuestas_facturas(sb, gastos):
    """Aplica la jerarquia FACTURA (real) > COTIZADO (real) > ICOCED
    (estimado) -- un precio real mas nuevo gana a uno real mas viejo, y
    cualquier precio real gana a un ICOCED, igual que en la macro.
    Tambien descarta gastos ya aplicados antes (por id_gasto en la
    bitacora)."""
    codigos = list({g["codigo"] for g in gastos})
    insumos_actuales = (
        sb.table("insumos")
        .select("codigo, descripcion, precio, fecha_cotizacion, origen_precio")
        .in_("codigo", codigos)
        .execute()
        .data
    )
    insumos_por_codigo = {i["codigo"]: i for i in insumos_actuales}

    ya_aplicados = (
        sb.table("bitacora_precios")
        .select("valor_nuevo")
        .eq("origen", "FACTURA")
        .not_.is_("insumo_codigo", "null")
        .execute()
        .data
    )
    # nota: la bitacora no guarda el id_gasto (no hay columna 'ref' en el
    # esquema base) -- para no reaplicar el mismo gasto dos veces en la
    # misma sesion, el filtrado real ocurre por insumo+fecha (ver abajo);
    # esto es best-effort, igual que antes de tener un ID persistido.

    propuestas = []
    for g in gastos:
        insumo = insumos_por_codigo.get(g["codigo"])
        if insumo is None:
            continue  # codigo no existe en insumos -- se ignora, igual que la macro
        origen_actual = (insumo.get("origen_precio") or "").upper()
        fecha_actual = insumo.get("fecha_cotizacion")
        if fecha_actual and isinstance(fecha_actual, str):
            fecha_actual = date.fromisoformat(fecha_actual[:10])

        aplicar = True
        if origen_actual in ("FACTURA", "COTIZADO") and fecha_actual and g["fecha"]:
            if g["fecha"] < fecha_actual:
                aplicar = False
        if not aplicar:
            continue

        precio_actual = float(insumo.get("precio") or 0)
        var_pct = ((g["precio_unitario"] / precio_actual - 1) * 100) if precio_actual else 0.0
        propuestas.append(
            {
                "id_gasto": g["id_gasto"],
                "codigo": g["codigo"],
                "descripcion": insumo.get("descripcion", ""),
                "precio_actual": precio_actual,
                "precio_factura": g["precio_unitario"],
                "var_pct": var_pct,
                "fecha": g["fecha"].isoformat() if g["fecha"] else None,
                "proveedor": g["proveedor"],
            }
        )
    return propuestas


def aplicar_precios_desde_facturas(sb, propuestas):
    """Actualiza insumos.precio con FUENTE='FACTURA', registra en
    bitacora (a nivel de insumo) y recalcula materiales de los APUs
    afectados -- igual que ActualizarPreciosDesdeFacturas en la macro."""
    aplicados = 0
    codigos_tocados = []
    for p in propuestas:
        sb.table("insumos").update(
            {
                "precio": p["precio_factura"],
                "fecha_cotizacion": p["fecha"],
                "origen_precio": "FACTURA",
                "proveedor": p["proveedor"] or None,
            }
        ).eq("codigo", p["codigo"]).execute()
        sb.table("bitacora_precios").insert(
            {
                "insumo_codigo": p["codigo"],
                "campo": "precio",
                "valor_anterior": p["precio_actual"],
                "valor_nuevo": p["precio_factura"],
                "origen": "FACTURA",
                "usuario": _usuario_bitacora(),
            }
        ).execute()
        aplicados += 1
        codigos_tocados.append(p["codigo"])
    apus_recalculados = recalcular_materiales_apus_afectados(sb, codigos_tocados, "FACTURA")
    return aplicados, apus_recalculados


def obtener_cargos_personal(sb):
    return (
        sb.table("cargos_personal")
        .select("cargo, tarifa_dia, tipo, activo")
        .order("cargo")
        .execute()
        .data
    )


def crear_cargo_personal(sb, cargo, tarifa_dia, tipo):
    """Crea un cargo/rol nuevo en cargos_personal (ej. un oficio que no
    estaba entre los 16 originales). Queda disponible de inmediato para
    agregarlo a la composicion de cualquier cuadrilla."""
    sb.table("cargos_personal").insert(
        {"cargo": cargo, "tarifa_dia": float(tarifa_dia), "tipo": tipo}
    ).execute()


def calcular_costo_dia_cuadrillas(sb):
    """Costo/dia de cada cuadrilla = suma(cantidad de cada cargo * tarifa_dia
    del cargo) -- la misma formula viva que describe schema.sql para
    cuadrillas.costo_dia (calculada aqui porque no es una columna fisica).
    El costo del MAESTRO se divide entre frentes_maestro de la cuadrilla
    (1 = ejecuta, esta todo el dia en esa cuadrilla; 4, por ejemplo, =
    dirige y rota entre varios frentes) -- igual que la formula de la
    hoja CUADRILLAS del Excel maestro (CREAR_CUADRILLAS.txt)."""
    cargos = {c["cargo"]: float(c["tarifa_dia"]) for c in obtener_cargos_personal(sb)}
    frentes = {
        c["codigo"]: max(int(c.get("frentes_maestro") or 1), 1)
        for c in sb.table("cuadrillas").select("codigo, frentes_maestro").execute().data
    }
    cuadrilla_cargos = (
        sb.table("cuadrilla_cargos").select("cuadrilla_codigo, cargo, cantidad").execute().data
    )
    costo_dia = {}
    for cc in cuadrilla_cargos:
        tarifa = cargos.get(cc["cargo"], 0.0)
        subtotal = float(cc["cantidad"]) * tarifa
        if cc["cargo"] == "MAESTRO":
            subtotal = subtotal / frentes.get(cc["cuadrilla_codigo"], 1)
        costo_dia[cc["cuadrilla_codigo"]] = costo_dia.get(cc["cuadrilla_codigo"], 0.0) + subtotal
    return costo_dia


def obtener_cuadrillas(sb):
    """Lista todas las cuadrillas (activas e inactivas) con su costo/dia
    calculado en vivo -- base del CRUD de Fase 2."""
    filas = (
        sb.table("cuadrillas")
        .select("codigo, nombre, frentes_maestro, uso_sugerido, activo")
        .order("codigo")
        .execute()
        .data
    )
    costo_dia = calcular_costo_dia_cuadrillas(sb)
    for f in filas:
        f["costo_dia"] = round(costo_dia.get(f["codigo"], 0.0), 0)
    return filas


def obtener_composicion_cuadrilla(sb, codigo):
    """Cargos y cantidades de una cuadrilla, con tarifa y subtotal (el
    del MAESTRO ya dividido por frentes_maestro) para el editor."""
    cuadrilla = sb.table("cuadrillas").select("frentes_maestro").eq("codigo", codigo).execute().data
    frentes = max(int((cuadrilla[0]["frentes_maestro"] if cuadrilla else 1) or 1), 1)
    filas = (
        sb.table("cuadrilla_cargos")
        .select("cargo, cantidad")
        .eq("cuadrilla_codigo", codigo)
        .order("cargo")
        .execute()
        .data
    )
    tarifas = {c["cargo"]: float(c["tarifa_dia"]) for c in obtener_cargos_personal(sb)}
    for f in filas:
        tarifa = tarifas.get(f["cargo"], 0.0)
        subtotal = float(f["cantidad"]) * tarifa
        if f["cargo"] == "MAESTRO":
            subtotal = subtotal / frentes
        f["tarifa_dia"] = tarifa
        f["subtotal"] = round(subtotal, 0)
        f["eliminar"] = False
    return filas


def crear_cuadrilla(sb, codigo, nombre, frentes_maestro, uso_sugerido):
    sb.table("cuadrillas").insert(
        {
            "codigo": codigo,
            "nombre": nombre,
            "frentes_maestro": int(frentes_maestro),
            "uso_sugerido": uso_sugerido or None,
        }
    ).execute()


def actualizar_cuadrilla(sb, codigo, nombre, frentes_maestro, uso_sugerido, activo):
    sb.table("cuadrillas").update(
        {
            "nombre": nombre,
            "frentes_maestro": int(frentes_maestro),
            "uso_sugerido": uso_sugerido or None,
            "activo": bool(activo),
        }
    ).eq("codigo", codigo).execute()


def guardar_composicion_cuadrilla(sb, codigo, df_cargos):
    """Sincroniza cuadrilla_cargos con lo editado en pantalla: borra los
    marcados como eliminar, agrega/actualiza cantidad de los demas."""
    import pandas as pd

    for _, fila in df_cargos.iterrows():
        cargo = fila.get("cargo")
        if cargo is None or (isinstance(cargo, float) and pd.isna(cargo)) or not str(cargo).strip():
            continue
        if fila.get("eliminar"):
            sb.table("cuadrilla_cargos").delete().eq("cuadrilla_codigo", codigo).eq(
                "cargo", cargo
            ).execute()
        else:
            sb.table("cuadrilla_cargos").upsert(
                {"cuadrilla_codigo": codigo, "cargo": cargo, "cantidad": float(fila["cantidad"] or 0)}
            ).execute()


def asignar_cuadrilla_apu(sb, apu_codigo, apu_original, cuadrilla_codigo):
    """Asigna, cambia o quita la cuadrilla de un APU. Si queda con
    cuadrilla y el APU ya tiene rendimiento_dia, recalcula
    mano_obra = costo_dia_cuadrilla / rendimiento_dia (igual que el
    recosteo de Fase 4) y lo deja en la bitacora -- equivalente a
    ASIGNAR_CUADRILLAS.txt / CAMBIAR_CUADRILLA.txt del Excel maestro."""
    cambios = {"cuadrilla_codigo": cuadrilla_codigo or None}
    rendimiento = float(apu_original.get("rendimiento_dia") or 0)
    mano_obra_anterior = float(apu_original.get("mano_obra") or 0)
    if cuadrilla_codigo and rendimiento:
        costo_dia = calcular_costo_dia_cuadrillas(sb).get(cuadrilla_codigo, 0.0)
        cambios["mano_obra"] = round(costo_dia / rendimiento, 0)
    sb.table("catalogo_apu").update(cambios).eq("codigo", apu_codigo).execute()
    if "mano_obra" in cambios:
        _registrar_cambio_apu(sb, apu_codigo, "mano_obra", mano_obra_anterior, cambios["mano_obra"])


def obtener_apus_con_cuadrilla(sb):
    return (
        sb.table("catalogo_apu")
        .select("codigo, descripcion, categoria, mano_obra, cuadrilla_codigo, rendimiento_dia")
        .not_.is_("cuadrilla_codigo", "null")
        .eq("es_candidato", False)
        .order("codigo")
        .execute()
        .data
    )


def calcular_impacto_recosteo(sb):
    """Recalcula mano_obra = costo_dia_cuadrilla / rendimiento_dia para
    cada APU que ya tenga cuadrilla_codigo y rendimiento_dia asignados
    -- equivalente a RECOSTEAR_MO, pero solo para lo que la Fase 2 ya
    permite enlazar (los demas APUs quedan igual, se corrigen a mano)."""
    costo_dia = calcular_costo_dia_cuadrillas(sb)
    filas = obtener_apus_con_cuadrilla(sb)
    impacto = []
    for f in filas:
        rendimiento = float(f.get("rendimiento_dia") or 0)
        cd = costo_dia.get(f["cuadrilla_codigo"])
        if not rendimiento or cd is None:
            continue
        viejo = float(f.get("mano_obra") or 0)
        nuevo = round(cd / rendimiento, 0)
        var_pct = ((nuevo - viejo) / viejo * 100) if viejo else 0.0
        impacto.append(
            {
                "codigo": f["codigo"],
                "categoria": f["categoria"],
                "descripcion": f["descripcion"],
                "mano_obra_viejo": viejo,
                "mano_obra_nuevo": nuevo,
                "var_pct": var_pct,
            }
        )
    return impacto


def aplicar_recosteo_mano_obra(sb, impacto):
    aplicados = 0
    for fila in impacto:
        if fila["mano_obra_nuevo"] == fila["mano_obra_viejo"]:
            continue
        sb.table("catalogo_apu").update({"mano_obra": fila["mano_obra_nuevo"]}).eq(
            "codigo", fila["codigo"]
        ).execute()
        sb.table("bitacora_precios").insert(
            {
                "apu_codigo": fila["codigo"],
                "campo": "mano_obra",
                "valor_anterior": fila["mano_obra_viejo"],
                "valor_nuevo": fila["mano_obra_nuevo"],
                "origen": "MANUAL",
                "usuario": _usuario_bitacora(),
            }
        ).execute()
        aplicados += 1
    return aplicados


def obtener_bitacora(sb, limite=200):
    return (
        sb.table("bitacora_precios")
        .select("*")
        .order("creado_en", desc=True)
        .limit(limite)
        .execute()
        .data
    )


def obtener_parametro(sb, clave, default=""):
    """Lee un valor de la tabla parametros (parte del esquema desde la
    Fase 0). Se usa aqui para recordar el ultimo % de ICOCED usado, para
    no tener que escribirlo desde cero cada vez."""
    fila = sb.table("parametros").select("valor").eq("clave", clave).execute().data
    return fila[0]["valor"] if fila else default


def guardar_parametro(sb, clave, valor, descripcion=None):
    registro = {"clave": clave, "valor": str(valor)}
    if descripcion is not None:
        registro["descripcion"] = descripcion
    sb.table("parametros").upsert(registro).execute()


# ---------------------------------------------------------------------
# Editor de receta de APU (que insumos usa, cantidades, rendimiento,
# equipo/transporte). Esto era la parte de "gestion de la biblioteca
# maestra de APUs" que el plan original dejaba en el Excel -- ahora que
# la Fase 4 ya trajo el desglose por insumo, se puede editar aqui.
# ---------------------------------------------------------------------
def obtener_categorias_apu(sb):
    """Lista de categorias (prefijo del codigo, ej. 'PRE', 'CIM', 'PIS')
    para poder filtrar antes de buscar -- son ~2300 APUs en total, muy
    dificil de recordar por codigo, mas facil por capitulo."""
    filas = sb.table("catalogo_apu").select("categoria").eq("es_candidato", False).limit(5000).execute().data
    return sorted({f["categoria"] for f in filas if f.get("categoria")})


def buscar_apus(sb, texto=None, categoria=None, limite=300):
    query = sb.table("catalogo_apu").select(
        "codigo, categoria, descripcion, unidad, equipo, materiales, transporte, "
        "mano_obra, personal_supervision, total, rendimiento_dia, cuadrilla_codigo, es_candidato"
    ).eq("es_candidato", False)
    if categoria:
        query = query.eq("categoria", categoria)
    if texto:
        query = query.or_(f"codigo.ilike.%{texto}%,descripcion.ilike.%{texto}%")
    return query.order("codigo").limit(limite).execute().data


def obtener_apu_detalle(sb, codigo):
    filas = sb.table("catalogo_apu").select("*").eq("codigo", codigo).execute().data
    return filas[0] if filas else None


def obtener_insumos_de_apu(sb, codigo):
    filas = (
        sb.table("apu_insumos")
        .select("insumo_codigo, cantidad, insumos(descripcion, unidad, precio)")
        .eq("apu_codigo", codigo)
        .order("insumo_codigo")
        .execute()
        .data
    )
    salida = []
    for f in filas:
        insumo = f.get("insumos") or {}
        precio = float(insumo.get("precio") or 0)
        cantidad = float(f["cantidad"] or 0)
        salida.append(
            {
                "insumo_codigo": f["insumo_codigo"],
                "descripcion": insumo.get("descripcion", ""),
                "unidad": insumo.get("unidad", ""),
                "precio": precio,
                "cantidad": cantidad,
                "subtotal": round(cantidad * precio, 2),
                "eliminar": False,
            }
        )
    return salida


def obtener_fijos_de_apu(sb, codigo):
    filas = (
        sb.table("apu_materiales_fijos")
        .select("id, descripcion, unidad, cantidad, precio_unitario")
        .eq("apu_codigo", codigo)
        .order("id")
        .execute()
        .data
    )
    for f in filas:
        f["subtotal"] = round(float(f["cantidad"] or 0) * float(f["precio_unitario"] or 0), 2)
        f["eliminar"] = False
    return filas


def obtener_equipo_de_apu(sb, codigo):
    """Lineas de equipo/herramienta del APU (seccion 'I. EQUIPO Y
    HERRAMIENTAS' del Excel) -- mismo patron que obtener_fijos_de_apu."""
    filas = (
        sb.table("apu_equipo_items")
        .select("id, descripcion, unidad, cantidad, precio_unitario")
        .eq("apu_codigo", codigo)
        .order("id")
        .execute()
        .data
    )
    for f in filas:
        f["subtotal"] = round(float(f["cantidad"] or 0) * float(f["precio_unitario"] or 0), 2)
        f["eliminar"] = False
    return filas


def obtener_transporte_de_apu(sb, codigo):
    """Lineas de transporte del APU (seccion 'III. TRANSPORTES' del
    Excel) -- mismo patron que obtener_fijos_de_apu."""
    filas = (
        sb.table("apu_transporte_items")
        .select("id, descripcion, unidad, cantidad, precio_unitario")
        .eq("apu_codigo", codigo)
        .order("id")
        .execute()
        .data
    )
    for f in filas:
        f["subtotal"] = round(float(f["cantidad"] or 0) * float(f["precio_unitario"] or 0), 2)
        f["eliminar"] = False
    return filas


def buscar_insumos_catalogo(sb, texto, limite=20):
    if not texto:
        return []
    return (
        sb.table("insumos")
        .select("codigo, descripcion, unidad, precio")
        .eq("activo", True)
        .or_(f"codigo.ilike.%{texto}%,descripcion.ilike.%{texto}%")
        .order("codigo")
        .limit(limite)
        .execute()
        .data
    )


def _siguiente_numero_codigo(codigos_existentes, prefijo):
    """Dado un listado de codigos 'PREFIJO-0001' y un prefijo, calcula el
    siguiente numero consecutivo y el ancho (cantidad de digitos) a usar
    -- toma el ancho del codigo mas reciente con ese prefijo para no
    romper el formato (ej. insumos usan 4 digitos, APUs usan 3)."""
    prefijo = prefijo.strip().upper()
    numeros = []
    ancho = 3
    for c in codigos_existentes:
        if not c:
            continue
        c = c.strip().upper()
        if "-" not in c:
            continue
        pre, _, num = c.rpartition("-")
        if pre != prefijo or not num.isdigit():
            continue
        numeros.append(int(num))
        ancho = max(ancho, len(num))
    siguiente = (max(numeros) + 1) if numeros else 1
    return siguiente, ancho


def _fetch_todas_las_filas(query):
    """Trae TODAS las filas de una consulta de Supabase paginando con
    .range(), en vez de un solo .execute() -- Supabase/PostgREST limita
    cada respuesta a 1000 filas por defecto, y con mas de 1000 filas para
    una misma tabla/prefijo (ej. mas de 1000 insumos 'MAT-') un solo
    .execute() se queda corto silenciosamente: no avisa que trunco nada,
    solo devuelve las primeras 1000 (en el orden que le de la gana al
    motor si no se pidio order(), o por el order() pedido si lo hay).
    Esto hacia que el consecutivo sugerido pisara codigos que ya existian
    mas alla de esa fila 1000, y que listados como el Semaforo de
    vigencia mostraran solo una porcion del catalogo real. Paginando de a
    1000 nos aseguramos de ver el listado completo."""
    paso = 1000
    filas = []
    inicio = 0
    while True:
        pagina = query.range(inicio, inicio + paso - 1).execute().data
        if not pagina:
            break
        filas.extend(pagina)
        if len(pagina) < paso:
            break
        inicio += paso
    return filas


def _fetch_todos_los_codigos(query):
    """Igual que _fetch_todas_las_filas pero devuelve solo el campo
    'codigo' de cada fila (para las funciones de siguiente-consecutivo,
    que solo necesitan los codigos)."""
    return [f["codigo"] for f in _fetch_todas_las_filas(query)]


def siguiente_codigo_insumo(sb, prefijo="MAT"):
    """Siguiente codigo libre para un insumo nuevo, ej. 'MAT-2201' si el
    ultimo insumo MAT- que existe es 'MAT-2200'. Evita que dos personas
    inventen el mismo codigo a mano y se sobreescriban sin darse cuenta."""
    prefijo = (prefijo or "MAT").strip().upper()
    query = sb.table("insumos").select("codigo").ilike("codigo", f"{prefijo}-%")
    codigos = _fetch_todos_los_codigos(query)
    siguiente, ancho = _siguiente_numero_codigo(codigos, prefijo)
    return f"{prefijo}-{siguiente:0{ancho}d}"


def siguiente_codigo_apu(sb, categoria):
    """Siguiente codigo libre para un APU nuevo dentro de un capitulo/
    categoria, ej. 'PIN-066' si el ultimo PIN- que existe es 'PIN-065'."""
    categoria = (categoria or "").strip().upper()
    if not categoria:
        return ""
    query = sb.table("catalogo_apu").select("codigo").eq("categoria", categoria)
    codigos = _fetch_todos_los_codigos(query)
    siguiente, ancho = _siguiente_numero_codigo(codigos, categoria)
    return f"{categoria}-{siguiente:0{ancho}d}"


def crear_insumo_nuevo(sb, codigo, descripcion, unidad, precio, proveedor=None, marcar_cotizado=True):
    """Da de alta un insumo que todavia no existe en la tabla insumos --
    antes de esto, solo se podian enlazar insumos ya cargados desde el
    Excel maestro."""
    codigo = codigo.strip().upper()
    if sb.table("insumos").select("codigo").eq("codigo", codigo).execute().data:
        raise ValueError(f"Ya existe un insumo con el codigo {codigo}.")
    registro = {
        "codigo": codigo,
        "descripcion": descripcion.strip(),
        "unidad": unidad.strip(),
        "precio": float(precio),
        "activo": True,
    }
    if proveedor:
        registro["proveedor"] = proveedor.strip()
    if marcar_cotizado and precio:
        registro["origen_precio"] = "COTIZADO"
        registro["fecha_cotizacion"] = date.today().isoformat()
    sb.table("insumos").insert(registro).execute()
    sb.table("bitacora_precios").insert(
        {
            "insumo_codigo": codigo,
            "campo": "creacion",
            "valor_anterior": 0,
            "valor_nuevo": float(precio or 0),
            "origen": "CREACION",
            "usuario": _usuario_bitacora(),
        }
    ).execute()
    return codigo


def crear_apu_nuevo(sb, codigo, descripcion, unidad):
    """Da de alta un APU nuevo (vacio, sin receta todavia) -- antes de
    esto solo se podian editar APUs que ya existian en el catalogo."""
    codigo = codigo.strip().upper()
    if "-" not in codigo:
        raise ValueError("El codigo debe tener el formato CATEGORIA-NUMERO, ej. PIN-068.")
    if sb.table("catalogo_apu").select("codigo").eq("codigo", codigo).execute().data:
        raise ValueError(f"Ya existe un APU con el codigo {codigo}.")
    sb.table("catalogo_apu").insert(
        {"codigo": codigo, "descripcion": descripcion.strip(), "unidad": unidad.strip(), "es_candidato": True}
    ).execute()
    sb.table("bitacora_precios").insert(
        {
            "apu_codigo": codigo,
            "campo": "creacion",
            "valor_anterior": 0,
            "valor_nuevo": 0,
            "origen": "CREACION",
            "usuario": _usuario_bitacora(),
        }
    ).execute()
    return codigo


def duplicar_apu(sb, codigo_origen, codigo_nuevo, descripcion_nueva=None, unidad_nueva=None):
    """Crea un APU nuevo copiando la receta completa de uno que ya existe:
    cuadrilla asignada, rendimiento, equipo, transporte, y todos los
    insumos y lineas de materiales fijas. Sirve para partir de un APU
    parecido y solo ajustarle un par de cosas (una cantidad, un insumo
    distinto, etc.) en vez de armar todo desde cero. Deja registro en la
    bitacora indicando de que APU se duplico."""
    codigo_origen = codigo_origen.strip().upper()
    codigo_nuevo = codigo_nuevo.strip().upper()

    apu_original = obtener_apu_detalle(sb, codigo_origen)
    if apu_original is None:
        raise ValueError(f"El APU de origen {codigo_origen} no existe.")
    if "-" not in codigo_nuevo:
        raise ValueError("El codigo nuevo debe tener el formato CATEGORIA-NUMERO, ej. PIN-068.")
    if sb.table("catalogo_apu").select("codigo").eq("codigo", codigo_nuevo).execute().data:
        raise ValueError(f"Ya existe un APU con el codigo {codigo_nuevo}.")

    # --- 1. fila principal: copia los valores de costo/receta general ---
    registro = {
        "codigo": codigo_nuevo,
        "descripcion": (descripcion_nueva or apu_original["descripcion"]).strip(),
        "unidad": (unidad_nueva or apu_original["unidad"]).strip(),
        "equipo": float(apu_original.get("equipo") or 0),
        "transporte": float(apu_original.get("transporte") or 0),
        "mano_obra": float(apu_original.get("mano_obra") or 0),
        "personal_supervision": float(apu_original.get("personal_supervision") or 0),
        "rendimiento_dia": apu_original.get("rendimiento_dia"),
        "cuadrilla_codigo": apu_original.get("cuadrilla_codigo"),
        "es_candidato": bool(apu_original.get("es_candidato")),
    }
    sb.table("catalogo_apu").insert(registro).execute()

    # --- 2. copiar insumos de la receta (apu_insumos) ---
    insumos_origen = (
        sb.table("apu_insumos")
        .select("insumo_codigo, cantidad")
        .eq("apu_codigo", codigo_origen)
        .execute()
        .data
    )
    if insumos_origen:
        sb.table("apu_insumos").insert(
            [
                {"apu_codigo": codigo_nuevo, "insumo_codigo": f["insumo_codigo"], "cantidad": f["cantidad"]}
                for f in insumos_origen
            ]
        ).execute()

    # --- 3. copiar lineas de materiales fijas (apu_materiales_fijos) ---
    fijos_origen = (
        sb.table("apu_materiales_fijos")
        .select("descripcion, unidad, cantidad, precio_unitario")
        .eq("apu_codigo", codigo_origen)
        .execute()
        .data
    )
    if fijos_origen:
        sb.table("apu_materiales_fijos").insert(
            [
                {
                    "apu_codigo": codigo_nuevo,
                    "descripcion": f["descripcion"],
                    "unidad": f["unidad"],
                    "cantidad": f["cantidad"],
                    "precio_unitario": f["precio_unitario"],
                }
                for f in fijos_origen
            ]
        ).execute()

    # --- 3b. copiar lineas de equipo (apu_equipo_items) ---
    equipo_origen = (
        sb.table("apu_equipo_items")
        .select("descripcion, unidad, cantidad, precio_unitario")
        .eq("apu_codigo", codigo_origen)
        .execute()
        .data
    )
    if equipo_origen:
        sb.table("apu_equipo_items").insert(
            [{"apu_codigo": codigo_nuevo, **f} for f in equipo_origen]
        ).execute()

    # --- 3c. copiar lineas de transporte (apu_transporte_items) ---
    transporte_origen = (
        sb.table("apu_transporte_items")
        .select("descripcion, unidad, cantidad, precio_unitario")
        .eq("apu_codigo", codigo_origen)
        .execute()
        .data
    )
    if transporte_origen:
        sb.table("apu_transporte_items").insert(
            [{"apu_codigo": codigo_nuevo, **f} for f in transporte_origen]
        ).execute()

    # --- 4. recalcular materiales del nuevo APU desde la vista (insumos + fijos ya copiados) ---
    if insumos_origen or fijos_origen:
        calculado = (
            sb.table("v_apu_materiales_calculado")
            .select("materiales_calculado")
            .eq("apu_codigo", codigo_nuevo)
            .execute()
            .data
        )
        materiales_nuevo = round(float(calculado[0]["materiales_calculado"]), 0) if calculado else 0.0
        sb.table("catalogo_apu").update({"materiales": materiales_nuevo}).eq("codigo", codigo_nuevo).execute()

    # --- 5. bitacora: queda registrado de que APU se duplico ---
    sb.table("bitacora_precios").insert(
        {
            "apu_codigo": codigo_nuevo,
            "campo": f"creacion (duplicado de {codigo_origen})",
            "valor_anterior": 0,
            "valor_nuevo": float(apu_original.get("total") or 0),
            "origen": "CREACION",
            "usuario": _usuario_bitacora(),
        }
    ).execute()

    return codigo_nuevo


def _registrar_cambio_apu(sb, apu_codigo, campo, viejo, nuevo, origen="MANUAL"):
    if viejo is None:
        viejo = 0
    if nuevo is None:
        nuevo = 0
    if round(float(viejo), 4) == round(float(nuevo), 4):
        return
    sb.table("bitacora_precios").insert(
        {
            "apu_codigo": apu_codigo,
            "campo": campo,
            "valor_anterior": float(viejo),
            "valor_nuevo": float(nuevo),
            "origen": origen,
            "usuario": _usuario_bitacora(),
        }
    ).execute()


def _sincronizar_items_por_lineas(sb, tabla, apu_codigo, df, valor_manual_respaldo):
    """Aplica a una tabla de lineas (apu_equipo_items o
    apu_transporte_items -- misma forma que apu_materiales_fijos) los
    cambios editados en pantalla: borra las marcadas, actualiza cantidad/
    precio de las demas. Devuelve el subtotal: si quedan lineas, es la
    suma de cantidad*precio_unitario de todas; si no quedan lineas (o
    nunca se detallo por lineas), devuelve el valor manual de respaldo
    (para no perder el numero que ya tenia el APU antes de este cambio)."""
    for _, fila in df.iterrows():
        if fila.get("eliminar"):
            sb.table(tabla).delete().eq("id", int(fila["id"])).execute()
        else:
            sb.table(tabla).update(
                {"cantidad": float(fila["cantidad"]), "precio_unitario": float(fila["precio_unitario"])}
            ).eq("id", int(fila["id"])).execute()

    restantes = sb.table(tabla).select("cantidad, precio_unitario").eq("apu_codigo", apu_codigo).execute().data
    if restantes:
        return round(sum(float(r["cantidad"] or 0) * float(r["precio_unitario"] or 0) for r in restantes), 0)
    return float(valor_manual_respaldo or 0)


def guardar_receta_apu(
    sb, apu_codigo, apu_original, df_insumos, df_fijos, df_equipo, df_transporte,
    equipo_manual, transporte_manual, rendimiento_dia, mano_obra_manual,
):
    """Aplica los cambios de la receta de un APU: sincroniza apu_insumos,
    apu_materiales_fijos, apu_equipo_items y apu_transporte_items con lo
    editado en pantalla, actualiza equipo/transporte/rendimiento_dia (y
    mano_obra si no tiene cuadrilla asignada), recalcula materiales desde
    v_apu_materiales_calculado, y deja todo registrado en la bitacora.
    Equipo y transporte se calculan por lineas (igual que materiales) si
    el APU ya tiene alguna linea cargada; si no tiene ninguna, se usa el
    valor manual (compatibilidad con los APU antiguos que solo traen un
    numero suelto de equipo/transporte, sin desglose)."""
    # --- insumos: eliminar marcados, actualizar cantidad cambiada ---
    for _, fila in df_insumos.iterrows():
        if fila.get("eliminar"):
            sb.table("apu_insumos").delete().eq("apu_codigo", apu_codigo).eq(
                "insumo_codigo", fila["insumo_codigo"]
            ).execute()
        else:
            sb.table("apu_insumos").update({"cantidad": float(fila["cantidad"])}).eq(
                "apu_codigo", apu_codigo
            ).eq("insumo_codigo", fila["insumo_codigo"]).execute()

    # --- lineas fijas: eliminar marcadas, actualizar cambiadas ---
    for _, fila in df_fijos.iterrows():
        if fila.get("eliminar"):
            sb.table("apu_materiales_fijos").delete().eq("id", int(fila["id"])).execute()
        else:
            sb.table("apu_materiales_fijos").update(
                {
                    "cantidad": float(fila["cantidad"]),
                    "precio_unitario": float(fila["precio_unitario"]),
                }
            ).eq("id", int(fila["id"])).execute()

    # --- materiales: recalcular desde la vista (insumos + fijos, ya actualizados) ---
    calculado = (
        sb.table("v_apu_materiales_calculado")
        .select("materiales_calculado")
        .eq("apu_codigo", apu_codigo)
        .execute()
        .data
    )
    materiales_nuevo = round(float(calculado[0]["materiales_calculado"]), 0) if calculado else 0.0
    _registrar_cambio_apu(sb, apu_codigo, "materiales", apu_original.get("materiales"), materiales_nuevo)

    # --- equipo y transporte: por lineas si las tiene, si no el valor manual ---
    equipo_nuevo = _sincronizar_items_por_lineas(sb, "apu_equipo_items", apu_codigo, df_equipo, equipo_manual)
    transporte_nuevo = _sincronizar_items_por_lineas(
        sb, "apu_transporte_items", apu_codigo, df_transporte, transporte_manual
    )

    # --- mano de obra: si tiene cuadrilla, se recalcula sola; si no, se edita a mano ---
    cuadrilla_codigo = apu_original.get("cuadrilla_codigo")
    if cuadrilla_codigo and rendimiento_dia:
        costo_dia = calcular_costo_dia_cuadrillas(sb).get(cuadrilla_codigo)
        mano_obra_nuevo = round(costo_dia / rendimiento_dia, 0) if costo_dia else float(apu_original.get("mano_obra") or 0)
    else:
        mano_obra_nuevo = float(mano_obra_manual)
    _registrar_cambio_apu(sb, apu_codigo, "mano_obra", apu_original.get("mano_obra"), mano_obra_nuevo)

    _registrar_cambio_apu(sb, apu_codigo, "equipo", apu_original.get("equipo"), equipo_nuevo)
    _registrar_cambio_apu(sb, apu_codigo, "transporte", apu_original.get("transporte"), transporte_nuevo)
    _registrar_cambio_apu(
        sb, apu_codigo, "rendimiento_dia", apu_original.get("rendimiento_dia"), rendimiento_dia
    )

    # --- si ya quedo con algun costo real, deja de ser "candidato" (incompleto) ---
    personal_supervision = float(apu_original.get("personal_supervision") or 0)
    total_nuevo = materiales_nuevo + mano_obra_nuevo + equipo_nuevo + transporte_nuevo + personal_supervision
    actualizacion = {
        "materiales": materiales_nuevo,
        "mano_obra": mano_obra_nuevo,
        "equipo": equipo_nuevo,
        "transporte": transporte_nuevo,
        "rendimiento_dia": float(rendimiento_dia) if rendimiento_dia else None,
    }
    if total_nuevo > 0 and apu_original.get("es_candidato"):
        actualizacion["es_candidato"] = False

    sb.table("catalogo_apu").update(actualizacion).eq("codigo", apu_codigo).execute()


def calcular_plazo_dias_sugerido(sb, items):
    """Plazo sugerido = suma de los dias de cada actividad del presupuesto
    (cantidad / rendimiento_dia del APU), tratando todo como secuencial --
    la opcion simple y conservadora que se definio como regla de negocio
    el 2026-08-03. Devuelve (dias_sugeridos, [codigos sin rendimiento
    definido que no se pudieron contar])."""
    codigos = sorted({it["apu_codigo"] for it in items if it.get("apu_codigo")})
    rendimientos = {}
    if codigos:
        filas = (
            sb.table("catalogo_apu")
            .select("codigo, rendimiento_dia")
            .in_("codigo", codigos)
            .execute()
            .data
        )
        rendimientos = {f["codigo"]: f.get("rendimiento_dia") for f in filas}

    total_dias = 0.0
    sin_rendimiento = []
    for it in items:
        cod = it.get("apu_codigo")
        rend = rendimientos.get(cod) if cod else None
        if rend and float(rend) > 0:
            total_dias += float(it["cantidad"]) / float(rend)
        else:
            sin_rendimiento.append(cod or it.get("descripcion_snapshot") or "(sin codigo)")
    return round(total_dias, 1), sin_rendimiento


# ---------------------------------------------------------------------
# Estado de sesion
# ---------------------------------------------------------------------
if "presupuesto_id" not in st.session_state:
    st.session_state.presupuesto_id = None

st.markdown(
    '<h1 style="margin-bottom:0">Presupuestos <span style="color:#9C7A2E">COMSAS</span></h1>',
    unsafe_allow_html=True,
)
st.caption("Fase 1 + 2 + 3 + 4 + 5 - conectado a la base de datos real (catalogo_apu, cuadrillas, etc.)")

tab_nuevo, tab_catalogo, tab_resumen, tab_aiu, tab_precios, tab_editor_apu = st.tabs(
    [
        "1. Presupuesto",
        "2. Agregar items del catalogo",
        "3. Resumen y total",
        "4. AIU y propuesta",
        "5. Mantenimiento de precios",
        "6. Editar receta de APU",
    ]
)

# ---------------------------------------------------------------------
# TAB 1: crear o elegir presupuesto
# ---------------------------------------------------------------------
with tab_nuevo:
    st.subheader("Crear presupuesto nuevo")
    st.caption(
        "El plazo ya NO se pregunta aqui -- se calcula solo, a partir del rendimiento "
        "de las actividades que agregues, en la pestaña '3. Resumen y total' (y ahi "
        "mismo lo puedes ajustar a mano si el que calcula la app queda corto o largo)."
    )
    with st.form("form_nuevo_presupuesto"):
        col1, col2 = st.columns(2)
        with col1:
            cliente = st.text_input("Cliente")
            proyecto = st.text_input("Proyecto")
        with col2:
            ubicacion = st.text_input("Ubicacion")
        crear = st.form_submit_button("Crear presupuesto", type="primary")

    if crear:
        if not cliente or not proyecto:
            st.error("Cliente y Proyecto son obligatorios.")
        else:
            resp = (
                sb.table("presupuestos")
                .insert(
                    {
                        "cliente": cliente,
                        "proyecto": proyecto,
                        "ubicacion": ubicacion,
                        "estado": "borrador",
                    }
                )
                .execute()
            )
            nuevo_id = resp.data[0]["id"]
            st.session_state.presupuesto_id = nuevo_id
            st.success(f"Presupuesto creado: {cliente} - {proyecto}")

    st.divider()
    st.subheader("O importar un presupuesto ya armado desde Excel")
    st.caption(
        "Trae cliente, proyecto, ubicacion y plazo directamente del Excel (hoja INICIO / "
        "CRONOGRAMA) y los items de la hoja PRESUPUESTO, todo en un solo paso -- para no "
        "volver a digitar datos que ya estan ahi y evitar errores de transcripcion."
    )
    archivo_gral = st.file_uploader(
        "Archivo Excel (.xlsx o .xlsm) del presupuesto ya armado",
        type=["xlsx", "xlsm"],
        key="importador_general_excel",
    )
    nombre_hoja_gral = st.text_input(
        "Nombre de la hoja de items", value="PRESUPUESTO", key="hoja_importar_general"
    )

    if archivo_gral is not None and st.button("Leer datos del Excel", key="btn_leer_general"):
        try:
            libro_gral = cargar_libro_excel(archivo_gral)
            st.session_state.import_datos_generales = leer_datos_generales_excel(libro_gral)
            st.session_state.import_items_general = parsear_excel_presupuesto(
                libro_gral, nombre_hoja_gral
            )
            if not st.session_state.import_items_general:
                st.warning(
                    "No encontre items para importar. Revisa el nombre de la hoja y que la "
                    "columna A tenga el encabezado 'COD APU'."
                )
        except Exception as e:
            st.error(f"No pude leer el archivo: {e}")
            st.session_state.import_datos_generales = None
            st.session_state.import_items_general = []

    datos_gral = st.session_state.get("import_datos_generales")
    items_gral = st.session_state.get("import_items_general", [])

    if datos_gral is not None:
        st.caption(
            "Estos datos vienen del Excel -- revisalos y corrige lo que haga falta antes "
            "de crear el presupuesto."
        )
        colg1, colg2 = st.columns(2)
        with colg1:
            cliente_g = st.text_input(
                "Cliente", value=datos_gral["cliente"], key="cliente_import"
            )
            proyecto_g = st.text_input(
                "Proyecto", value=datos_gral["proyecto"], key="proyecto_import"
            )
        with colg2:
            ubicacion_g = st.text_input(
                "Ubicacion", value=datos_gral["ubicacion"], key="ubicacion_import"
            )
            plazo_g = st.number_input(
                "Plazo (meses)",
                min_value=0.0,
                step=0.5,
                value=float(datos_gral["plazo_meses"]),
                key="plazo_import",
            )

        if items_gral:
            import pandas as pd

            codigos_encontrados_g = verificar_en_catalogo(
                sb, [f["codigo"] for f in items_gral]
            )
            df_prev_g = pd.DataFrame(
                [
                    {
                        "Capitulo": f["capitulo"],
                        "Codigo": f["codigo"],
                        "Descripcion": f["descripcion"],
                        "Unidad": f["unidad"],
                        "Cantidad": f["cantidad"],
                        "Precio unitario": f["precio_unitario"],
                        "Subtotal": f["cantidad"] * f["precio_unitario"],
                        "En catalogo": "Si" if f["codigo"] in codigos_encontrados_g else "No",
                    }
                    for f in items_gral
                ]
            )
            st.dataframe(df_prev_g, width="stretch", hide_index=True)
            total_import_g = df_prev_g["Subtotal"].sum()
            st.write(f"**{len(items_gral)} item(s) detectados · Total: {money(total_import_g)}**")

            sin_match_g = int((df_prev_g["En catalogo"] == "No").sum())
            if sin_match_g:
                st.warning(
                    f"{sin_match_g} item(s) con codigo que no esta en el catalogo actual. "
                    "Se importan igual (con su precio del Excel), pero sin vinculo al catalogo."
                )

            if st.button(
                "Crear presupuesto e importar todo", type="primary", key="btn_crear_importar"
            ):
                if not cliente_g or not proyecto_g:
                    st.error("Cliente y Proyecto son obligatorios (revisa los campos de arriba).")
                else:
                    resp_g = (
                        sb.table("presupuestos")
                        .insert(
                            {
                                "cliente": cliente_g,
                                "proyecto": proyecto_g,
                                "ubicacion": ubicacion_g,
                                "plazo_meses": plazo_g,
                                "plazo_dias": round(plazo_g * 30, 1),
                                "estado": "borrador",
                            }
                        )
                        .execute()
                    )
                    nuevo_id_g = resp_g.data[0]["id"]
                    insertados_g = importar_filas_a_presupuesto(
                        sb, nuevo_id_g, items_gral, codigos_encontrados_g
                    )
                    st.session_state.presupuesto_id = nuevo_id_g
                    st.session_state.import_datos_generales = None
                    st.session_state.import_items_general = []
                    st.success(
                        f"Presupuesto creado ({cliente_g} - {proyecto_g}) con "
                        f"{insertados_g} item(s) importado(s)."
                    )
                    st.rerun()

    st.divider()
    st.subheader("O continuar / editar un presupuesto existente")
    st.caption(
        "Incluye presupuestos ya generados: si el cliente pide agregar o quitar "
        "actividades, lo reabres aqui en vez de rehacerlo desde cero."
    )
    existentes = (
        sb.table("presupuestos")
        .select("id, cliente, proyecto, estado, creado_en")
        .neq("estado", "archivado")
        .order("creado_en", desc=True)
        .limit(50)
        .execute()
        .data
    )
    if existentes:
        opciones = {
            f"[{b['estado']}] {b['cliente']} - {b['proyecto']} ({b['creado_en'][:10]})": b["id"]
            for b in existentes
        }
        elegido = st.selectbox("Presupuestos existentes", list(opciones.keys()))
        if st.button("Abrir este presupuesto"):
            st.session_state.presupuesto_id = opciones[elegido]
            st.success("Presupuesto cargado.")
    else:
        st.info("No hay presupuestos todavia.")

    if st.session_state.presupuesto_id:
        st.info(f"Presupuesto activo (id): {st.session_state.presupuesto_id}")

# ---------------------------------------------------------------------
# TAB 2: buscar y agregar items del catalogo
# ---------------------------------------------------------------------
with tab_catalogo:
    if not st.session_state.presupuesto_id:
        st.warning("Primero crea o elige un presupuesto en la pestaña 1.")
    else:
        with st.expander("Importar un presupuesto ya armado desde Excel"):
            st.caption(
                "Sube la copia en Excel de un presupuesto existente (hoja 'PRESUPUESTO' del "
                "MODULO_PRESUPUESTOS_COMSAS_V4, con columnas COD APU | ITEM | ACTIVIDAD | UN | "
                "CANT | VR UNITARIO | VR TOTAL). Se usan los precios que ya estaban en el Excel "
                "-- no se recalculan con el catalogo actual -- para que el total importado "
                "coincida exactamente con el original."
            )
            archivo = st.file_uploader(
                "Archivo Excel (.xlsx o .xlsm)", type=["xlsx", "xlsm"], key="importador_excel"
            )
            nombre_hoja = st.text_input(
                "Nombre de la hoja", value="PRESUPUESTO", key="hoja_importar"
            )

            if archivo is not None and st.button("Leer archivo"):
                try:
                    libro = cargar_libro_excel(archivo)
                    st.session_state.importacion_filas = parsear_excel_presupuesto(
                        libro, nombre_hoja
                    )
                    if not st.session_state.importacion_filas:
                        st.warning(
                            "No encontre items para importar. Revisa el nombre de la hoja "
                            "y que la columna A tenga el encabezado 'COD APU'."
                        )
                except Exception as e:
                    st.error(f"No pude leer el archivo: {e}")
                    st.session_state.importacion_filas = []

            filas_import = st.session_state.get("importacion_filas", [])
            if filas_import:
                import pandas as pd

                codigos_encontrados = verificar_en_catalogo(
                    sb, [f["codigo"] for f in filas_import]
                )
                df_prev = pd.DataFrame(
                    [
                        {
                            "Capitulo": f["capitulo"],
                            "Codigo": f["codigo"],
                            "Descripcion": f["descripcion"],
                            "Unidad": f["unidad"],
                            "Cantidad": f["cantidad"],
                            "Precio unitario": f["precio_unitario"],
                            "Subtotal": f["cantidad"] * f["precio_unitario"],
                            "En catalogo": "Si" if f["codigo"] in codigos_encontrados else "No",
                        }
                        for f in filas_import
                    ]
                )
                st.dataframe(df_prev, width="stretch", hide_index=True)
                total_import = df_prev["Subtotal"].sum()
                st.write(f"**{len(filas_import)} item(s) detectados · Total: {money(total_import)}**")

                sin_match = int((df_prev["En catalogo"] == "No").sum())
                if sin_match:
                    st.warning(
                        f"{sin_match} item(s) con codigo que no esta en el catalogo actual. "
                        "Se importan igual (con su precio del Excel), pero sin vinculo al "
                        "catalogo -- no se actualizaran si luego cambia un precio maestro."
                    )

                if st.button("Importar estos items al presupuesto activo", type="primary"):
                    insertados = importar_filas_a_presupuesto(
                        sb, st.session_state.presupuesto_id, filas_import, codigos_encontrados
                    )
                    st.session_state.importacion_filas = []
                    st.success(
                        f"{insertados} item(s) importado(s). Ve a la pestaña 3 para ver el total."
                    )
                    st.rerun()

        st.divider()
        st.subheader("Buscar en el catalogo")
        busqueda = st.text_input("Buscar por codigo o descripcion", key="busq")
        col_a, col_b = st.columns([1, 3])
        with col_a:
            solo_activos = st.checkbox("Solo APUs completos (no candidatos)", value=True)

        query = sb.table("catalogo_apu").select(
            "codigo, categoria, descripcion, unidad, total, es_candidato"
        )
        if busqueda:
            query = query.or_(
                f"codigo.ilike.%{busqueda}%,descripcion.ilike.%{busqueda}%"
            )
        if solo_activos:
            query = query.eq("es_candidato", False)
        resultados = query.order("codigo").limit(100).execute().data

        st.caption(f"{len(resultados)} resultados (máximo 100 mostrados)")

        cantidades = {}
        for item in resultados:
            c1, c2, c3, c4 = st.columns([1, 4, 1, 1])
            with c1:
                st.text(item["codigo"])
            with c2:
                st.text(f"{item['descripcion'][:90]}  ·  {item['unidad']}")
            with c3:
                st.text(money(item["total"]))
            with c4:
                cantidades[item["codigo"]] = st.number_input(
                    "Cant.", min_value=0.0, step=1.0, key=f"cant_{item['codigo']}",
                    label_visibility="collapsed",
                )

        if st.button("Agregar items con cantidad > 0", type="primary"):
            agregados = 0
            for item in resultados:
                cant = cantidades.get(item["codigo"], 0)
                if cant and cant > 0:
                    # capitulo = categoria (prefijo del codigo), se crea si no existe
                    cap = (
                        sb.table("presupuesto_capitulos")
                        .select("id")
                        .eq("presupuesto_id", st.session_state.presupuesto_id)
                        .eq("codigo", item["categoria"])
                        .execute()
                        .data
                    )
                    if cap:
                        capitulo_id = cap[0]["id"]
                    else:
                        nuevo_cap = (
                            sb.table("presupuesto_capitulos")
                            .insert(
                                {
                                    "presupuesto_id": st.session_state.presupuesto_id,
                                    "codigo": item["categoria"],
                                    "nombre": item["categoria"],
                                }
                            )
                            .execute()
                        )
                        capitulo_id = nuevo_cap.data[0]["id"]

                    sb.table("presupuesto_items").insert(
                        {
                            "presupuesto_id": st.session_state.presupuesto_id,
                            "capitulo_id": capitulo_id,
                            "apu_codigo": item["codigo"],
                            "descripcion_snapshot": item["descripcion"],
                            "unidad_snapshot": item["unidad"],
                            "cantidad": cant,
                            "precio_unitario_snapshot": item["total"],
                        }
                    ).execute()
                    agregados += 1
            if agregados:
                st.success(f"{agregados} item(s) agregado(s) al presupuesto.")
            else:
                st.warning("No marcaste ninguna cantidad mayor a 0.")

# ---------------------------------------------------------------------
# TAB 3: resumen, editar cantidades, eliminar items, y total
# ---------------------------------------------------------------------
with tab_resumen:
    if not st.session_state.presupuesto_id:
        st.warning("Primero crea o elige un presupuesto en la pestaña 1.")
    else:
        items = (
            sb.table("presupuesto_items")
            .select("*, presupuesto_capitulos(nombre)")
            .eq("presupuesto_id", st.session_state.presupuesto_id)
            .order("id")
            .execute()
            .data
        )
        if not items:
            st.info("Este presupuesto todavia no tiene items. Agrega algunos en la pestaña 2.")
        else:
            import pandas as pd

            st.subheader("Editar items (agrupado por capitulo, igual que en el Excel)")
            st.caption(
                "Cambia la cantidad directamente en la tabla, o marca 'Eliminar' "
                "para quitar una actividad que el cliente ya no pidio. Al final, "
                "dale clic a 'Guardar cambios'."
            )

            por_capitulo = {}
            for it in items:
                nombre_cap = (it.get("presupuesto_capitulos") or {}).get("nombre", "Sin capitulo")
                por_capitulo.setdefault(nombre_cap, []).append(it)

            ediciones = []  # [(items_originales_del_capitulo, dataframe_editado), ...]
            for nombre_cap, items_cap in por_capitulo.items():
                subtotal_cap = sum(it["cantidad"] * it["precio_unitario_snapshot"] for it in items_cap)
                st.markdown(f"**{nombre_cap}**")
                df_cap = pd.DataFrame(
                    [
                        {
                            "id": it["id"],
                            "Codigo": it["apu_codigo"],
                            "Descripcion": it["descripcion_snapshot"],
                            "Unidad": it["unidad_snapshot"],
                            "Cantidad": float(it["cantidad"]),
                            "Precio unitario": float(it["precio_unitario_snapshot"]),
                            "Eliminar": False,
                        }
                        for it in items_cap
                    ]
                )
                clave_segura = re.sub(r"[^a-zA-Z0-9]+", "_", nombre_cap)
                editado_cap = st.data_editor(
                    df_cap,
                    column_config={
                        "id": None,  # oculto, solo para identificar la fila al guardar
                        "Precio unitario": st.column_config.NumberColumn(format="$ %d", disabled=True),
                        "Cantidad": st.column_config.NumberColumn(min_value=0.0, step=1.0),
                        "Codigo": st.column_config.TextColumn(disabled=True),
                        "Descripcion": st.column_config.TextColumn(disabled=True),
                        "Unidad": st.column_config.TextColumn(disabled=True),
                    },
                    hide_index=True,
                    width="stretch",
                    key=f"editor_items_{clave_segura}",
                )
                st.caption(f"Sub-Total {nombre_cap}: {money(subtotal_cap)}")
                ediciones.append((items_cap, editado_cap))

            if st.button("Guardar cambios", type="primary"):
                eliminados, actualizados = 0, 0
                for items_cap, editado_cap in ediciones:
                    originales = {it["id"]: it for it in items_cap}
                    for _, fila in editado_cap.iterrows():
                        item_id = int(fila["id"])
                        if fila["Eliminar"]:
                            sb.table("presupuesto_items").delete().eq("id", item_id).execute()
                            eliminados += 1
                        elif float(fila["Cantidad"]) != float(originales[item_id]["cantidad"]):
                            sb.table("presupuesto_items").update(
                                {"cantidad": float(fila["Cantidad"])}
                            ).eq("id", item_id).execute()
                            actualizados += 1
                st.success(f"Listo: {eliminados} eliminado(s), {actualizados} cantidad(es) actualizada(s).")
                st.rerun()

            st.divider()
            st.subheader("Totales")
            total_general = 0
            for nombre_cap, filas_cap in por_capitulo.items():
                subtotal = sum(it["cantidad"] * it["precio_unitario_snapshot"] for it in filas_cap)
                st.write(f"**Subtotal {nombre_cap}: {money(subtotal)}**")
                total_general += subtotal

            st.markdown(f"## Costo directo total: {money(total_general)}")
            st.caption(
                "Este total es el costo directo, equivalente a la hoja PRESUPUESTO de tu Excel. "
                "El AIU y la propuesta se calculan en la pestaña 4."
            )

            sb.table("presupuestos").update(
                {"costo_directo": total_general}
            ).eq("id", st.session_state.presupuesto_id).execute()

            st.divider()
            st.subheader("Plazo de ejecucion")
            plazo_sugerido, codigos_sin_rendimiento = calcular_plazo_dias_sugerido(sb, items)
            st.caption(
                f"Plazo sugerido segun rendimientos: **{plazo_sugerido:g} dias** "
                "(suma de cantidad / rendimiento-dia de cada actividad, una tras otra). "
                "Es un punto de partida -- ajustalo abajo si tu criterio dice que la app "
                "se quedo corta o se paso."
            )
            if codigos_sin_rendimiento:
                st.caption(
                    f"{len(codigos_sin_rendimiento)} item(s) sin rendimiento/dia definido en "
                    "su APU no se contaron para este calculo: "
                    + ", ".join(codigos_sin_rendimiento[:15])
                    + ("..." if len(codigos_sin_rendimiento) > 15 else "")
                )
            presu_plazo_actual = (
                sb.table("presupuestos")
                .select("plazo_dias")
                .eq("id", st.session_state.presupuesto_id)
                .single()
                .execute()
                .data
            )
            plazo_guardado = (presu_plazo_actual or {}).get("plazo_dias")
            valor_inicial_plazo = float(plazo_guardado) if plazo_guardado else plazo_sugerido
            plazo_dias_final = st.number_input(
                "Plazo (dias) -- editable",
                min_value=0.0,
                step=1.0,
                value=valor_inicial_plazo,
                key="plazo_dias_editable",
            )
            if st.button("Guardar plazo"):
                sb.table("presupuestos").update(
                    {"plazo_dias": plazo_dias_final, "plazo_meses": round(plazo_dias_final / 30, 2)}
                ).eq("id", st.session_state.presupuesto_id).execute()
                st.success(f"Plazo guardado: {plazo_dias_final:g} dias.")

            st.divider()
            st.subheader("Refrescar precios de este presupuesto")
            st.caption(
                "Trae el precio actual del catalogo para los items que ya tienen apu_codigo "
                "(los importados o agregados a mano sin codigo no se tocan). Esto SOLO afecta "
                "este presupuesto -- ningun otro presupuesto ni el catalogo maestro se modifican."
            )
            estado_presu = (
                sb.table("presupuestos").select("estado").eq("id", st.session_state.presupuesto_id).single().execute().data
            )
            if estado_presu and estado_presu.get("estado") in ("enviado", "aprobado"):
                st.warning(
                    "Este presupuesto ya fue marcado como enviado/aprobado -- si refrescas "
                    "precios aqui, el valor interno cambia pero el documento que ya le "
                    "entregaste al cliente NO se actualiza solo. Genera la propuesta de "
                    "nuevo en la pestaña 4 si necesitas que coincidan."
                )

            if st.button("Ver cambios de precio disponibles"):
                st.session_state.refresco_cambios = preparar_refresco_precios(sb, st.session_state.presupuesto_id)

            cambios_refresco = st.session_state.get("refresco_cambios")
            if cambios_refresco is not None:
                if not cambios_refresco:
                    st.info("No hay diferencias -- todos los items enlazados a un APU ya tienen el precio actual del catalogo.")
                else:
                    import pandas as pd

                    df_refresco = pd.DataFrame(
                        [
                            {
                                "Codigo": c["apu_codigo"],
                                "Descripcion": c["descripcion"],
                                "Cantidad": c["cantidad"],
                                "Precio actual (presupuesto)": c["precio_viejo"],
                                "Precio nuevo (catalogo)": c["precio_nuevo"],
                                "Var. subtotal": (c["precio_nuevo"] - c["precio_viejo"]) * c["cantidad"],
                            }
                            for c in cambios_refresco
                        ]
                    )
                    st.dataframe(df_refresco, width="stretch", hide_index=True)
                    variacion_total = df_refresco["Var. subtotal"].sum()
                    st.write(
                        f"**{len(cambios_refresco)} item(s) cambiarian · Variacion en el costo directo: {money(variacion_total)}**"
                    )

                    if st.button("Aplicar refresco de precios a este presupuesto", type="primary"):
                        nuevo_total = aplicar_refresco_precios(sb, st.session_state.presupuesto_id, cambios_refresco)
                        st.session_state.refresco_cambios = None
                        st.success(
                            f"{len(cambios_refresco)} item(s) actualizado(s). Nuevo costo directo: {money(nuevo_total)}."
                        )
                        st.rerun()

# ---------------------------------------------------------------------
# TAB 4: AIU, propuesta en Word y Excel de respaldo
# ---------------------------------------------------------------------
with tab_aiu:
    if not st.session_state.presupuesto_id:
        st.warning("Primero crea o elige un presupuesto en la pestaña 1.")
    else:
        presupuesto = (
            sb.table("presupuestos")
            .select("*")
            .eq("id", st.session_state.presupuesto_id)
            .single()
            .execute()
            .data
        )
        items, por_capitulo = obtener_items_y_capitulos(sb, st.session_state.presupuesto_id)

        if not items:
            st.info("Este presupuesto todavia no tiene items. Agrega algunos en la pestaña 2.")
        else:
            costo_real = presupuesto.get("costo_directo") or sum(
                float(it["cantidad"]) * float(it["precio_unitario_snapshot"]) for it in items
            )

            estado_actual = presupuesto.get("estado") or "borrador"
            colE1, colE2 = st.columns([1, 2])
            with colE1:
                st.metric("Estado del presupuesto", estado_actual.upper())
            with colE2:
                st.write("")
                if estado_actual == "borrador":
                    st.caption("Pasa a GENERADO automaticamente cuando generes la propuesta en Word (abajo).")
                elif estado_actual == "generado":
                    if st.button("Marcar como enviado al cliente"):
                        sb.table("presupuestos").update({"estado": "enviado"}).eq(
                            "id", st.session_state.presupuesto_id
                        ).execute()
                        st.success("Marcado como enviado.")
                        st.rerun()
                elif estado_actual == "enviado":
                    if st.button("Marcar como aprobado"):
                        sb.table("presupuestos").update({"estado": "aprobado"}).eq(
                            "id", st.session_state.presupuesto_id
                        ).execute()
                        st.success("Marcado como aprobado.")
                        st.rerun()
                elif estado_actual == "aprobado":
                    st.caption("Presupuesto aprobado.")
            st.divider()

            st.subheader("1. Margen real de la empresa")
            st.caption(
                "Este margen es aparte del AIU -- NO se le muestra al cliente como tal. Se "
                "suma al costo de cada APU (de COSTOS) ANTES de calcular el AIU, y ese costo "
                "aumentado es el que la app usa como 'Costo directo' para todo lo demas "
                "(AIU, propuesta, Excel de respaldo). Por defecto es el mismo para todos los "
                "proyectos (parametro global), pero se puede ajustar solo para este proyecto."
            )
            margen_real_pct_default = obtener_margen_real_pct(sb, presupuesto)
            margen_real_pct_ui = st.number_input(
                "% Margen real de la empresa", min_value=0.0, max_value=100.0, step=1.0,
                value=margen_real_pct_default * 100, format="%.2f",
            )
            margen_real_pct = margen_real_pct_ui / 100

            items_presentacion = aplicar_margen_items(items, margen_real_pct)
            por_capitulo_presentacion = agrupar_por_capitulo(items_presentacion)
            costo_directo = sum(
                float(it["cantidad"]) * float(it["precio_unitario_snapshot"]) for it in items_presentacion
            )

            colm1, colm2 = st.columns(2)
            with colm1:
                st.metric("Costo real (lo que le cuesta a la empresa)", money(costo_real))
            with colm2:
                st.metric("Costo directo (con margen -- base del AIU)", money(costo_directo))

            st.divider()

            st.subheader("2. Administracion, Imprevistos, Utilidad (AIU)")
            st.caption(
                "Se calcula sobre el costo directo YA con el margen real incluido, igual que "
                "en PLANTILLA_AIU.xlsx. Los porcentajes quedan guardados en el presupuesto."
            )
            st.write(f"**Costo directo:** {money(costo_directo)}")

            st.caption("Los porcentajes se digitan como numero entero (12 = 12%, no 0.12).")
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                administracion_pct_ui = st.number_input(
                    "% Administracion", min_value=0.0, max_value=100.0, step=1.0,
                    value=float(presupuesto.get("administracion_pct") or 0.12) * 100,
                    format="%.2f",
                )
            with col2:
                imprevistos_pct_ui = st.number_input(
                    "% Imprevistos", min_value=0.0, max_value=100.0, step=1.0,
                    value=float(presupuesto.get("imprevistos_pct") or 0.01) * 100,
                    format="%.2f",
                )
            with col3:
                utilidad_pct_ui = st.number_input(
                    "% Utilidad", min_value=0.0, max_value=100.0, step=1.0,
                    value=float(presupuesto.get("utilidad_pct") or 0.05) * 100,
                    format="%.2f",
                )
            with col4:
                iva_utilidad_pct_ui = st.number_input(
                    "% IVA sobre utilidad", min_value=0.0, max_value=100.0, step=1.0,
                    value=float(presupuesto.get("iva_utilidad_pct") or 0.19) * 100,
                    format="%.2f",
                )

            administracion_pct = administracion_pct_ui / 100
            imprevistos_pct = imprevistos_pct_ui / 100
            utilidad_pct = utilidad_pct_ui / 100
            iva_utilidad_pct = iva_utilidad_pct_ui / 100

            aiu = calcular_aiu(costo_directo, administracion_pct, imprevistos_pct, utilidad_pct, iva_utilidad_pct)

            r1, r2, r3 = st.columns(3)
            with r1:
                st.metric("Administracion", money(aiu["administracion"]))
                st.metric("Imprevistos", money(aiu["imprevistos"]))
            with r2:
                st.metric("Utilidad", money(aiu["utilidad"]))
                st.metric("IVA utilidad", money(aiu["iva_utilidad"]))
            with r3:
                st.metric("Total AIU (%)", f"{aiu['aiu_total_pct'] * 100:.2f}%")
                st.metric("Valor del contrato", money(aiu["valor_total"]))

            st.divider()
            st.subheader("3. Datos para la propuesta")
            colp1, colp2 = st.columns(2)
            with colp1:
                atencion = st.text_input(
                    "Dirigido a (nombre y cargo)", value=presupuesto.get("atencion") or ""
                )
                vigencia_dias = st.number_input(
                    "Vigencia de la propuesta (dias calendario)", min_value=1, step=1,
                    value=int(presupuesto.get("vigencia_dias") or 30),
                )
            with colp2:
                st.caption("Condiciones de pago (% del valor total, en numero entero)")
                pcol1, pcol2, pcol3, pcol4 = st.columns(4)
                with pcol1:
                    anticipo_pct_ui = st.number_input(
                        "Anticipo", min_value=0.0, max_value=100.0, step=5.0,
                        value=float(presupuesto.get("anticipo_pct") or 0.45) * 100, format="%.0f",
                    )
                with pcol2:
                    pago2_pct_ui = st.number_input(
                        "2do pago", min_value=0.0, max_value=100.0, step=5.0,
                        value=float(presupuesto.get("pago2_pct") or 0.45) * 100, format="%.0f",
                    )
                with pcol3:
                    avance2_pct_ui = st.number_input(
                        "Avance para 2do pago", min_value=0.0, max_value=100.0, step=5.0,
                        value=float(presupuesto.get("avance2_pct") or 0.50) * 100, format="%.0f",
                    )
                with pcol4:
                    pagofin_pct_ui = st.number_input(
                        "Saldo final", min_value=0.0, max_value=100.0, step=5.0,
                        value=float(presupuesto.get("pagofin_pct") or 0.10) * 100, format="%.0f",
                    )

                anticipo_pct = anticipo_pct_ui / 100
                pago2_pct = pago2_pct_ui / 100
                avance2_pct = avance2_pct_ui / 100
                pagofin_pct = pagofin_pct_ui / 100

            if st.button("Guardar AIU y datos de propuesta", type="primary"):
                sb.table("presupuestos").update(
                    {
                        "administracion_pct": administracion_pct,
                        "imprevistos_pct": imprevistos_pct,
                        "utilidad_pct": utilidad_pct,
                        "iva_utilidad_pct": iva_utilidad_pct,
                        "atencion": atencion,
                        "vigencia_dias": int(vigencia_dias),
                        "anticipo_pct": anticipo_pct,
                        "pago2_pct": pago2_pct,
                        "avance2_pct": avance2_pct,
                        "pagofin_pct": pagofin_pct,
                        "aiu_total": aiu["aiu_total_valor"],
                        "valor_total": aiu["valor_total"],
                        "margen_real_pct": margen_real_pct,
                        "costo_directo_presentacion": costo_directo,
                    }
                ).eq("id", st.session_state.presupuesto_id).execute()
                presupuesto.update(
                    {
                        "administracion_pct": administracion_pct,
                        "imprevistos_pct": imprevistos_pct,
                        "utilidad_pct": utilidad_pct,
                        "iva_utilidad_pct": iva_utilidad_pct,
                        "atencion": atencion,
                        "vigencia_dias": int(vigencia_dias),
                        "anticipo_pct": anticipo_pct,
                        "pago2_pct": pago2_pct,
                        "avance2_pct": avance2_pct,
                        "pagofin_pct": pagofin_pct,
                    }
                )
                st.success("Guardado.")

            st.divider()
            st.subheader("4. Generar documentos")

            import os

            carpeta_plantillas = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            plantilla_path = os.path.join(carpeta_plantillas, "PLANTILLA PROPUESTA COMSAS.docx")
            # Archivo "portador" de la macro Nlet, SOLO con esa macro (sin
            # ninguna hoja del modulo maestro) -- para que el respaldo no
            # cargue datos internos si algun dia se comparte por error.
            # Se crea una unica vez en Excel (ver instrucciones aparte).
            plantilla_macro_path = os.path.join(carpeta_plantillas, "PLANTILLA_MACRO_LETRAS.xlsm")

            if "carpeta_guardado" not in st.session_state:
                st.session_state.carpeta_guardado = os.path.expanduser("~")

            st.caption(
                f"Al generar se abrira el explorador de carpetas de Windows (ultima carpeta "
                f"usada: {st.session_state.carpeta_guardado})."
            )

            cold, colx = st.columns(2)
            with cold:
                st.caption("Propuesta en Word (usa PLANTILLA PROPUESTA COMSAS.docx)")
                if st.button("Generar propuesta en Word"):
                    if not os.path.exists(plantilla_path):
                        st.error(f"No encuentro la plantilla en: {plantilla_path}")
                    else:
                        buffer = generar_propuesta_docx(
                            plantilla_path, presupuesto, items_presentacion, por_capitulo_presentacion,
                            costo_directo, aiu,
                        )
                        nombre = sanitizar_nombre_archivo(
                            f"PROPUESTA - {presupuesto.get('proyecto', 'presupuesto')[:60]}.docx"
                        )
                        carpeta_elegida = elegir_carpeta_con_dialogo(st.session_state.carpeta_guardado)
                        if carpeta_elegida:
                            st.session_state.carpeta_guardado = carpeta_elegida
                            ruta = os.path.join(carpeta_elegida, nombre)
                            try:
                                with open(ruta, "wb") as fh:
                                    fh.write(buffer.getvalue())
                                st.success(f"Guardada en: {ruta}")
                            except Exception as e:
                                st.error(f"No pude guardar el archivo en esa carpeta ({e}). Descargala abajo.")
                        else:
                            st.info(
                                "No elegiste carpeta (o el dialogo no pudo abrirse -- esto solo "
                                "funciona corriendo la app en tu propia maquina). Descargala abajo."
                            )
                        st.download_button(
                            "Descargar propuesta (.docx)",
                            data=buffer,
                            file_name=nombre,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )

                        # Generar la propuesta es lo que marca que el presupuesto
                        # dejo de ser un borrador interno. Solo avanza el estado,
                        # nunca lo regresa (si ya estaba enviado/aprobado, se deja igual).
                        if (presupuesto.get("estado") or "borrador") == "borrador":
                            sb.table("presupuestos").update({"estado": "generado"}).eq(
                                "id", st.session_state.presupuesto_id
                            ).execute()
                            presupuesto["estado"] = "generado"
                            st.info("Estado del presupuesto actualizado a 'generado'.")
            with colx:
                st.caption("Excel de respaldo del presupuesto (letras en vivo con macro Nlet)")
                if st.button("Generar Excel de respaldo"):
                    usa_macro = os.path.exists(plantilla_macro_path)
                    if usa_macro:
                        buffer_xlsx = generar_excel_respaldo_macro(
                            sb, presupuesto, items_presentacion, por_capitulo_presentacion, costo_directo,
                            aiu, plantilla_macro_path, margen_real_pct,
                        )
                        extension = "xlsm"
                        mime_xlsx = "application/vnd.ms-excel.sheet.macroEnabled.12"
                    else:
                        st.warning(
                            f"No encontre {plantilla_macro_path} -- genero el respaldo sin macro "
                            "(el valor en letras queda fijo, no se recalcula solo)."
                        )
                        buffer_xlsx = generar_excel_respaldo(
                            sb, presupuesto, items_presentacion, por_capitulo_presentacion, costo_directo,
                            aiu, margen_real_pct,
                        )
                        extension = "xlsx"
                        mime_xlsx = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

                    nombre_xlsx = sanitizar_nombre_archivo(
                        f"RESPALDO - {presupuesto.get('proyecto', 'presupuesto')[:60]}.{extension}"
                    )
                    carpeta_elegida_x = elegir_carpeta_con_dialogo(st.session_state.carpeta_guardado)
                    if carpeta_elegida_x:
                        st.session_state.carpeta_guardado = carpeta_elegida_x
                        ruta_xlsx = os.path.join(carpeta_elegida_x, nombre_xlsx)
                        try:
                            with open(ruta_xlsx, "wb") as fh:
                                fh.write(buffer_xlsx.getvalue())
                            st.success(f"Guardado en: {ruta_xlsx}")
                            if usa_macro:
                                st.caption(
                                    "Al abrirlo, Excel va a pedir 'Habilitar contenido' -- es normal, "
                                    "es la macro Nlet que calcula el valor en letras."
                                )
                        except Exception as e:
                            st.error(f"No pude guardar el archivo en esa carpeta ({e}). Descargalo abajo.")
                    else:
                        st.info(
                            "No elegiste carpeta (o el dialogo no pudo abrirse -- esto solo "
                            "funciona corriendo la app en tu propia maquina). Descargalo abajo."
                        )
                    st.download_button(
                        f"Descargar respaldo (.{extension})",
                        data=buffer_xlsx,
                        file_name=nombre_xlsx,
                        mime=mime_xlsx,
                    )

            st.divider()
            st.caption(
                "Excel de manejo interno (Cronograma, Lista de materiales, Flujo de caja, "
                "Dashboard y APUs) -- solo para el equipo, NO se le manda al cliente. El "
                "cronograma y el flujo de caja ya son diarios (no por mes/semana)."
            )
            usar_curva_s_ui = st.checkbox(
                "Repartir el cronograma con curva S (en vez de reparto lineal por dia)",
                value=False,
                key="usar_curva_s",
            )
            if st.button("Generar Excel de manejo interno"):
                buffer_interno = generar_excel_manejo_interno(
                    sb, presupuesto, items_presentacion, por_capitulo_presentacion, costo_directo, aiu,
                    usar_curva_s_ui,
                )
                nombre_interno = sanitizar_nombre_archivo(
                    f"MANEJO INTERNO - {presupuesto.get('proyecto', 'presupuesto')[:60]}.xlsx"
                )
                carpeta_elegida_i = elegir_carpeta_con_dialogo(st.session_state.carpeta_guardado)
                if carpeta_elegida_i:
                    st.session_state.carpeta_guardado = carpeta_elegida_i
                    ruta_interno = os.path.join(carpeta_elegida_i, nombre_interno)
                    try:
                        with open(ruta_interno, "wb") as fh:
                            fh.write(buffer_interno.getvalue())
                        st.success(f"Guardado en: {ruta_interno}")
                    except Exception as e:
                        st.error(f"No pude guardar el archivo en esa carpeta ({e}). Descargalo abajo.")
                else:
                    st.info(
                        "No elegiste carpeta (o el dialogo no pudo abrirse -- esto solo "
                        "funciona corriendo la app en tu propia maquina). Descargalo abajo."
                    )
                st.download_button(
                    "Descargar Excel de manejo interno (.xlsx)",
                    data=buffer_interno,
                    file_name=nombre_interno,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )

# ---------------------------------------------------------------------
# TAB 5: mantenimiento de precios del catalogo (Fase 4)
# ---------------------------------------------------------------------
with tab_precios:
    st.caption(
        "Mantenimiento del catalogo completo (no de un presupuesto en particular). "
        "Equivalente a MANTENIMIENTO_PRECIOS.txt y RECOSTEAR_MANO_OBRA.txt del Excel maestro."
    )

    sub_semaforo, sub_icoced, sub_facturas, sub_cuadrillas, sub_mo, sub_bitacora = st.tabs(
        [
            "Semaforo de vigencia",
            "Escalar por ICOCED",
            "Precios desde facturas",
            "Cuadrillas",
            "Recosteo de mano de obra",
            "Bitacora",
        ]
    )

    # --- Semaforo de vigencia (a nivel de insumo, hoja MATERIALES) ---
    with sub_semaforo:
        st.subheader("Semaforo de vigencia de precios de insumos")
        st.caption(
            "VIGENTE (<=90 dias) / POR VENCER (91-180) / VENCIDO (>180) / "
            "ESTIMADO (ultimo cambio fue por ICOCED) / SIN FECHA (nunca se ha cotizado). "
            "Vive a nivel de insumo (materia prima), no de APU -- un APU puede tener "
            "varios insumos con vigencia distinta."
        )
        import pandas as pd

        filas_sem = obtener_insumos_para_semaforo(sb)
        if not filas_sem:
            st.info("No hay insumos activos en el catalogo.")
        else:
            df_sem = pd.DataFrame(filas_sem)
            conteo = df_sem["vigencia"].value_counts()
            cols_metric = st.columns(5)
            for col, estado in zip(
                cols_metric, ["VIGENTE", "POR VENCER", "VENCIDO", "ESTIMADO", "SIN FECHA"]
            ):
                col.metric(estado, int(conteo.get(estado, 0)))

            filtro_estado = st.multiselect(
                "Filtrar por vigencia",
                ["VIGENTE", "POR VENCER", "VENCIDO", "ESTIMADO", "SIN FECHA"],
                default=["VENCIDO", "POR VENCER", "SIN FECHA"],
            )
            busq_insumo = st.text_input("Buscar por codigo o descripcion", key="busq_insumo_sem")
            df_mostrar = df_sem[df_sem["vigencia"].isin(filtro_estado)] if filtro_estado else df_sem
            if busq_insumo:
                mask = df_mostrar["codigo"].str.contains(busq_insumo, case=False, na=False) | df_mostrar[
                    "descripcion"
                ].str.contains(busq_insumo, case=False, na=False)
                df_mostrar = df_mostrar[mask]
            st.caption(f"{len(df_mostrar)} insumo(s)")
            st.dataframe(
                df_mostrar[
                    ["codigo", "descripcion", "unidad", "precio", "proveedor", "fecha_cotizacion", "origen_precio", "vigencia"]
                ],
                width="stretch",
                hide_index=True,
            )

    # --- Escalar por ICOCED (a nivel de insumo) ---
    with sub_icoced:
        st.subheader("Escalar precios de insumos por indice ICOCED (DANE)")
        st.caption(
            "Sube el % del subindice de MATERIALES del boletin ICOCED del mes (lo publica el "
            "DANE, normalmente a mediados del mes siguiente). Solo se aplica a los insumos "
            "que NO se refrescaron ya este mes (no hay doble conteo). Jerarquia: un precio "
            "real (COTIZADO/FACTURA) manda; ICOCED solo estima mientras llega un precio real. "
            "Al aplicar, se recalcula automaticamente el 'materiales' de cada APU que use "
            "alguno de estos insumos."
        )
        st.caption(
            "El campo ya viene con el ultimo % que usaste (guardado como parametro) -- si no "
            "ha salido un boletin nuevo del DANE, puedes dejarlo igual; si ya sabes el nuevo "
            "numero, solo cambialo aqui."
        )
        pct_guardado = obtener_parametro(sb, "icoced_pct_mensual", "0")
        pct_icoced = st.number_input(
            "Variacion del mes (%)", step=0.1, format="%.2f", value=float(pct_guardado), key="pct_icoced"
        )

        with st.expander("Buscar el % en el boletin del DANE (beta)"):
            st.caption(
                "Intenta descargar el anexo mensual del ICOCED directo del DANE y encontrar "
                "la variacion de 'Materiales'. Es un formato que el DANE puede cambiar sin "
                "aviso, asi que esto es SOLO una sugerencia -- revisala contra el boletin "
                "antes de usarla. Si no encuentra nada, bajalo tu mismo aqui: "
                "https://www.dane.gov.co/index.php/estadisticas-por-tema/precios-y-costos/"
                "indice-de-costos-de-la-construccion-de-edificaciones-icoced"
            )
            if st.button("Buscar en el DANE"):
                with st.spinner("Descargando y leyendo el anexo del DANE..."):
                    resultado_dane = buscar_variacion_icoced_dane()
                st.session_state.resultado_dane_icoced = resultado_dane
                if resultado_dane is None:
                    st.warning(
                        "No pude descargar o leer el anexo automaticamente (puede que el "
                        "boletin de este mes aun no salga, o que el DANE haya cambiado el "
                        "formato). Bajalo del link de arriba y escribe el % a mano."
                    )

            resultado_dane = st.session_state.get("resultado_dane_icoced")
            if resultado_dane:
                import pandas as pd

                st.caption(f"Anexo leido: {resultado_dane['mes']}{resultado_dane['anio']} -- {resultado_dane['url']}")
                st.caption(
                    "Busca la columna 'Mensual' (no 'Año corrido' ni 'Doce meses') en la fila "
                    "de 'Materiales' -- si el encabezado detectado no dice claramente cual es "
                    "cual, abre el anexo (link arriba) y usa la celda como referencia."
                )
                df_dane = pd.DataFrame(resultado_dane["candidatos"])
                st.dataframe(df_dane, width="stretch", hide_index=True)

                opciones = {
                    f"{c['hoja']} · {c['celda']} · {c['encabezado']} · {c['etiqueta_fila']} = {c['valor']:g}%": c["valor"]
                    for c in resultado_dane["candidatos"]
                }
                etiqueta_elegida = st.selectbox(
                    "Elige la fila que corresponda a 'Materiales, variacion mensual, total nacional'",
                    list(opciones.keys()),
                    key="opcion_dane_elegida",
                )
                if st.button("Usar este valor en el campo de arriba"):
                    st.session_state.pct_icoced = float(opciones[etiqueta_elegida])
                    st.session_state.resultado_dane_icoced = None
                    st.rerun()

        if st.button("Ver candidatos a escalar"):
            insumos_todos = (
                sb.table("insumos")
                .select("codigo, descripcion, precio, fecha_cotizacion, origen_precio")
                .eq("activo", True)
                .execute()
                .data
            )
            mes_actual = (date.today().year, date.today().month)
            candidatos = [f for f in insumos_todos if es_escalable_icoced_insumo(f, mes_actual)]
            st.session_state.icoced_candidatos = candidatos

        candidatos = st.session_state.get("icoced_candidatos", [])
        if candidatos:
            import pandas as pd

            st.write(f"**{len(candidatos)} insumo(s) candidatos** (precio > 0 y no refrescados este mes).")
            if pct_icoced != 0:
                df_prev_icoced = pd.DataFrame(
                    [
                        {
                            "Codigo": f["codigo"],
                            "Descripcion": f["descripcion"],
                            "Precio actual": float(f["precio"] or 0),
                            "Precio nuevo": round(float(f["precio"] or 0) * (1 + pct_icoced / 100), 0),
                        }
                        for f in candidatos
                    ]
                )
                st.dataframe(df_prev_icoced, width="stretch", hide_index=True)

                if not es_admin():
                    st.caption("Solo un administrador puede aplicar este escalamiento.")
                if st.button(
                    "Aplicar escalamiento a todos los candidatos",
                    type="primary",
                    disabled=not es_admin(),
                ):
                    n_aplicados, n_apus = aplicar_escalamiento_icoced(sb, candidatos, pct_icoced)
                    guardar_parametro(
                        sb, "icoced_pct_mensual", pct_icoced,
                        "Ultimo % de variacion mensual ICOCED (materiales) usado para escalar precios de insumos",
                    )
                    st.session_state.icoced_candidatos = []
                    st.success(
                        f"{n_aplicados} precio(s) de insumo escalado(s) +{pct_icoced:g}% y "
                        f"registrados en la bitacora. {n_apus} APU(s) recalculado(s). "
                        "Quedan marcados como ESTIMADO hasta que llegue un precio real."
                    )
                    st.rerun()
            else:
                st.info("Escribe un porcentaje distinto de 0 para ver la vista previa.")
        else:
            st.caption("Dale clic a 'Ver candidatos a escalar' para empezar.")

    # --- Precios desde facturas (puente Control de Costos -> insumos) ---
    with sub_facturas:
        st.subheader("Actualizar precios de insumos desde el Control de Costos")
        st.caption(
            "Sube el Control de Costos del proyecto (hoja 'GASTOS', con COD INSUMO + CANT + "
            "BASE + FECHA + PROVEEDOR). El precio real (BASE/CANT) manda sobre COTIZADO e "
            "ICOCED -- este es el puente equivalente a ActualizarPreciosDesdeFacturas de la "
            "macro. Al aplicar, se recalcula automaticamente el 'materiales' de cada APU "
            "afectado."
        )
        archivo_gastos = st.file_uploader(
            "Control de Costos (.xlsx o .xlsm)", type=["xlsx", "xlsm"], key="archivo_gastos"
        )

        if archivo_gastos is not None and st.button("Leer GASTOS y ver propuestas"):
            try:
                gastos = parsear_gastos_control_costos(archivo_gastos)
                if not gastos:
                    st.warning(
                        "No encontre gastos validos (revisa que GASTOS tenga COD INSUMO en "
                        "la columna AI y CANT en AJ, desde la fila 4)."
                    )
                    st.session_state.propuestas_facturas = []
                else:
                    st.session_state.propuestas_facturas = preparar_propuestas_facturas(sb, gastos)
                    st.session_state.gastos_leidos = len(gastos)
            except Exception as e:
                st.error(f"No pude leer el archivo: {e}")
                st.session_state.propuestas_facturas = []

        propuestas = st.session_state.get("propuestas_facturas", [])
        if propuestas:
            import pandas as pd

            st.write(
                f"**{len(propuestas)} precio(s) propuesto(s)** de {st.session_state.get('gastos_leidos', '?')} "
                "gasto(s) leidos (los demas ya estaban aplicados, o un precio real mas nuevo ya manda)."
            )
            df_prop = pd.DataFrame(propuestas)
            grandes = df_prop["var_pct"].abs() > 40
            if grandes.any():
                st.warning(
                    f"{int(grandes.sum())} propuesta(s) con variacion mayor al 40% -- revisalas "
                    "antes de aplicar (columna 'var_pct')."
                )
            st.dataframe(
                df_prop[["codigo", "descripcion", "precio_actual", "precio_factura", "var_pct", "fecha", "proveedor"]],
                width="stretch",
                hide_index=True,
            )

            if not es_admin():
                st.caption("Solo un administrador puede aplicar estos precios.")
            if st.button(
                "Aplicar todos los precios propuestos",
                type="primary",
                disabled=not es_admin(),
            ):
                n_aplicados, n_apus = aplicar_precios_desde_facturas(sb, propuestas)
                st.session_state.propuestas_facturas = []
                st.success(
                    f"{n_aplicados} precio(s) de insumo actualizado(s) desde facturas y "
                    f"registrados en la bitacora. {n_apus} APU(s) recalculado(s)."
                )
                st.rerun()
        else:
            st.caption("Sube el archivo y dale clic a 'Leer GASTOS y ver propuestas' para empezar.")

    # --- Recosteo de mano de obra ---
    # --- Cuadrillas: crear y editar composicion (Fase 2) ---
    with sub_cuadrillas:
        st.subheader("Cuadrillas (unidades basicas de mano de obra)")
        st.caption(
            "Crea cuadrillas nuevas o edita la composicion (cargo y cantidad) de las "
            "existentes -- equivalente a CREAR_CUADRILLAS.txt del Excel maestro. El "
            "costo/dia sale solo de las tarifas de personal (pestaña 'Recosteo de "
            "mano de obra'). Para asignarle una cuadrilla a un APU (nuevo o "
            "existente), hazlo desde '6. Editar receta de APU'."
        )
        import pandas as pd

        cuadrillas_actuales = obtener_cuadrillas(sb)
        df_cuadrillas = pd.DataFrame(cuadrillas_actuales)
        st.markdown("**Cuadrillas existentes**")
        if df_cuadrillas.empty:
            st.info("Todavia no hay cuadrillas creadas.")
        else:
            st.dataframe(
                df_cuadrillas[
                    ["codigo", "nombre", "frentes_maestro", "uso_sugerido", "costo_dia", "activo"]
                ],
                width="stretch",
                hide_index=True,
            )

        st.divider()
        st.markdown("**Crear o editar una cuadrilla**")
        opciones_cuadrilla = {f"{c['codigo']} · {c['nombre']}": c["codigo"] for c in cuadrillas_actuales}
        elegida = st.selectbox(
            "Elige una cuadrilla para editarla, o crea una nueva",
            ["(nueva cuadrilla)"] + list(opciones_cuadrilla.keys()),
            key="select_cuadrilla_editor",
        )

        if elegida == "(nueva cuadrilla)":
            col_n1, col_n2, col_n3 = st.columns([1, 2, 1])
            with col_n1:
                codigo_nuevo = st.text_input("Codigo (ej. UBXXX)", key="codigo_cuadrilla_nueva")
            with col_n2:
                nombre_nuevo = st.text_input("Nombre", key="nombre_cuadrilla_nueva")
            with col_n3:
                frentes_nuevo = st.number_input(
                    "Frentes maestro", min_value=1, step=1, value=1, key="frentes_cuadrilla_nueva"
                )
            uso_nuevo = st.text_input("Uso sugerido (opcional)", key="uso_cuadrilla_nueva")
            st.caption(
                "Frentes maestro: 1 si el maestro EJECUTA (esta todo el dia en esa "
                "cuadrilla, ej. pintura, enchapes) -- el numero de frentes que rota "
                "si solo DIRIGE varias cuadrillas a la vez (ej. 4 en obra civil, "
                "muros, demoliciones); ahi su costo se divide entre esos frentes."
            )
            if not es_gestor():
                st.caption("Solo un administrador o cotizador puede crear cuadrillas.")
            if st.button("Crear cuadrilla", type="primary", disabled=not es_gestor()):
                codigo_norm = codigo_nuevo.strip().upper()
                if not codigo_norm or not nombre_nuevo:
                    st.error("Codigo y nombre son obligatorios.")
                elif codigo_norm in opciones_cuadrilla.values():
                    st.error("Ya existe una cuadrilla con ese codigo.")
                else:
                    crear_cuadrilla(sb, codigo_norm, nombre_nuevo, frentes_nuevo, uso_nuevo)
                    st.success(f"Cuadrilla {codigo_norm} creada. Ahora agrega sus cargos abajo.")
                    st.session_state.select_cuadrilla_editor = f"{codigo_norm} · {nombre_nuevo}"
                    st.rerun()
        else:
            codigo_activo = opciones_cuadrilla[elegida]
            cuadrilla_actual = next(c for c in cuadrillas_actuales if c["codigo"] == codigo_activo)

            col_e1, col_e2, col_e3, col_e4 = st.columns([2, 1, 2, 1])
            with col_e1:
                nombre_ui = st.text_input(
                    "Nombre", value=cuadrilla_actual["nombre"], key=f"nombre_{codigo_activo}"
                )
            with col_e2:
                frentes_ui = st.number_input(
                    "Frentes maestro", min_value=1, step=1,
                    value=int(cuadrilla_actual["frentes_maestro"]), key=f"frentes_{codigo_activo}",
                )
            with col_e3:
                uso_ui = st.text_input(
                    "Uso sugerido", value=cuadrilla_actual.get("uso_sugerido") or "",
                    key=f"uso_{codigo_activo}",
                )
            with col_e4:
                activo_ui = st.checkbox(
                    "Activa", value=bool(cuadrilla_actual["activo"]), key=f"activo_{codigo_activo}"
                )

            if not es_gestor():
                st.caption("Solo un administrador o cotizador puede editar cuadrillas.")
            if st.button("Guardar datos de la cuadrilla", disabled=not es_gestor()):
                actualizar_cuadrilla(sb, codigo_activo, nombre_ui, frentes_ui, uso_ui, activo_ui)
                st.success("Datos actualizados.")
                st.rerun()

            st.markdown(f"**Composicion -- costo/dia actual: {money(cuadrilla_actual['costo_dia'])}**")
            composicion = obtener_composicion_cuadrilla(sb, codigo_activo)
            cargos_disponibles = [
                c["cargo"] for c in obtener_cargos_personal(sb) if c["tipo"] == "operativo"
            ]
            df_comp = (
                pd.DataFrame(composicion)
                if composicion
                else pd.DataFrame(columns=["cargo", "cantidad", "tarifa_dia", "subtotal", "eliminar"])
            )
            df_comp_editado = st.data_editor(
                df_comp,
                column_config={
                    "cargo": st.column_config.SelectboxColumn("Cargo", options=cargos_disponibles),
                    "cantidad": st.column_config.NumberColumn(
                        "Cantidad", min_value=0.0, step=0.5, format="%.2f"
                    ),
                    "tarifa_dia": st.column_config.NumberColumn(
                        "Tarifa/dia", format="$ %d", disabled=True
                    ),
                    "subtotal": st.column_config.NumberColumn(
                        "Subtotal/dia", format="$ %d", disabled=True
                    ),
                    "eliminar": st.column_config.CheckboxColumn("Eliminar"),
                },
                hide_index=True,
                width="stretch",
                num_rows="dynamic",
                key=f"editor_composicion_{codigo_activo}",
            )
            st.caption(
                "Para agregar un cargo: agrega una fila, elige el cargo y la cantidad, "
                "y guarda -- la tarifa/subtotal se completa sola despues de guardar."
            )

            if not es_gestor():
                st.caption("Solo un administrador o cotizador puede editar la composicion de una cuadrilla.")
            if st.button("Guardar composicion", type="primary", disabled=not es_gestor()):
                guardar_composicion_cuadrilla(sb, codigo_activo, df_comp_editado)
                st.success("Composicion actualizada -- el costo/dia se recalculo solo.")
                st.rerun()

    with sub_mo:
        st.subheader("Cambio de tarifas de personal y recosteo de mano de obra")
        st.caption(
            "Edita la tarifa/dia de los cargos que cambiaron (ej. reajuste salarial anual) "
            "y guarda. Despues, recalcula el impacto: mano_obra de cada APU se recompone "
            "como (costo/dia de su cuadrilla) / (rendimiento del APU), igual que "
            "RECOSTEAR_MO en el Excel -- pero solo para los APUs que ya tienen cuadrilla "
            "y rendimiento asignados (Fase 2)."
        )
        import pandas as pd

        cargos_actuales = obtener_cargos_personal(sb)

        with st.expander("Crear un cargo/rol nuevo"):
            st.caption(
                "Para un oficio que no esta entre los cargos actuales (ej. un "
                "ayudante electrico). Una vez creado, queda disponible para "
                "agregarlo a la composicion de cualquier cuadrilla."
            )
            col_c1, col_c2, col_c3 = st.columns([2, 1, 1])
            with col_c1:
                cargo_nuevo_nombre = st.text_input("Nombre del cargo (ej. AYUDANTE ELECTRICO)", key="cargo_nuevo_nombre")
            with col_c2:
                cargo_nuevo_tarifa = st.number_input(
                    "Tarifa/dia ($)", min_value=0.0, step=1000.0, key="cargo_nuevo_tarifa"
                )
            with col_c3:
                cargo_nuevo_tipo = st.selectbox(
                    "Tipo", ["operativo", "supervision"], key="cargo_nuevo_tipo"
                )
            st.caption(
                "'operativo' = va en cuadrillas (mano de obra del APU). "
                "'supervision' = SISO/Residente/Director, se carga aparte en PARAMETROS, "
                "no en cuadrillas."
            )
            if not es_gestor():
                st.caption("Solo un administrador o cotizador puede crear cargos.")
            if st.button("Crear cargo", disabled=not es_gestor()):
                nombre_norm = cargo_nuevo_nombre.strip().upper()
                existentes = {c["cargo"] for c in cargos_actuales}
                if not nombre_norm:
                    st.error("El nombre del cargo es obligatorio.")
                elif nombre_norm in existentes:
                    st.error("Ya existe un cargo con ese nombre.")
                elif cargo_nuevo_tarifa <= 0:
                    st.error("La tarifa/dia debe ser mayor a 0.")
                else:
                    crear_cargo_personal(sb, nombre_norm, cargo_nuevo_tarifa, cargo_nuevo_tipo)
                    st.success(f"Cargo {nombre_norm} creado.")
                    st.rerun()

        df_cargos = pd.DataFrame(cargos_actuales)
        st.markdown("**1. Tarifas de personal**")
        df_cargos_editado = st.data_editor(
            df_cargos,
            column_config={
                "cargo": st.column_config.TextColumn(disabled=True),
                "tipo": st.column_config.TextColumn(disabled=True),
                "tarifa_dia": st.column_config.NumberColumn(format="$ %d"),
                "activo": st.column_config.CheckboxColumn(disabled=True),
            },
            hide_index=True,
            width="stretch",
            key="editor_cargos_personal",
        )

        if not es_gestor():
            st.caption("Solo un administrador o cotizador puede guardar cambios de tarifas.")
        if st.button("Guardar tarifas nuevas", disabled=not es_gestor()):
            cambios_tarifa = 0
            originales_cargo = {c["cargo"]: c for c in cargos_actuales}
            for _, fila in df_cargos_editado.iterrows():
                cargo = fila["cargo"]
                nueva_tarifa = float(fila["tarifa_dia"])
                if nueva_tarifa != float(originales_cargo[cargo]["tarifa_dia"]):
                    sb.table("cargos_personal").update({"tarifa_dia": nueva_tarifa}).eq(
                        "cargo", cargo
                    ).execute()
                    cambios_tarifa += 1
            if cambios_tarifa:
                st.success(f"{cambios_tarifa} tarifa(s) actualizada(s).")
            else:
                st.info("No hubo cambios en las tarifas.")

        st.divider()
        st.markdown("**2. Impacto en el catalogo (no guarda todavia)**")
        if st.button("Calcular impacto del recosteo"):
            st.session_state.impacto_recosteo = calcular_impacto_recosteo(sb)

        impacto = st.session_state.get("impacto_recosteo")
        if impacto is not None:
            if not impacto:
                st.warning(
                    "No hay APUs con cuadrilla_codigo y rendimiento_dia asignados todavia "
                    "-- el recalculo automatico no tiene sobre que aplicar. Asigna una "
                    "cuadrilla desde '6. Editar receta de APU', o corrige mano_obra a mano "
                    "en los APUs afectados."
                )
            else:
                df_impacto = pd.DataFrame(impacto)
                cambian = df_impacto[df_impacto["mano_obra_viejo"] != df_impacto["mano_obra_nuevo"]]
                st.dataframe(
                    df_impacto[
                        ["codigo", "categoria", "descripcion", "mano_obra_viejo", "mano_obra_nuevo", "var_pct"]
                    ],
                    width="stretch",
                    hide_index=True,
                )
                viejo_total = df_impacto["mano_obra_viejo"].sum()
                nuevo_total = df_impacto["mano_obra_nuevo"].sum()
                var_total = ((nuevo_total - viejo_total) / viejo_total * 100) if viejo_total else 0.0
                st.write(
                    f"**Total mano de obra -- viejo: {money(viejo_total)} · nuevo: {money(nuevo_total)} "
                    f"· variacion: {var_total:.1f}%**"
                )
                st.write(f"**{len(cambian)} APU(s) van a cambiar** de {len(df_impacto)} evaluados.")

                if not es_gestor():
                    st.caption("Solo un administrador o cotizador puede confirmar el recosteo.")
                if len(cambian) and st.button(
                    "Confirmar y guardar recosteo", type="primary", disabled=not es_gestor()
                ):
                    n_aplicados = aplicar_recosteo_mano_obra(sb, impacto)
                    st.session_state.impacto_recosteo = None
                    st.success(
                        f"{n_aplicados} APU(s) actualizados y registrados en la bitacora."
                    )
                    st.rerun()

    # --- Bitacora ---
    with sub_bitacora:
        st.subheader("Bitacora de cambios de precios")
        st.caption("Auditoria: quien cambio que precio (usuario y origen), y cuando.")
        import pandas as pd

        registros = obtener_bitacora(sb)
        if not registros:
            st.info("Todavia no hay cambios registrados en la bitacora.")
        else:
            df_bit = pd.DataFrame(registros)
            if "usuario" not in df_bit.columns:
                df_bit["usuario"] = None
            df_bit["usuario"] = df_bit["usuario"].fillna("(sin registrar)")
            col_f1, col_f2, col_f3 = st.columns(3)
            with col_f1:
                filtro_campo = st.multiselect(
                    "Campo", sorted(df_bit["campo"].unique()), default=list(df_bit["campo"].unique())
                )
            with col_f2:
                filtro_origen = st.multiselect(
                    "Origen", sorted(df_bit["origen"].dropna().unique()),
                    default=list(df_bit["origen"].dropna().unique()),
                )
            with col_f3:
                filtro_usuario = st.multiselect(
                    "Usuario", sorted(df_bit["usuario"].unique()), default=list(df_bit["usuario"].unique())
                )
            df_bit_filtrado = df_bit[
                df_bit["campo"].isin(filtro_campo)
                & df_bit["origen"].isin(filtro_origen)
                & df_bit["usuario"].isin(filtro_usuario)
            ]
            columnas_bit = [
                "creado_en", "usuario", "apu_codigo", "insumo_codigo",
                "campo", "valor_anterior", "valor_nuevo", "origen",
            ]
            columnas_bit = [c for c in columnas_bit if c in df_bit_filtrado.columns]
            st.dataframe(
                df_bit_filtrado[columnas_bit],
                width="stretch",
                hide_index=True,
            )

# ---------------------------------------------------------------------
# TAB 6: editor de receta de APU (insumos, cantidades, rendimiento,
# equipo, transporte) -- lo que antes solo se podia cambiar en el Excel
# maestro.
# ---------------------------------------------------------------------
with tab_editor_apu:
    st.caption(
        "Edita la 'receta' de un APU: que insumos usa y en que cantidad, el rendimiento "
        "por dia, y los valores de equipo/transporte. Al guardar, 'materiales' se "
        "recalcula solo desde los insumos (y las lineas fijas); 'mano_obra' se recalcula "
        "solo si el APU ya tiene una cuadrilla asignada -- si no, se edita directamente."
    )

    col_crear1, col_crear2, col_crear3 = st.columns(3)
    with col_crear1:
        with st.expander("Crear un insumo nuevo"):
            st.caption("Para cuando el material que necesitas todavia no esta en el catalogo de insumos.")
            nuevo_ins_prefijo = st.text_input(
                "Prefijo del codigo (ej. MAT)", value="MAT", key="nuevo_insumo_prefijo",
            ).strip().upper() or "MAT"
            codigo_sugerido_ins = siguiente_codigo_insumo(sb, nuevo_ins_prefijo)
            nuevo_ins_codigo = st.text_input(
                "Codigo (autocompletado con el siguiente consecutivo -- se puede ajustar)",
                value=codigo_sugerido_ins,
                key=f"nuevo_insumo_codigo_{nuevo_ins_prefijo}",
            )
            st.caption(f"Siguiente consecutivo libre para '{nuevo_ins_prefijo}': {codigo_sugerido_ins}")
            nuevo_ins_desc = st.text_input("Descripcion", key="nuevo_insumo_desc")
            col_a, col_b = st.columns(2)
            with col_a:
                nuevo_ins_unidad = st.text_input("Unidad (ej. Un, Kg, m3)", key="nuevo_insumo_unidad")
            with col_b:
                nuevo_ins_precio = st.number_input("Precio", min_value=0.0, step=100.0, key="nuevo_insumo_precio")
            nuevo_ins_proveedor = st.text_input("Proveedor (opcional)", key="nuevo_insumo_proveedor")
            if not es_gestor():
                st.caption("Solo un administrador o cotizador puede crear insumos.")
            if st.button("Crear insumo", key="btn_crear_insumo", disabled=not es_gestor()):
                if not nuevo_ins_codigo or not nuevo_ins_desc or not nuevo_ins_unidad:
                    st.error("Codigo, descripcion y unidad son obligatorios.")
                else:
                    try:
                        codigo_creado = crear_insumo_nuevo(
                            sb, nuevo_ins_codigo, nuevo_ins_desc, nuevo_ins_unidad,
                            nuevo_ins_precio, nuevo_ins_proveedor or None,
                        )
                        st.success(f"Insumo {codigo_creado} creado. Ya lo puedes buscar y agregar a una receta abajo.")
                    except Exception as e:
                        st.error(str(e))

    with col_crear2:
        with st.expander("Crear un APU nuevo"):
            st.caption(
                "Arranca vacio (queda marcado como 'candidato' hasta que le agregues "
                "insumos y guardes la receta)."
            )
            categorias_apu_nuevo = obtener_categorias_apu(sb)
            nueva_apu_categoria = st.selectbox(
                "Capitulo", categorias_apu_nuevo + ["(otro capitulo nuevo)"], key="nuevo_apu_categoria",
            )
            if nueva_apu_categoria == "(otro capitulo nuevo)":
                nueva_apu_categoria = st.text_input(
                    "Prefijo del capitulo nuevo (ej. XYZ)", key="nuevo_apu_categoria_custom",
                ).strip().upper()
            codigo_sugerido_apu = siguiente_codigo_apu(sb, nueva_apu_categoria) if nueva_apu_categoria else ""
            nuevo_apu_codigo = st.text_input(
                "Codigo (autocompletado con el siguiente consecutivo del capitulo -- se puede ajustar)",
                value=codigo_sugerido_apu,
                key=f"nuevo_apu_codigo_{nueva_apu_categoria}",
            )
            if nueva_apu_categoria:
                st.caption(f"Siguiente consecutivo libre para '{nueva_apu_categoria}': {codigo_sugerido_apu}")
            nuevo_apu_desc = st.text_input("Descripcion", key="nuevo_apu_desc")
            nuevo_apu_unidad = st.text_input("Unidad (ej. m2, Un, Glb)", key="nuevo_apu_unidad")
            if not es_gestor():
                st.caption("Solo un administrador o cotizador puede crear APUs nuevos.")
            if st.button("Crear APU", key="btn_crear_apu", disabled=not es_gestor()):
                if not nuevo_apu_codigo or not nuevo_apu_desc or not nuevo_apu_unidad:
                    st.error("Codigo, descripcion y unidad son obligatorios.")
                else:
                    try:
                        codigo_creado = crear_apu_nuevo(sb, nuevo_apu_codigo, nuevo_apu_desc, nuevo_apu_unidad)
                        st.session_state.editor_apu_codigo = codigo_creado
                        st.success(f"APU {codigo_creado} creado. Ya lo puedes editar abajo.")
                        st.rerun()
                    except Exception as e:
                        st.error(str(e))

    with col_crear3:
        with st.expander("Duplicar un APU existente"):
            st.caption(
                "Copia la cuadrilla, rendimiento, equipo, transporte y todos los "
                "insumos/lineas fijas de un APU que ya existe a uno nuevo -- para "
                "partir de uno parecido y solo ajustarle un par de cosas."
            )
            categorias_dup = obtener_categorias_apu(sb)
            col_dup_cat, col_dup_txt = st.columns(2)
            with col_dup_cat:
                categoria_dup = st.selectbox(
                    "Capitulo (opcional)", ["(todos)"] + categorias_dup, key="dup_apu_categoria"
                )
            with col_dup_txt:
                busq_dup = st.text_input("Buscar por codigo o descripcion", key="dup_apu_busqueda")
            resultados_dup = buscar_apus(
                sb, texto=busq_dup or None,
                categoria=None if categoria_dup == "(todos)" else categoria_dup,
            )
            if not resultados_dup:
                st.info("No encontre APUs con esos filtros.")
            else:
                opciones_dup = {
                    f"{a['codigo']} · {a['descripcion'][:70]} ({a['unidad']})": a["codigo"]
                    for a in resultados_dup
                }
                elegido_dup_etiqueta = st.selectbox(
                    "APU de origen (el que quieres copiar)", list(opciones_dup.keys()), key="dup_apu_origen_select",
                )
                codigo_origen_dup = opciones_dup[elegido_dup_etiqueta]
                apu_origen_dup = obtener_apu_detalle(sb, codigo_origen_dup)
                codigo_sugerido_dup = siguiente_codigo_apu(sb, apu_origen_dup["categoria"])
                nuevo_codigo_dup = st.text_input(
                    "Codigo del APU nuevo (autocompletado con el siguiente consecutivo -- se puede ajustar)",
                    value=codigo_sugerido_dup,
                    key=f"dup_apu_codigo_{codigo_origen_dup}",
                )
                st.caption(f"Siguiente consecutivo libre para '{apu_origen_dup['categoria']}': {codigo_sugerido_dup}")
                nueva_desc_dup = st.text_input(
                    "Descripcion del APU nuevo", value=apu_origen_dup["descripcion"], key=f"dup_apu_desc_{codigo_origen_dup}",
                )
                nueva_unidad_dup = st.text_input(
                    "Unidad del APU nuevo", value=apu_origen_dup["unidad"], key=f"dup_apu_unidad_{codigo_origen_dup}",
                )
                if not es_gestor():
                    st.caption("Solo un administrador o cotizador puede duplicar un APU.")
                if st.button("Duplicar APU", key="btn_duplicar_apu", disabled=not es_gestor()):
                    if not nuevo_codigo_dup or not nueva_desc_dup or not nueva_unidad_dup:
                        st.error("Codigo, descripcion y unidad son obligatorios.")
                    else:
                        try:
                            codigo_creado_dup = duplicar_apu(
                                sb, codigo_origen_dup, nuevo_codigo_dup, nueva_desc_dup, nueva_unidad_dup,
                            )
                            st.session_state.editor_apu_codigo = codigo_creado_dup
                            st.success(
                                f"APU {codigo_creado_dup} creado a partir de {codigo_origen_dup}, con su misma "
                                "cuadrilla, rendimiento, equipo, transporte e insumos. Ya lo puedes ajustar abajo."
                            )
                            st.rerun()
                        except Exception as e:
                            st.error(str(e))

    st.divider()
    st.caption(
        "Son ~2.300 APUs -- filtra primero por capitulo (opcional) y despues elige de la "
        "lista desplegable, o escribe parte del codigo/descripcion para acortarla."
    )
    col_cat, col_txt = st.columns([1, 2])
    with col_cat:
        categorias_disponibles = obtener_categorias_apu(sb)
        categoria_elegida = st.selectbox(
            "Capitulo (opcional)", ["(todos)"] + categorias_disponibles, key="categoria_apu_editor"
        )
    with col_txt:
        busq_apu = st.text_input("Buscar por codigo o descripcion (opcional)", key="busq_apu_editor")

    resultados_apu = buscar_apus(
        sb,
        texto=busq_apu or None,
        categoria=None if categoria_elegida == "(todos)" else categoria_elegida,
    )
    if not resultados_apu:
        st.info("No encontre APUs con esos filtros.")
    else:
        st.caption(f"{len(resultados_apu)} APU(s) -- si son muchos, agrega un capitulo o texto para acortar la lista.")
        opciones_apu = {
            f"{a['codigo']} · {a['descripcion'][:70]} ({a['unidad']})": a["codigo"] for a in resultados_apu
        }
        elegido_apu = st.selectbox(
            "Elige el APU (puedes escribir dentro de este campo para buscar en la lista)",
            list(opciones_apu.keys()),
            key="select_apu_editor",
        )
        if st.button("Abrir esta receta"):
            st.session_state.editor_apu_codigo = opciones_apu[elegido_apu]
            st.rerun()

    apu_codigo_activo = st.session_state.get("editor_apu_codigo")
    if apu_codigo_activo:
        apu = obtener_apu_detalle(sb, apu_codigo_activo)
        if apu is None:
            st.warning("Ese APU ya no existe en el catalogo.")
        else:
            st.divider()
            st.subheader(f"{apu['codigo']} · {apu['descripcion']}")
            st.caption(f"Unidad: {apu['unidad']} · Categoria: {apu['categoria']}")

            with st.expander("Editar descripcion / unidad de este APU"):
                col_nom1, col_nom2 = st.columns([2, 1])
                with col_nom1:
                    descripcion_ui = st.text_input(
                        "Descripcion", value=apu["descripcion"], key=f"descripcion_editor_{apu_codigo_activo}",
                    )
                with col_nom2:
                    unidad_ui = st.text_input(
                        "Unidad", value=apu["unidad"], key=f"unidad_editor_{apu_codigo_activo}",
                    )
                if not es_gestor():
                    st.caption("Solo un administrador o cotizador puede renombrar un APU.")
                if st.button(
                    "Guardar descripcion / unidad", key=f"btn_guardar_nombre_{apu_codigo_activo}",
                    disabled=not es_gestor(),
                ):
                    if not descripcion_ui.strip() or not unidad_ui.strip():
                        st.error("Descripcion y unidad no pueden quedar vacias.")
                    else:
                        sb.table("catalogo_apu").update(
                            {"descripcion": descripcion_ui.strip(), "unidad": unidad_ui.strip()}
                        ).eq("codigo", apu_codigo_activo).execute()
                        st.success("Descripcion/unidad actualizadas.")
                        st.rerun()

            import pandas as pd

            st.markdown("---")
            st.markdown("### I. EQUIPO Y HERRAMIENTAS")
            equipo_items = obtener_equipo_de_apu(sb, apu_codigo_activo)
            df_equipo_editor = None
            if equipo_items:
                df_equipo_base = pd.DataFrame(equipo_items)
                df_equipo_editor = st.data_editor(
                    df_equipo_base,
                    column_config={
                        "id": None,
                        "descripcion": st.column_config.TextColumn("Descripcion"),
                        "unidad": st.column_config.TextColumn("Unidad"),
                        "cantidad": st.column_config.NumberColumn("Cant/Rend", min_value=0.0, step=0.0001, format="%.4f"),
                        "precio_unitario": st.column_config.NumberColumn("Precio unitario", format="$ %d"),
                        "subtotal": st.column_config.NumberColumn("VR parcial", format="$ %d", disabled=True),
                        "eliminar": st.column_config.CheckboxColumn("Eliminar"),
                    },
                    hide_index=True,
                    width="stretch",
                    key=f"editor_equipo_{apu_codigo_activo}",
                )
            else:
                st.caption("Esta receta no tiene equipo desglosado por lineas todavia.")

            with st.expander("Agregar una linea de equipo/herramienta"):
                desc_equipo = st.text_input("Descripcion", key="desc_equipo_nueva")
                col_q1, col_q2, col_q3 = st.columns(3)
                with col_q1:
                    unidad_equipo = st.text_input("Unidad", key="unidad_equipo_nueva")
                with col_q2:
                    cantidad_equipo = st.number_input(
                        "Cantidad/Rendimiento", min_value=0.0, step=0.0001, format="%.4f", key="cantidad_equipo_nueva"
                    )
                with col_q3:
                    precio_equipo = st.number_input(
                        "Precio unitario", min_value=0.0, step=100.0, key="precio_equipo_nueva"
                    )
                if not es_gestor():
                    st.caption("Solo un administrador o cotizador puede agregar lineas de equipo.")
                if st.button("Agregar esta linea de equipo", key="btn_agregar_equipo", disabled=not es_gestor()):
                    if not desc_equipo:
                        st.error("La descripcion es obligatoria.")
                    else:
                        sb.table("apu_equipo_items").insert(
                            {
                                "apu_codigo": apu_codigo_activo,
                                "descripcion": desc_equipo,
                                "unidad": unidad_equipo,
                                "cantidad": float(cantidad_equipo),
                                "precio_unitario": float(precio_equipo),
                            }
                        ).execute()
                        st.success("Linea de equipo agregada.")
                        st.rerun()

            if not equipo_items:
                equipo_manual_ui = st.number_input(
                    "Valor equipo total ($, mientras no lo detalles por lineas)",
                    min_value=0.0, step=100.0, value=float(apu.get("equipo") or 0), key="equipo_editor",
                )
            else:
                equipo_manual_ui = float(apu.get("equipo") or 0)
                subtotal_equipo_actual = round(
                    sum(float(f["cantidad"] or 0) * float(f["precio_unitario"] or 0) for f in equipo_items), 0
                )
                st.caption(f"Subtotal equipo = {money(subtotal_equipo_actual)}")

            st.markdown("### II. MATERIALES")
            st.caption(
                "Los insumos que componen esta actividad (equivalente a la seccion "
                "'II. MATERIALES' del APU en Excel)."
            )
            st.markdown("**Insumos de la receta**")
            insumos_apu = obtener_insumos_de_apu(sb, apu_codigo_activo)
            df_insumos_editor = None
            if insumos_apu:
                df_insumos_base = pd.DataFrame(insumos_apu)
                df_insumos_editor = st.data_editor(
                    df_insumos_base,
                    column_config={
                        "insumo_codigo": st.column_config.TextColumn("Codigo", disabled=True),
                        "descripcion": st.column_config.TextColumn("Descripcion", disabled=True),
                        "unidad": st.column_config.TextColumn("Unidad", disabled=True),
                        "precio": st.column_config.NumberColumn("Precio insumo", format="$ %d", disabled=True),
                        "cantidad": st.column_config.NumberColumn("Cantidad", min_value=0.0, step=0.0001, format="%.4f"),
                        "subtotal": st.column_config.NumberColumn("Subtotal", format="$ %d", disabled=True),
                        "eliminar": st.column_config.CheckboxColumn("Eliminar"),
                    },
                    hide_index=True,
                    width="stretch",
                    key=f"editor_insumos_{apu_codigo_activo}",
                )
            else:
                st.caption("Esta receta no tiene insumos enlazados todavia.")

            with st.expander("Agregar un insumo a esta receta"):
                busq_ins = st.text_input("Buscar insumo por codigo o descripcion", key="busq_insumo_editor")
                resultados_ins = buscar_insumos_catalogo(sb, busq_ins) if busq_ins else []
                if busq_ins and not resultados_ins:
                    st.info("No encontre insumos con ese texto.")
                elif resultados_ins:
                    opciones_ins = {
                        f"{i['codigo']} · {i['descripcion'][:70]} ({i['unidad']}, {money(i['precio'])})": i["codigo"]
                        for i in resultados_ins
                    }
                    elegido_ins = st.selectbox("Resultados", list(opciones_ins.keys()), key="select_insumo_editor")
                    cantidad_nueva = st.number_input(
                        "Cantidad por unidad de APU", min_value=0.0001, step=0.0001, format="%.4f",
                        value=1.0, key="cantidad_insumo_nuevo",
                    )
                    if not es_gestor():
                        st.caption("Solo un administrador o cotizador puede agregar insumos a una receta.")
                    if st.button("Agregar este insumo a la receta", disabled=not es_gestor()):
                        sb.table("apu_insumos").upsert(
                            {
                                "apu_codigo": apu_codigo_activo,
                                "insumo_codigo": opciones_ins[elegido_ins],
                                "cantidad": float(cantidad_nueva),
                            }
                        ).execute()
                        st.success("Insumo agregado. Ajusta la cantidad abajo si hace falta y guarda los cambios.")
                        st.rerun()

            st.markdown("**Lineas de materiales fijas** (no enlazadas a un insumo -- ej. concreto)")
            fijos_apu = obtener_fijos_de_apu(sb, apu_codigo_activo)
            df_fijos_editor = None
            if fijos_apu:
                df_fijos_base = pd.DataFrame(fijos_apu)
                df_fijos_editor = st.data_editor(
                    df_fijos_base,
                    column_config={
                        "id": None,
                        "descripcion": st.column_config.TextColumn("Descripcion"),
                        "unidad": st.column_config.TextColumn("Unidad"),
                        "cantidad": st.column_config.NumberColumn("Cantidad", min_value=0.0, step=0.0001, format="%.4f"),
                        "precio_unitario": st.column_config.NumberColumn("Precio unitario", format="$ %d"),
                        "subtotal": st.column_config.NumberColumn("Subtotal", format="$ %d", disabled=True),
                        "eliminar": st.column_config.CheckboxColumn("Eliminar"),
                    },
                    hide_index=True,
                    width="stretch",
                    key=f"editor_fijos_{apu_codigo_activo}",
                )
            else:
                st.caption("Esta receta no tiene lineas fijas.")

            with st.expander("Agregar una linea fija (sin enlazar a un insumo)"):
                desc_fija = st.text_input("Descripcion", key="desc_fija_nueva")
                col_f1, col_f2, col_f3 = st.columns(3)
                with col_f1:
                    unidad_fija = st.text_input("Unidad", key="unidad_fija_nueva")
                with col_f2:
                    cantidad_fija = st.number_input(
                        "Cantidad", min_value=0.0, step=0.0001, format="%.4f", key="cantidad_fija_nueva"
                    )
                with col_f3:
                    precio_fija = st.number_input(
                        "Precio unitario", min_value=0.0, step=100.0, key="precio_fija_nueva"
                    )
                if not es_gestor():
                    st.caption("Solo un administrador o cotizador puede agregar lineas fijas a una receta.")
                if st.button("Agregar esta linea fija", disabled=not es_gestor()):
                    if not desc_fija:
                        st.error("La descripcion es obligatoria.")
                    else:
                        sb.table("apu_materiales_fijos").insert(
                            {
                                "apu_codigo": apu_codigo_activo,
                                "descripcion": desc_fija,
                                "unidad": unidad_fija,
                                "cantidad": float(cantidad_fija),
                                "precio_unitario": float(precio_fija),
                            }
                        ).execute()
                        st.success("Linea agregada.")
                        st.rerun()

            st.markdown("---")
            st.markdown("### III. TRANSPORTES")
            transporte_items = obtener_transporte_de_apu(sb, apu_codigo_activo)
            df_transporte_editor = None
            if transporte_items:
                df_transporte_base = pd.DataFrame(transporte_items)
                df_transporte_editor = st.data_editor(
                    df_transporte_base,
                    column_config={
                        "id": None,
                        "descripcion": st.column_config.TextColumn("Descripcion"),
                        "unidad": st.column_config.TextColumn("Unidad"),
                        "cantidad": st.column_config.NumberColumn("Cant/Rend", min_value=0.0, step=0.0001, format="%.4f"),
                        "precio_unitario": st.column_config.NumberColumn("Precio unitario", format="$ %d"),
                        "subtotal": st.column_config.NumberColumn("VR parcial", format="$ %d", disabled=True),
                        "eliminar": st.column_config.CheckboxColumn("Eliminar"),
                    },
                    hide_index=True,
                    width="stretch",
                    key=f"editor_transporte_{apu_codigo_activo}",
                )
            else:
                st.caption("Esta receta no tiene transporte desglosado por lineas todavia.")

            with st.expander("Agregar una linea de transporte"):
                desc_transp = st.text_input("Descripcion", key="desc_transporte_nueva")
                col_t1, col_t2, col_t3 = st.columns(3)
                with col_t1:
                    unidad_transp = st.text_input("Unidad", key="unidad_transporte_nueva")
                with col_t2:
                    cantidad_transp = st.number_input(
                        "Cantidad/Rendimiento", min_value=0.0, step=0.0001, format="%.4f", key="cantidad_transporte_nueva"
                    )
                with col_t3:
                    precio_transp = st.number_input(
                        "Precio unitario", min_value=0.0, step=100.0, key="precio_transporte_nueva"
                    )
                if not es_gestor():
                    st.caption("Solo un administrador o cotizador puede agregar lineas de transporte.")
                if st.button("Agregar esta linea de transporte", key="btn_agregar_transporte", disabled=not es_gestor()):
                    if not desc_transp:
                        st.error("La descripcion es obligatoria.")
                    else:
                        sb.table("apu_transporte_items").insert(
                            {
                                "apu_codigo": apu_codigo_activo,
                                "descripcion": desc_transp,
                                "unidad": unidad_transp,
                                "cantidad": float(cantidad_transp),
                                "precio_unitario": float(precio_transp),
                            }
                        ).execute()
                        st.success("Linea de transporte agregada.")
                        st.rerun()

            if not transporte_items:
                transporte_manual_ui = st.number_input(
                    "Valor transporte total ($, mientras no lo detalles por lineas)",
                    min_value=0.0, step=100.0, value=float(apu.get("transporte") or 0), key="transporte_editor",
                )
            else:
                transporte_manual_ui = float(apu.get("transporte") or 0)
                subtotal_transporte_actual = round(
                    sum(float(f["cantidad"] or 0) * float(f["precio_unitario"] or 0) for f in transporte_items), 0
                )
                st.caption(f"Subtotal transporte = {money(subtotal_transporte_actual)}")

            st.markdown("---")
            st.markdown("### IV. MANO DE OBRA")
            cuadrillas_lista = obtener_cuadrillas(sb)
            opciones_cuad_apu = {"(sin cuadrilla -- mano de obra manual)": None}
            opciones_cuad_apu.update(
                {f"{c['codigo']} · {c['nombre']} ({money(c['costo_dia'])}/dia)": c["codigo"] for c in cuadrillas_lista}
            )
            etiquetas_cuad_apu = list(opciones_cuad_apu.keys())
            actual_codigo = apu.get("cuadrilla_codigo")
            indice_actual = 0
            for i, (etiqueta, codigo) in enumerate(opciones_cuad_apu.items()):
                if codigo == actual_codigo:
                    indice_actual = i
                    break
            cuadrilla_elegida_etiqueta = st.selectbox(
                "Cuadrilla (equivalente a la hoja CUADRILLAS del Excel)",
                etiquetas_cuad_apu,
                index=indice_actual,
                key=f"select_cuadrilla_apu_{apu_codigo_activo}",
            )
            cuadrilla_elegida_codigo = opciones_cuad_apu[cuadrilla_elegida_etiqueta]
            if cuadrilla_elegida_codigo != actual_codigo:
                st.caption(
                    "Cambio sin guardar todavia. Si el APU ya tiene rendimiento/dia, la "
                    "mano de obra se recalcula sola al guardar (costo/dia de la cuadrilla "
                    "/ rendimiento) -- igual que ASIGNAR_CUADRILLAS.txt / "
                    "CAMBIAR_CUADRILLA.txt del Excel maestro."
                )
                if not es_gestor():
                    st.caption("Solo un administrador o cotizador puede reasignar la cuadrilla de un APU.")
                if st.button(
                    "Guardar cuadrilla asignada",
                    key=f"guardar_cuadrilla_{apu_codigo_activo}",
                    disabled=not es_gestor(),
                ):
                    asignar_cuadrilla_apu(sb, apu_codigo_activo, apu, cuadrilla_elegida_codigo)
                    st.success("Cuadrilla actualizada.")
                    st.rerun()
            elif actual_codigo:
                st.caption(
                    "La mano de obra se recalcula sola a partir del rendimiento "
                    "(no se edita directo mientras tenga cuadrilla asignada)."
                )

            rendimiento_ui = st.number_input(
                "Rendimiento (unidades/dia que hace esa cuadrilla)", min_value=0.0, step=0.01, format="%.4f",
                value=float(apu.get("rendimiento_dia") or 0), key="rendimiento_editor",
            )

            if apu.get("cuadrilla_codigo"):
                st.caption(f"Mano de obra actual (calculada = costo/dia de la cuadrilla ÷ rendimiento): {money(apu.get('mano_obra'))}")
                mano_obra_ui = float(apu.get("mano_obra") or 0)
            else:
                mano_obra_ui = st.number_input(
                    "Mano de obra ($, sin cuadrilla asignada -- se edita a mano)",
                    min_value=0.0, step=100.0, value=float(apu.get("mano_obra") or 0),
                    key="mano_obra_editor",
                )

            st.markdown("---")

            def _subtotal_editable(df, col_precio):
                if df is None:
                    return 0.0
                restante = df.loc[~df["eliminar"]]
                return float((restante["cantidad"] * restante[col_precio]).sum())

            equipo_preview = _subtotal_editable(df_equipo_editor, "precio_unitario") if df_equipo_editor is not None else equipo_manual_ui
            transporte_preview = (
                _subtotal_editable(df_transporte_editor, "precio_unitario") if df_transporte_editor is not None else transporte_manual_ui
            )
            materiales_preview = _subtotal_editable(df_insumos_editor, "precio") + _subtotal_editable(df_fijos_editor, "precio_unitario")
            total_preview = equipo_preview + materiales_preview + transporte_preview + mano_obra_ui + float(apu.get("personal_supervision") or 0)
            st.write(
                f"**Vista previa -- Equipo: {money(equipo_preview)} · Materiales: {money(materiales_preview)} · "
                f"Transporte: {money(transporte_preview)} · Total APU: {money(total_preview)}**"
            )
            st.caption("El total real se recalcula al guardar (por redondeo puede variar unos pesos).")

            if not es_gestor():
                st.caption("Solo un administrador o cotizador puede guardar cambios de la receta de un APU.")
            if st.button("Guardar cambios de la receta", type="primary", disabled=not es_gestor()):
                guardar_receta_apu(
                    sb,
                    apu_codigo_activo,
                    apu,
                    df_insumos_editor if df_insumos_editor is not None else pd.DataFrame(columns=["insumo_codigo", "cantidad", "eliminar"]),
                    df_fijos_editor if df_fijos_editor is not None else pd.DataFrame(columns=["id", "cantidad", "precio_unitario", "eliminar"]),
                    df_equipo_editor if df_equipo_editor is not None else pd.DataFrame(columns=["id", "cantidad", "precio_unitario", "eliminar"]),
                    df_transporte_editor if df_transporte_editor is not None else pd.DataFrame(columns=["id", "cantidad", "precio_unitario", "eliminar"]),
                    equipo_manual_ui,
                    transporte_manual_ui,
                    rendimiento_ui,
                    mano_obra_ui,
                )
                st.success(f"Receta de {apu_codigo_activo} actualizada y registrada en la bitacora.")
                st.rerun()
